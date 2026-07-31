"""Tworzy przykładowy szablon operatu (szablony/operat_wzor.docx) wraz z opisem pól.

To tylko rusztowanie do testów — docelowy szablon brat robi sam w Wordzie,
wklejając swoje formatki i wstawiając w nich tagi {{ }}. Uruchomienie:

    python narzedzia/utworz_wzor_szablonu.py
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from app.config import SZABLONY  # noqa: E402


def akapit(dokument, tekst, rozmiar=11, pogrubienie=False, wyrownanie=None, odstep_przed=0):
    element = dokument.add_paragraph()
    element.paragraph_format.space_before = Pt(odstep_przed)
    bieg = element.add_run(tekst)
    bieg.bold = pogrubienie
    bieg.font.size = Pt(rozmiar)
    if wyrownanie is not None:
        element.alignment = wyrownanie
    return element


def zbuduj() -> Path:
    dokument = Document()
    styl = dokument.styles["Normal"]
    styl.font.name = "Times New Roman"
    styl.font.size = Pt(11)

    SRODEK = WD_ALIGN_PARAGRAPH.CENTER
    PRAWA = WD_ALIGN_PARAGRAPH.RIGHT
    JUSTUJ = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- strona tytułowa ---
    akapit(dokument, "{{ firma_nazwa }}", 10)
    akapit(dokument, "{{ firma_adres }}", 10)
    akapit(dokument, "tel. {{ firma_telefon }} · {{ firma_email }}", 10)
    akapit(dokument, "{{ miejscowosc }}, dnia {{ data_dokumentu }}", 10, wyrownanie=PRAWA,
           odstep_przed=6)

    akapit(dokument, "OPERAT TECHNICZNY", 20, pogrubienie=True, wyrownanie=SRODEK, odstep_przed=48)
    akapit(dokument, "{{ cel_opracowania }}", 13, wyrownanie=SRODEK, odstep_przed=8)

    akapit(dokument, "Nr roboty (KERG): {{ nr_roboty }}", 12, pogrubienie=True,
           wyrownanie=SRODEK, odstep_przed=36)
    akapit(dokument, "Nr operatu: {{ nr_operatu }}", 12, wyrownanie=SRODEK)

    akapit(dokument, "Województwo: {{ polozenie_wojewodztwo }}    Powiat: {{ polozenie_powiat }}",
           12, wyrownanie=SRODEK, odstep_przed=24)
    akapit(dokument, "Jednostka ewidencyjna: {{ polozenie_gmina_teryt }} {{ polozenie_gmina }}",
           12, wyrownanie=SRODEK)
    akapit(dokument, "Obręb: {{ polozenie_obreb_numer }} {{ polozenie_obreb }} "
                     "({{ polozenie_obreb_teryt }})", 12, wyrownanie=SRODEK)
    akapit(dokument, "Działki nr: {{ dzialki }}", 12, wyrownanie=SRODEK)

    akapit(dokument, "Zamawiający: {{ zamawiajacy }}", 11, wyrownanie=SRODEK, odstep_przed=36)
    akapit(dokument, "{{ zamawiajacy_adres }}", 11, wyrownanie=SRODEK)

    akapit(dokument, "Wykonawca prac geodezyjnych:", 11, wyrownanie=SRODEK, odstep_przed=48)
    akapit(dokument, "{{ geodeta_imie_nazwisko }}", 11, pogrubienie=True, wyrownanie=SRODEK)
    akapit(dokument, "uprawnienia zawodowe nr {{ uprawnienia_nr }}", 11, wyrownanie=SRODEK)

    dokument.add_page_break()

    # --- sprawozdanie techniczne ---
    akapit(dokument, "SPRAWOZDANIE TECHNICZNE", 14, pogrubienie=True, wyrownanie=SRODEK)

    akapit(dokument, "1. Cel i zakres opracowania", 12, pogrubienie=True, odstep_przed=14)
    akapit(dokument, "{{ cel_opracowania }} dla działek nr {{ dzialki }}. "
                     "Położenie: {{ polozenie }}.", wyrownanie=JUSTUJ)
    akapit(dokument, "{{ zakres_opis }}", wyrownanie=JUSTUJ)

    akapit(dokument, "2. Zgłoszenie pracy geodezyjnej", 12, pogrubienie=True, odstep_przed=14)
    akapit(dokument, "Praca została zgłoszona w {{ osrodek }} w dniu {{ data_zgloszenia }}, "
                     "identyfikator zgłoszenia {{ nr_roboty }}.", wyrownanie=JUSTUJ)

    akapit(dokument, "3. Materiały wykorzystane do opracowania", 12, pogrubienie=True,
           odstep_przed=14)
    akapit(dokument, "{{ materialy }}", wyrownanie=JUSTUJ)

    akapit(dokument, "4. Zastosowana technologia i sprzęt", 12, pogrubienie=True, odstep_przed=14)
    akapit(dokument, "Pomiar wykonano metodą {{ metoda_pomiaru }} w układzie współrzędnych "
                     "{{ uklad_wspolrzednych }}, układ wysokościowy {{ uklad_wysokosciowy }}.",
           wyrownanie=JUSTUJ)
    akapit(dokument, "Użyty sprzęt: {{ sprzet }}.", wyrownanie=JUSTUJ)
    akapit(dokument, "Prace polowe wykonano w dniach {{ data_prac_od }} – {{ data_prac_do }}.",
           wyrownanie=JUSTUJ)

    akapit(dokument, "5. Wykaz współrzędnych punktów", 12, pogrubienie=True, odstep_przed=14)

    # Uwaga: wiersz zawierający {%tr ... %} znika przy generowaniu w całości,
    # więc znaczniki pętli muszą stać w osobnych wierszach nad i pod wierszem z danymi.
    tabela = dokument.add_table(rows=4, cols=5)
    tabela.style = "Table Grid"
    for indeks, naglowek in enumerate(["Nr punktu", "X [m]", "Y [m]", "H [m]", "Uwagi"]):
        komorka = tabela.rows[0].cells[indeks]
        komorka.text = ""
        bieg = komorka.paragraphs[0].add_run(naglowek)
        bieg.bold = True
        bieg.font.size = Pt(10)

    tabela.rows[1].cells[0].text = "{%tr for punkt in punkty %}"
    wiersz = tabela.rows[2].cells
    wiersz[0].text = "{{ punkt.numer }}"
    wiersz[1].text = "{{ punkt.x }}"
    wiersz[2].text = "{{ punkt.y }}"
    wiersz[3].text = "{{ punkt.h }}"
    wiersz[4].text = "{{ punkt.uwagi }}"
    tabela.rows[3].cells[0].text = "{%tr endfor %}"
    for kolumna, szerokosc in zip(tabela.columns, [3, 3.2, 3.2, 2.4, 4]):
        for komorka in kolumna.cells:
            komorka.width = Cm(szerokosc)

    # {%p if %} kasuje cały akapit ze znacznikiem — dlatego warunek i treść
    # muszą być w osobnych akapitach.
    akapit(dokument, "{%p if uwagi %}", 12, pogrubienie=True, odstep_przed=14)
    akapit(dokument, "6. Uwagi", 12, pogrubienie=True)
    akapit(dokument, "{{ uwagi }}", wyrownanie=JUSTUJ)
    akapit(dokument, "{%p endif %}")

    akapit(dokument, "{{ geodeta_imie_nazwisko }}", 11, wyrownanie=PRAWA, odstep_przed=48)
    akapit(dokument, "uprawnienia nr {{ uprawnienia_nr }}", 10, wyrownanie=PRAWA)

    plik = SZABLONY / "operat_wzor.docx"
    dokument.save(plik)
    return plik


OPIS_POL = {
    "nazwa": "Operat techniczny (wzór testowy)",
    "opis": "Strona tytułowa + sprawozdanie techniczne + wykaz współrzędnych. "
            "Do podmiany na własny szablon.",
    "wzor_nazwy": "Operat_{nr_roboty}_{obreb}",
    "licznik": "operat",
    "pola": [
        {"klucz": "nr_roboty", "etykieta": "Nr roboty / KERG", "wymagane": True,
         "grupa": "Robota", "szerokosc": "trzecia"},
        {"klucz": "nr_operatu", "etykieta": "Nr operatu", "typ": "auto_numer",
         "domyslnie": "{numer3}/{rok}", "grupa": "Robota", "szerokosc": "trzecia"},
        {"klucz": "data_dokumentu", "etykieta": "Data dokumentu", "typ": "date",
         "domyslnie": "dzisiaj", "grupa": "Robota", "szerokosc": "trzecia"},
        {"klucz": "cel_opracowania", "etykieta": "Cel opracowania", "typ": "select",
         "grupa": "Robota", "szerokosc": "polowa",
         "opcje": ["Mapa do celów projektowych",
                   "Mapa z projektem podziału nieruchomości",
                   "Wznowienie znaków granicznych / wyznaczenie punktów granicznych",
                   "Geodezyjna inwentaryzacja powykonawcza",
                   "Tyczenie obiektu budowlanego",
                   "Ustalenie przebiegu granic działek ewidencyjnych"]},
        {"klucz": "osrodek", "etykieta": "Ośrodek dokumentacji (PODGiK)",
         "grupa": "Robota", "szerokosc": "polowa"},
        {"klucz": "data_zgloszenia", "etykieta": "Data zgłoszenia pracy", "typ": "date",
         "grupa": "Robota", "szerokosc": "trzecia"},

        {"klucz": "powiat", "etykieta": "Powiat", "grupa": "Położenie", "szerokosc": "trzecia"},
        {"klucz": "gmina", "etykieta": "Gmina", "grupa": "Położenie", "szerokosc": "trzecia"},
        {"klucz": "obreb", "etykieta": "Obręb", "grupa": "Położenie", "szerokosc": "trzecia"},
        {"klucz": "dzialki", "etykieta": "Numery działek", "grupa": "Położenie",
         "szerokosc": "polowa", "podpowiedz": "np. 123/4, 123/5, 124"},
        {"klucz": "miejscowosc", "etykieta": "Miejscowość (nagłówek dokumentu)",
         "grupa": "Położenie", "szerokosc": "polowa"},

        {"klucz": "zamawiajacy", "etykieta": "Zamawiający", "grupa": "Zamawiający",
         "szerokosc": "polowa"},
        {"klucz": "zamawiajacy_adres", "etykieta": "Adres zamawiającego",
         "grupa": "Zamawiający", "szerokosc": "polowa"},

        {"klucz": "zakres_opis", "etykieta": "Opis zakresu prac", "typ": "textarea",
         "grupa": "Opis techniczny"},
        {"klucz": "materialy", "etykieta": "Materiały wykorzystane do opracowania",
         "typ": "textarea", "grupa": "Opis techniczny",
         "podpowiedz": "np. mapa zasadnicza, operaty jednostkowe, dane osnowy"},
        {"klucz": "metoda_pomiaru", "etykieta": "Metoda pomiaru", "typ": "select",
         "grupa": "Opis techniczny", "szerokosc": "trzecia",
         "opcje": ["GNSS RTK/RTN", "biegunową", "ortogonalną", "mieszaną (GNSS + tachimetria)"]},
        {"klucz": "uklad_wspolrzednych", "etykieta": "Układ współrzędnych", "typ": "select",
         "grupa": "Opis techniczny", "szerokosc": "trzecia",
         "opcje": ["PL-2000 strefa 5", "PL-2000 strefa 6", "PL-2000 strefa 7",
                   "PL-2000 strefa 8", "PL-1992"]},
        {"klucz": "uklad_wysokosciowy", "etykieta": "Układ wysokościowy", "typ": "select",
         "grupa": "Opis techniczny", "szerokosc": "trzecia",
         "opcje": ["PL-EVRF2007-NH", "PL-KRON86-NH"]},
        {"klucz": "sprzet", "etykieta": "Sprzęt pomiarowy", "grupa": "Opis techniczny",
         "szerokosc": "polowa", "podpowiedz": "model odbiornika / tachimetru i nr fabryczny"},
        {"klucz": "data_prac_od", "etykieta": "Prace polowe od", "typ": "date",
         "grupa": "Opis techniczny", "szerokosc": "trzecia"},
        {"klucz": "data_prac_do", "etykieta": "Prace polowe do", "typ": "date",
         "grupa": "Opis techniczny", "szerokosc": "trzecia"},

        {"klucz": "punkty", "etykieta": "Wykaz współrzędnych", "typ": "tabela",
         "grupa": "Wykaz współrzędnych",
         "kolumny": [{"klucz": "numer", "etykieta": "Nr punktu"},
                     {"klucz": "x", "etykieta": "X [m]"},
                     {"klucz": "y", "etykieta": "Y [m]"},
                     {"klucz": "h", "etykieta": "H [m]"},
                     {"klucz": "uwagi", "etykieta": "Uwagi"}]},

        {"klucz": "uwagi", "etykieta": "Uwagi końcowe", "typ": "textarea", "grupa": "Uwagi"},

        {"klucz": "geodeta_imie_nazwisko", "zrodlo": "ustawienia"},
        {"klucz": "uprawnienia_nr", "zrodlo": "ustawienia"},
        {"klucz": "firma_nazwa", "zrodlo": "ustawienia"},
        {"klucz": "firma_adres", "zrodlo": "ustawienia"},
        {"klucz": "firma_telefon", "zrodlo": "ustawienia"},
        {"klucz": "firma_email", "zrodlo": "ustawienia"},
    ],
}


if __name__ == "__main__":
    plik = zbuduj()
    opis = plik.with_suffix(".json")
    opis.write_text(json.dumps(OPIS_POL, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Utworzono:\n  {plik}\n  {opis}")
