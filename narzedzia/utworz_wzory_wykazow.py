"""Tworzy szkielety dwóch wykazów zmian danych ewidencyjnych.

    szablony/wykaz_zmian_budynku_wzor.docx   -> wykaz_zmian_budynku.docx
    szablony/wykaz_zmian_dzialki_wzor.docx   -> wykaz_zmian_dzialki.docx

Oba dokładają się do operatu checkboxem i **nie mają własnych pól** — korzystają
z danych wpisanych przy spisie treści. Checkbox w formularzu pojawia się sam, bo lista
„Dokumenty do wygenerowania” bierze się z plików leżących w `szablony/`.

**Wszystko poza znacznikami to zaślepki do podmiany.** Brat wkleja swoje formatki
i przenosi do nich tagi; tutaj chodzi o gotowe nazwy i działający układ.

    python narzedzia/utworz_wzory_wykazow.py

Istniejących plików nie nadpisuje bez `--nadpisz`, żeby jedno uruchomienie z rozpędu
nie skasowało prawdziwej formatki.
"""
import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from app.config import SZABLONY  # noqa: E402

SRODEK = WD_ALIGN_PARAGRAPH.CENTER

# Nagłówki tabeli — do podmiany na te, których naprawdę wymaga ośrodek.
KOLUMNY_BUDYNKU = ["Lp.", "Oznaczenie budynku", "Stan dotychczasowy", "Stan po zmianie"]
KOLUMNY_DZIALKI = ["Lp.", "Nr działki", "Stan dotychczasowy", "Stan po zmianie"]

WYKAZY = [
    {
        "plik": "wykaz_zmian_budynku_wzor.docx",
        "tytul": "WYKAZ ZMIAN DANYCH EWIDENCYJNYCH DOTYCZĄCYCH BUDYNKU",
        "nazwa": "Wykaz zmian danych ewidencyjnych dotyczących budynku",
        "kolumny": KOLUMNY_BUDYNKU,
        "wiersz": ["1.", "{{ nr_dzialki }}", "—", "—"],
    },
    {
        "plik": "wykaz_zmian_dzialki_wzor.docx",
        "tytul": "WYKAZ ZMIAN DANYCH EWIDENCYJNYCH DOTYCZĄCYCH DZIAŁKI EWIDENCYJNEJ",
        "nazwa": "Wykaz zmian danych ewidencyjnych dotyczących działki ewidencyjnej",
        "kolumny": KOLUMNY_DZIALKI,
        "wiersz": ["1.", "{{ nr_dzialki }}", "—", "—"],
    },
]


def akapit(dokument, tekst, rozmiar=11, pogrubienie=False, wyrownanie=None, odstep_przed=0):
    element = dokument.add_paragraph()
    element.paragraph_format.space_before = Pt(odstep_przed)
    bieg = element.add_run(tekst)
    bieg.bold = pogrubienie
    bieg.font.size = Pt(rozmiar)
    if wyrownanie is not None:
        element.alignment = wyrownanie
    return element


def zbuduj(opis: dict) -> Path:
    dokument = Document()
    styl = dokument.styles["Normal"]
    styl.font.name = "Times New Roman"
    styl.font.size = Pt(11)

    akapit(dokument, opis["tytul"], 14, pogrubienie=True, wyrownanie=SRODEK, odstep_przed=24)
    akapit(dokument, "do operatu nr {{ nr_operatu }}", 11, wyrownanie=SRODEK)

    akapit(dokument, "Nr roboty: {{ nr_roboty }}", 11, odstep_przed=24)
    akapit(dokument, "Jednostka ewidencyjna: [{{ polozenie_gmina_teryt }}] "
                     "{{ polozenie_gmina }}", 11)
    akapit(dokument, "Obręb: [{{ polozenie_obreb_teryt }}] {{ polozenie_obreb }}", 11)
    akapit(dokument, "Działki nr: {{ nr_dzialki }}", 11)

    tabela = dokument.add_table(rows=2, cols=len(opis["kolumny"]))
    tabela.style = "Table Grid"
    for komorka, naglowek in zip(tabela.rows[0].cells, opis["kolumny"]):
        komorka.text = ""
        bieg = komorka.paragraphs[0].add_run(naglowek)
        bieg.bold = True
        bieg.font.size = Pt(10)
    for komorka, tresc in zip(tabela.rows[1].cells, opis["wiersz"]):
        komorka.text = ""
        komorka.paragraphs[0].add_run(tresc).font.size = Pt(10)

    akapit(dokument, "Sporządził: {{ data_zakonczenia }}", 10, odstep_przed=24)

    plik = SZABLONY / opis["plik"]
    dokument.save(plik)
    return plik


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--nadpisz", action="store_true",
                        help="nadpisz istniejące .docx (UWAGA: skasuje prawdziwe formatki)")
    argumenty = parser.parse_args()

    for opis in WYKAZY:
        docelowy = SZABLONY / opis["plik"]
        if docelowy.exists() and not argumenty.nadpisz:
            print(f"{docelowy.name} już istnieje — zostawiam bez zmian.")
        else:
            print("Utworzono:", zbuduj(opis).name)

        # własnych pól nie ma: wszystko przychodzi z formularza operatu
        docelowy.with_suffix(".json").write_text(
            json.dumps({"nazwa": opis["nazwa"],
                        "opis": "Dokłada się do operatu. Korzysta z danych "
                                "wpisanych przy spisie treści.",
                        "pola": []}, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8")
        print("  opis pól:", docelowy.with_suffix(".json").name)
