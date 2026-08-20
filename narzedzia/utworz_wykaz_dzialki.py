"""Formatka wykazu zmian danych działki — zbudowana z formatki **wykazu budynku**.

Oba wykazy mają wyglądać jak jeden komplet: pionowa kartka, ta sama tabela („L.p. |
Oznaczenie atrybutu | STAN DOTYCHCZASOWY | STAN NOWY"), ten sam nagłówek, ta sama stopka
i **osobna strona na każdą pozycję** — jedna działka to jedna strona, dokładnie tak jak
jeden budynek. Dlatego nie budujemy tego dokumentu od zera: bierzemy formatkę budynku,
podmieniamy tytuł, dokładamy wiersz nagłówka z numerem działki i przebudowujemy wiersze
tabeli na atrybuty działki. Gdy brat kiedyś zmieni wygląd wykazu budynku, ten skrypt
odtworzy działkę w tym samym stylu.

    python narzedzia/utworz_wykaz_dzialki.py --wyjscie szablony/wykaz_zmian_dzialki_wzor.docx

Po zbudowaniu puść `ujednolic_wyglad.py` — przeliczy przystanki tabulatorów w nagłówku
i w podpisie, bo doszedł tam dłuższy wiersz.

Numer działki jest **podpolem każdego wykazu** (`dzialka.dzialka`), a nie jednym polem
na cały dokument: skoro każda działka ma własną stronę, to i własny numer w nagłówku.
Identyfikatorem jest dopiero całość — obręb, kropka i ten numer — i tak stoi w nagłówku.

Użytki mają **cztery kolumny obok siebie** w każdym stanie: OFU, OZU, OZK i PPU (pole
powierzchni użytków gruntowych i klas bonitacyjnych w obszarze działki). Wcześniej PPU
było osobnym, czwartym wierszem tabeli — a jeden użytek to jeden komplet czterech
wartości, więc czytało się je w dwóch odległych miejscach. Teraz stoją w jednej linijce,
a kilka użytków to kilka linijek w tych samych czterech kolumnach.
"""
from __future__ import annotations

import argparse
import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

KORZEN = Path(__file__).resolve().parent.parent
ZRODLO = KORZEN / "szablony" / "wykaz_zmian_budynku_wzor.docx"

TYTUL = ("WYKAZ ZMIAN DANYCH EWIDENCYJNYCH ", "DOTYCZĄCYCH DZIAŁKI EWIDENCYJNEJ")

# Wiersz nagłówka dokumentu z identyfikatorem działki. Numer bierze się z pętli, bo każda
# działka ma własną stronę; obręb jest wspólny dla całego operatu.
WIERSZ_IDENTYFIKATORA = ("Identyfikator działki ewidencyjnej:",
                         "[{{ polozenie_obreb_teryt }}.{{ dzialka.dzialka }}]")

# Szerokości kolumn (twipy). Suma musi wyjść **dokładnie** na szerokość tekstu strony —
# tabela szersza choćby o włos wychodzi poza margines, a Word łamie ją wtedy na drugą
# stronę. PPU dostaje więcej miejsca niż OFU/OZU/OZK, bo trzyma liczby typu „0.4013”.
SZEROKOSC_LP = 543
# 2188: cała tabela wychodzi wtedy na 8931 twipów, dokładnie tyle co tabela
# w wykazie budynku — oba dokumenty leżą w operacie obok siebie i różnicę
# szerokości widać na pierwszy rzut oka
SZEROKOSC_OZNACZENIA = 2188
SZEROKOSCI_UZYTKOW = {"ofu": 700, "ozu": 700, "ozk": 700, "pow_uzytkow": 1000}

# Podkolumny stanu — jeden użytek to jedna linijka we wszystkich czterech naraz.
UZYTKI = [("OFU", "ofu"), ("OZU", "ozu"), ("OZK", "ozk"), ("PPU", "pow_uzytkow")]

# Wiersze tabeli: (nazwa atrybutu, klucz). Wartość idzie w komórce scalonej przez cztery
# podkolumny stanu — te rozdzielają się dopiero w wierszu użytków.
ATRYBUTY: list[tuple[str, str]] = [
    ("Numer działki", "numer"),
    ("Pole powierzchni ewidencyjnej działki [ha]", "pow_ewidencyjna"),
]
WIERSZ_UZYTKOW = "Użytki gruntowe i klasy bonitacyjne w działce"


def _tekst_komorki(tc, tekst: str) -> None:
    """Zostawia w komórce jeden akapit z jednym biegiem o podanej treści.

    Formatowanie bierze się z komórki wzorcowej — dlatego klonujemy prawdziwe komórki
    formatki budynku, zamiast składać tabelę od zera.
    """
    akapity = tc.findall(qn("w:p"))
    for zbedny in akapity[1:]:
        tc.remove(zbedny)
    akapit = akapity[0]
    biegi = akapit.findall(qn("w:r"))
    for zbedny in biegi[1:]:
        akapit.remove(zbedny)
    if not tekst:
        for bieg in biegi:
            akapit.remove(bieg)
        return
    if not biegi:
        bieg = OxmlElement("w:r")
        akapit.append(bieg)
    else:
        bieg = biegi[0]
    for stary in bieg.findall(qn("w:t")):
        bieg.remove(stary)
    for stary in bieg.findall(qn("w:br")):
        bieg.remove(stary)
    w_t = OxmlElement("w:t")
    w_t.set(qn("xml:space"), "preserve")
    w_t.text = tekst
    bieg.append(w_t)


# Kolejność elementów w `tcPr` jest częścią schematu OOXML — Word odrzuca plik, w którym
# coś stoi nie na swoim miejscu, a LibreOffice łyka to bez słowa (pułapka 12d).
KOLEJNOSC_TCPR = ("w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge",
                  "w:tcBorders", "w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
                  "w:tcFitText", "w:vAlign", "w:hideMark")


def _ustaw_w_tcpr(tcPr, tag: str, **atrybuty):
    """Element `tag` w `tcPr` — istniejący albo dołożony we właściwe miejsce."""
    element = tcPr.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        dalsze = [n.split(":")[1] for n in KOLEJNOSC_TCPR[KOLEJNOSC_TCPR.index(tag) + 1:]]
        for dziecko in tcPr:
            if dziecko.tag.split("}")[1] in dalsze:
                dziecko.addprevious(element)
                break
        else:
            tcPr.append(element)
    for nazwa, wartosc in atrybuty.items():
        element.set(qn(f"w:{nazwa}"), str(wartosc))
    return element


def _komorka(wzorzec, tekst: str, szerokosc: int = 0, *, span: int = 0, vmerge: str = ""):
    tc = copy.deepcopy(wzorzec)
    tcPr = tc.find(qn("w:tcPr"))

    stary_span = tcPr.find(qn("w:gridSpan"))
    if span:
        _ustaw_w_tcpr(tcPr, "w:gridSpan", val=span)
    elif stary_span is not None:
        tcPr.remove(stary_span)

    stary_merge = tcPr.find(qn("w:vMerge"))
    if stary_merge is not None:
        tcPr.remove(stary_merge)
    if vmerge == "restart":
        _ustaw_w_tcpr(tcPr, "w:vMerge", val="restart")
    elif vmerge == "dalej":
        _ustaw_w_tcpr(tcPr, "w:vMerge")

    if szerokosc:
        _ustaw_w_tcpr(tcPr, "w:tcW", w=szerokosc, type="dxa")

    _tekst_komorki(tc, tekst)
    return tc


def _wiersz(wzorzec_tr, komorki):
    tr = copy.deepcopy(wzorzec_tr)
    for stara in tr.findall(qn("w:tc")):
        tr.remove(stara)
    for tc in komorki:
        tr.append(tc)
    return tr


def _zamien_w_akapicie(akapit, co: str, na: str) -> bool:
    """Podmiana tekstu w akapicie, którego Word potrafi trzymać w kilku biegach."""
    biegi = akapit.findall(qn("w:r"))
    tekst = "".join(t.text or "" for b in biegi for t in b.iter(qn("w:t")))
    if not biegi or co not in tekst:
        return False
    # zostawiamy pierwszy bieg (z jego krojem i kolorem), reszta idzie precz — Word tnie
    # tekst na biegi w przypadkowych miejscach, więc podmiana „w miejscu” bywa krucha
    for zbedny in biegi[1:]:
        akapit.remove(zbedny)
    bieg = biegi[0]
    for stary in bieg.findall(qn("w:t")):
        bieg.remove(stary)
    w_t = OxmlElement("w:t")
    w_t.set(qn("xml:space"), "preserve")
    w_t.text = tekst.replace(co, na)
    bieg.append(w_t)
    return True


def _akapit_etykieta_wartosc(wzorzec, etykieta: str, wartosc: str):
    """Klon wiersza nagłówka „Etykieta:<tab>wartość” z podmienioną treścią obu części.

    Etykieta i wartość są w **osobnych biegach** i różnią się formatowaniem — wartość jest
    pogrubiona. Sklejenie całości w pierwszy bieg (tak robiła pierwsza wersja) gubiło to
    pogrubienie i wiersz z numerem działki wyglądał inaczej niż trzy nad nim.
    """
    akapit = copy.deepcopy(wzorzec)
    biegi = akapit.findall(qn("w:r"))
    tabulator = next(i for i, b in enumerate(biegi) if b.find(qn("w:tab")) is not None)

    for zbedny in akapit.findall(qn("w:proofErr")):
        akapit.remove(zbedny)
    for zbedny in biegi[1:tabulator] + biegi[tabulator + 2:]:
        akapit.remove(zbedny)

    for bieg, tekst in ((biegi[0], etykieta), (biegi[tabulator + 1], wartosc)):
        for stary in bieg.findall(qn("w:t")):
            bieg.remove(stary)
        w_t = OxmlElement("w:t")
        w_t.set(qn("xml:space"), "preserve")
        w_t.text = tekst
        bieg.append(w_t)
    return akapit


def zbuduj(zrodlo: Path, wyjscie: Path) -> Path:
    dokument = Document(str(zrodlo))
    body = dokument.element.body

    # --- pętla i tytuł --------------------------------------------------------
    for akapit in body.iter(qn("w:p")):
        _zamien_w_akapicie(akapit, "{%p for wykaz in wykazy_budynkow %}",
                           "{%p for dzialka in wykazy_dzialek %}")
        _zamien_w_akapicie(akapit, "WYKAZ ZMIAN DANYCH EWIDENCYJNYCH ", TYTUL[0])
        _zamien_w_akapicie(akapit, "DOTYCZĄCYCH BUDYNKU", TYTUL[1])

    # --- wiersz nagłówka z numerem działki -------------------------------------
    wzorzec_naglowka = next(
        p for p in body.iter(qn("w:p"))
        if "Obręb ewidencyjny" in "".join(t.text or "" for t in p.iter(qn("w:t"))))
    wzorzec_naglowka.addnext(_akapit_etykieta_wartosc(
        wzorzec_naglowka, WIERSZ_IDENTYFIKATORA[0], WIERSZ_IDENTYFIKATORA[1]))

    # --- tabela ---------------------------------------------------------------
    # Wzorce bierzemy z prawdziwych komórek formatki budynku: krój, rozmiary, kreski
    # i wyśrodkowanie są tam już takie, jakie mają być w obu wykazach.
    tabela = dokument.tables[0]._tbl
    wiersze = tabela.findall(qn("w:tr"))
    naglowek_wz, prosty_wz = wiersze[0], wiersze[1]
    komorki_naglowka = naglowek_wz.findall(qn("w:tc"))
    komorki_proste = prosty_wz.findall(qn("w:tc"))
    lp_wz, opis_wz, dane_wz = komorki_proste[0], komorki_proste[1], komorki_proste[3]
    naglowek_lp, naglowek_opisu, naglowek_stanu = komorki_naglowka[0], komorki_naglowka[1], \
        komorki_naglowka[2]

    szerokosc_stanu = sum(SZEROKOSCI_UZYTKOW.values())
    nowe: list = []

    # --- nagłówek: dwa wiersze, bo stany dzielą się na cztery podkolumny -------
    nowe.append(_wiersz(naglowek_wz, [
        _komorka(naglowek_lp, "L.p.", SZEROKOSC_LP, vmerge="restart"),
        _komorka(naglowek_opisu, "Oznaczenie atrybutu działki", SZEROKOSC_OZNACZENIA,
                 vmerge="restart"),
        _komorka(naglowek_stanu, "STAN DOTYCHCZASOWY", szerokosc_stanu, span=4),
        _komorka(naglowek_stanu, "STAN NOWY", szerokosc_stanu, span=4),
    ]))
    nowe.append(_wiersz(naglowek_wz, [
        _komorka(naglowek_lp, "", SZEROKOSC_LP, vmerge="dalej"),
        _komorka(naglowek_opisu, "", SZEROKOSC_OZNACZENIA, vmerge="dalej"),
        *[_komorka(naglowek_stanu, skrot, SZEROKOSCI_UZYTKOW[klucz])
          for _ in range(2) for skrot, klucz in UZYTKI],
    ]))

    # --- atrybuty bez podziału na użytki: wartość na całą szerokość stanu ------
    for numer, (nazwa, klucz) in enumerate(ATRYBUTY, start=1):
        nowe.append(_wiersz(prosty_wz, [
            _komorka(lp_wz, f"{numer}.", SZEROKOSC_LP),
            _komorka(opis_wz, nazwa, SZEROKOSC_OZNACZENIA),
            _komorka(dane_wz, "{{ dzialka.%s_dotychczas }}" % klucz, szerokosc_stanu, span=4),
            _komorka(dane_wz, "{{r dzialka.%s_nowy }}" % klucz, szerokosc_stanu, span=4),
        ]))

    # --- użytki: cztery kolumny obok siebie, po jednej linijce na użytek -------
    nowe.append(_wiersz(prosty_wz, [
        _komorka(lp_wz, f"{len(ATRYBUTY) + 1}.", SZEROKOSC_LP),
        _komorka(opis_wz, WIERSZ_UZYTKOW, SZEROKOSC_OZNACZENIA),
        *[_komorka(dane_wz, "{{ dzialka.%s_dotychczas }}" % klucz, SZEROKOSCI_UZYTKOW[klucz])
          for _, klucz in UZYTKI],
        *[_komorka(dane_wz, "{{r dzialka.%s_nowy }}" % klucz, SZEROKOSCI_UZYTKOW[klucz])
          for _, klucz in UZYTKI],
    ]))

    for wiersz in wiersze:
        tabela.remove(wiersz)
    grid = tabela.find(qn("w:tblGrid"))
    for kolumna in grid.findall(qn("w:gridCol")):
        grid.remove(kolumna)
    for szerokosc in [SZEROKOSC_LP, SZEROKOSC_OZNACZENIA,
                      *[SZEROKOSCI_UZYTKOW[k] for _, k in UZYTKI] * 2]:
        kolumna = OxmlElement("w:gridCol")
        kolumna.set(qn("w:w"), str(szerokosc))
        grid.append(kolumna)
    _ustaw_w_tcpr(tabela.find(qn("w:tblPr")), "w:tblW",
                  w=SZEROKOSC_LP + SZEROKOSC_OZNACZENIA + 2 * szerokosc_stanu, type="dxa")
    for wiersz in nowe:
        tabela.append(wiersz)

    wyjscie.parent.mkdir(parents=True, exist_ok=True)
    dokument.save(str(wyjscie))
    return wyjscie


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--wyjscie", required=True, type=Path)
    parser.add_argument("--zrodlo", type=Path, default=ZRODLO)
    argumenty = parser.parse_args()

    dokument = Document(str(argumenty.zrodlo))
    if "wykazy_budynkow" not in dokument.element.body.xml:
        print(f"{argumenty.zrodlo} to nie jest formatka wykazu budynku — "
              "z niej bierze się układ tabeli i indeksy wierszy wzorcowych.",
              file=sys.stderr)
        return 1

    plik = zbuduj(argumenty.zrodlo, argumenty.wyjscie)
    print(f"{plik}: {len(Document(str(plik)).tables[0].rows)} wierszy tabeli, "
          "jedna działka na stronę")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
