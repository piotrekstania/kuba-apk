"""Wariant wykazu zmian danych działki na kartce **pionowej**.

Dzisiejsza formatka jest pozioma, bo trzyma trzynaście kolumn: L.p. plus sześć kolumn
stanu dotychczasowego i sześć stanu nowego, obok siebie. Ten wariant układa to inaczej:
kolumn jest osiem, a **każda działka zajmuje dwa wiersze** — „dotychczasowy" nad „nowym".
Dzięki temu wartości przed i po leżą jedno pod drugim, w tej samej kolumnie, i porównuje
się je okiem w pionie, zamiast wodzić palcem przez pół szerokiej kartki. Każda kolumna
jest przy tym **szersza** niż dziś, mimo węższej strony.

Skrypt nie tworzy dokumentu od zera: bierze **poziomą** formatkę i podmienia w niej samą
tabelę, zostawiając nagłówek, logo, podpis i stopkę takimi, jakie są.

**Konwersja jest już zrobiona** — od 19.08.2026 pionowy wykaz jest formatką domyślną
(commit `2b6596a` to ostatni z poziomą). Skrypt zostaje jako opis tego, co się wtedy
stało, i na wypadek, gdyby brat przysłał kolejną poziomą tabelę do przełożenia; sam
odmawia pracy, gdy dostanie plik, który jest już pionowy — inaczej klonowałby komórki
spod błędnych indeksów i wychodziłby z tego dokument bez sensu. Poziomy wzór wyjmiesz
z historii:

    git show 2b6596a:szablony/wykaz_zmian_dzialki_wzor.docx > /tmp/poziomy.docx

    python narzedzia/utworz_wykaz_dzialki_pionowy.py --wyjscie /tmp/wykaz_pionowy.docx

Po zbudowaniu puść na pliku `ujednolic_wyglad.py` — przeliczy przystanki tabulatorów
w nagłówku, w podpisie i w stopce, bo te liczą się z **szerokości strony**, a ta się
właśnie zmieniła.
"""
from __future__ import annotations

import argparse
import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips

KORZEN = Path(__file__).resolve().parent.parent
ZRODLO = KORZEN / "szablony" / "wykaz_zmian_dzialki_wzor.docx"

# A4 pionowo, w twipach. Marginesy zostają takie jak w pozostałych formatkach, więc
# na tabelę zostaje 9105 twipów (16,1 cm) — i to wystarcza, bo kolumn jest osiem,
# a nie trzynaście. Tabela szersza choćby o włos wychodzi poza margines i Word łamie
# ją na drugą stronę, więc szerokości muszą się sumować dokładnie.
STRONA_SZEROKOSC, STRONA_WYSOKOSC = 11906, 16838

# Szerokości kolumn. Dla porównania w dzisiejszej, poziomej formatce: numer działki 851,
# pole powierzchni 1416, OFU/OZU/OZK po ~709, pole użytków 1701 — czyli każda kolumna
# tego wariantu jest szersza albo taka sama, mimo węższej kartki.
KOLUMNY = {
    "lp": 563,
    # 1300 twipów: pogrubione „Dotychczasowy" mierzy 1022 twipy przy 8 pt, plus
    # marginesy komórki (2 × 108) i zapas na to, że Word łamie wiersz odrobinę
    # wcześniej, niż wynika z sumy szerokości znaków (pułapka 12g)
    "stan": 1300,
    "numer": 1100,
    "pole": 1500,
    "ofu": 709,
    "ozu": 709,
    "ozk": 709,
    "uzytki": 2515,
}
SZEROKOSC_TABELI = sum(KOLUMNY.values())

# Kolejność elementów w `tcPr` jest częścią schematu OOXML — Word odrzuca plik, w którym
# coś stoi nie na swoim miejscu, a LibreOffice łyka to bez słowa (patrz pułapka 12d
# w CLAUDE.md). Stąd wstawianie „przed pierwszym z następców", a nie na koniec.
KOLEJNOSC_TCPR = ("w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge",
                  "w:tcBorders", "w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
                  "w:tcFitText", "w:vAlign", "w:hideMark")

# Wcięcie bloku podpisu — tyle samo co w wykazie budynku, czyli w drugiej pionowej
# formatce operatu. Przystanek tabulatora zostawiamy pusty: przeliczy go
# `ujednolic_wyglad.py`, mierząc najdłuższą etykietę prawdziwym plikiem czcionki.
WCIECIE_PODPISU = 3401

GRUBA_KRESKA = 18        # 2,25 pt — oddziela jedną działkę od następnej
CIENKA_KRESKA = 4        # 0,5 pt — tyle mają wszystkie kreski w formatce

# Nagłówek tabeli jest wytłuszczony w całości, tak jak w wykazie budynku — łącznie
# z „Dotychczasowy" i „Nowy" w kolumnie stanu, bo to one rozdzielają parę wierszy.
POGRUBIONE = ("w:i", "w:iCs", "w:caps", "w:smallCaps", "w:strike", "w:dstrike",
              "w:outline", "w:shadow", "w:emboss", "w:imprint", "w:noProof",
              "w:snapToGrid", "w:vanish", "w:webHidden", "w:color", "w:spacing", "w:w",
              "w:kern", "w:position", "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect",
              "w:bdr", "w:shd", "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em",
              "w:lang", "w:eastAsianLayout", "w:specVanish", "w:oMath")


def _wstaw_wg_schematu(rodzic, element) -> None:
    """Wkłada element w miejsce, którego wymaga schemat OOXML."""
    nazwa = element.tag.split("}")[1]
    dalsze = KOLEJNOSC_TCPR[KOLEJNOSC_TCPR.index(f"w:{nazwa}") + 1:]
    for dziecko in rodzic:
        if dziecko.tag.split("}")[1] in [d.split(":")[1] for d in dalsze]:
            dziecko.addprevious(element)
            return
    rodzic.append(element)


def _ustaw(rodzic, tag: str, **atrybuty):
    """Element `tag` w `rodzic` — istniejący albo dołożony we właściwe miejsce."""
    element = rodzic.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        _wstaw_wg_schematu(rodzic, element)
    for nazwa, wartosc in atrybuty.items():
        element.set(qn(f"w:{nazwa}"), str(wartosc))
    return element


def _pogrub(bieg, pogrubienie: bool) -> None:
    """Włącza albo zdejmuje pogrubienie biegu, z zachowaniem kolejności w `rPr`."""
    rPr = bieg.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        bieg.insert(0, rPr)
    for tag in ("w:b", "w:bCs"):
        istniejacy = rPr.find(qn(tag))
        if not pogrubienie:
            if istniejacy is not None:
                rPr.remove(istniejacy)
            continue
        if istniejacy is not None:
            continue
        element = OxmlElement(tag)
        for dziecko in rPr:
            if dziecko.tag.split("}")[1] in [n.split(":")[1] for n in POGRUBIONE]:
                dziecko.addprevious(element)
                break
        else:
            rPr.append(element)


def _tekst_komorki(tc, tekst: str, pogrubienie: bool | None = None) -> None:
    """Zostawia w komórce jeden akapit z jednym biegiem o podanej treści.

    Formatowanie (krój, rozmiar, wyśrodkowanie) bierze się z komórki wzorcowej —
    dlatego klonujemy prawdziwe komórki formatki, zamiast budować tabelę od zera.
    """
    akapity = tc.findall(qn("w:p"))
    for zbedny in akapity[1:]:
        tc.remove(zbedny)
    akapit = akapity[0]
    biegi = akapit.findall(qn("w:r"))
    for zbedny in biegi[1:]:
        akapit.remove(zbedny)
    if not tekst:
        for bieg in biegi:
            akapit.remove(bieg)
        return
    if not biegi:                       # komórka wzorcowa była pusta
        bieg = OxmlElement("w:r")
        akapit.append(bieg)
    else:
        bieg = biegi[0]
    for w_t in bieg.findall(qn("w:t")):
        bieg.remove(w_t)
    for w_br in bieg.findall(qn("w:br")):
        bieg.remove(w_br)
    w_t = OxmlElement("w:t")
    w_t.set(qn("xml:space"), "preserve")
    w_t.text = tekst
    bieg.append(w_t)
    if pogrubienie is not None:
        _pogrub(bieg, pogrubienie)


def _komorka(wzorzec, tekst: str, szerokosc: int, *, span: int = 0, vmerge: str = "",
             dol: int = 0, pogrubienie: bool | None = None):
    """Kopia komórki wzorcowej z podmienioną treścią i szerokością."""
    tc = copy.deepcopy(wzorzec)
    tcPr = tc.find(qn("w:tcPr"))
    _ustaw(tcPr, "w:tcW", w=szerokosc, type="dxa")

    stary_span = tcPr.find(qn("w:gridSpan"))
    if span:
        _ustaw(tcPr, "w:gridSpan", val=span)
    elif stary_span is not None:
        tcPr.remove(stary_span)

    stary_merge = tcPr.find(qn("w:vMerge"))
    if stary_merge is not None:
        tcPr.remove(stary_merge)
    if vmerge == "restart":
        _ustaw(tcPr, "w:vMerge", val="restart")
    elif vmerge == "dalej":
        _ustaw(tcPr, "w:vMerge")

    if dol:
        brzegi = _ustaw(tcPr, "w:tcBorders")
        dolna = brzegi.find(qn("w:bottom"))
        if dolna is None:
            dolna = OxmlElement("w:bottom")
            brzegi.append(dolna)
        dolna.set(qn("w:val"), "single")
        dolna.set(qn("w:sz"), str(dol))
        dolna.set(qn("w:space"), "0")
        dolna.set(qn("w:color"), "000000")

    _tekst_komorki(tc, tekst, pogrubienie)
    return tc


def _wiersz(wzorzec_tr, komorki, *, naglowek: bool = False):
    """Wiersz zbudowany z gotowych komórek, z zachowaniem właściwości wzorca."""
    tr = copy.deepcopy(wzorzec_tr)
    for stara in tr.findall(qn("w:tc")):
        tr.remove(stara)
    trPr = tr.find(qn("w:trPr"))
    if naglowek:
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))
    elif trPr is not None and trPr.find(qn("w:tblHeader")) is not None:
        trPr.remove(trPr.find(qn("w:tblHeader")))
    for tc in komorki:
        tr.append(tc)
    return tr


def _na_pionowa(dokument) -> None:
    """A4 w pionie. Marginesy zostają — dzięki temu logo i stopka siedzą tam, gdzie w reszcie."""
    sectPr = dokument.element.body.find(qn("w:sectPr"))
    pgSz = sectPr.find(qn("w:pgSz"))
    pgSz.set(qn("w:w"), str(STRONA_SZEROKOSC))
    pgSz.set(qn("w:h"), str(STRONA_WYSOKOSC))
    if pgSz.get(qn("w:orient")):
        del pgSz.attrib[qn("w:orient")]


def _przenies_podpis(dokument) -> None:
    """Blok podpisu ustawiony pod **poziomą** kartkę zjeżdża w pionie poza margines.

    Wcięcie miał 8759 twipów, bo tyle było sensu przy 24-centymetrowej szerokości tekstu;
    na pionowej kartce zostaje z niego kolumna szeroka na jeden znak i podpis łamie się
    po literze. Ustawiamy wcięcie takie jak w wykazie budynku i **kasujemy przystanki
    tabulatora** — wtedy `ujednolic_wyglad.py` policzy kolumnę wartości sam. Bez tego
    kasowania uzna blok za już wyrównany (bo wszystkie wiersze mają ten sam przystanek)
    i w ogóle go nie tknie.
    """
    for akapit in dokument.paragraphs:
        if not any(etykieta in akapit.text for etykieta in
                   ("Sporządził", "Kierownik prac", "Data sporządzenia")):
            continue
        ustawienia = akapit.paragraph_format
        ustawienia.left_indent = Twips(WCIECIE_PODPISU)
        ustawienia.tab_stops.clear_all()


def zbuduj(zrodlo: Path, wyjscie: Path) -> Path:
    dokument = Document(str(zrodlo))
    if len(dokument.tables[0].columns) == len(KOLUMNY):
        raise SystemExit(
            f"{zrodlo} jest już pionowy — ten skrypt przekłada tabelę **poziomą** "
            "(13 kolumn) i klonuje komórki spod indeksów tamtej. Poziomy wzór:\n"
            "    git show 2b6596a:szablony/wykaz_zmian_dzialki_wzor.docx > /tmp/poziomy.docx")
    _na_pionowa(dokument)
    _przenies_podpis(dokument)

    stara = dokument.tables[0]._tbl
    wiersze = stara.findall(qn("w:tr"))
    # wzorce komórek — bierzemy prawdziwe komórki formatki, żeby nie odtwarzać ręcznie
    # kroju, rozmiaru, wyśrodkowania i kresek
    naglowek_wz = wiersze[1].findall(qn("w:tc"))[1]      # „Numer działki"
    grupa_wz = wiersze[1].findall(qn("w:tc"))[3]         # „Użytki gruntowe…", scalona na 3
    pod_wz = wiersze[2].findall(qn("w:tc"))[3]           # „OFU"
    dane_wz = wiersze[4].findall(qn("w:tc"))[1]          # komórka z danymi działki
    sterujaca_wz = wiersze[3].findall(qn("w:tc"))[0]     # komórka ze znacznikiem `{%tr %}`
    lp_wz = wiersze[4].findall(qn("w:tc"))[0]            # L.p. w wierszu danych

    K = KOLUMNY
    nowe: list = []

    # --- nagłówek: dwa wiersze, bo OFU/OZU/OZK siedzą pod wspólnym tytułem -----
    nowe.append(_wiersz(wiersze[1], [
        _komorka(naglowek_wz, "L.p.", K["lp"], vmerge="restart", pogrubienie=True),
        _komorka(naglowek_wz, "Stan", K["stan"], vmerge="restart", pogrubienie=True),
        _komorka(naglowek_wz, "Numer działki", K["numer"], vmerge="restart",
                 pogrubienie=True),
        _komorka(naglowek_wz, "Pole powierzchni ewidencyjnej działki [ha]", K["pole"],
                 vmerge="restart", pogrubienie=True),
        _komorka(grupa_wz, "Użytki gruntowe i klasy bonitacyjne w działce",
                 K["ofu"] + K["ozu"] + K["ozk"], span=3, pogrubienie=True),
        _komorka(naglowek_wz,
                 "Pole powierzchni użytków gruntowych i klas bonitacyjnych "
                 "w obszarze działki [ha]", K["uzytki"], vmerge="restart",
                 pogrubienie=True),
    ], naglowek=True))

    nowe.append(_wiersz(wiersze[2], [
        _komorka(naglowek_wz, "", K["lp"], vmerge="dalej"),
        _komorka(naglowek_wz, "", K["stan"], vmerge="dalej"),
        _komorka(naglowek_wz, "", K["numer"], vmerge="dalej"),
        _komorka(naglowek_wz, "", K["pole"], vmerge="dalej"),
        _komorka(pod_wz, "OFU", K["ofu"], pogrubienie=True),
        _komorka(pod_wz, "OZU", K["ozu"], pogrubienie=True),
        _komorka(pod_wz, "OZK", K["ozk"], pogrubienie=True),
        _komorka(naglowek_wz, "", K["uzytki"], vmerge="dalej"),
    ], naglowek=True))

    # --- wiersz sterujący pętli ----------------------------------------------
    # `{%tr for %}` kasuje **cały wiersz**, w którym stoi, i powtarza wszystko między
    # sobą a `{%tr endfor %}` — czyli u nas dwa wiersze naraz (patrz pułapka 1).
    nowe.append(_wiersz(wiersze[3], [
        _komorka(sterujaca_wz, "{%tr for dzialka in wykazy_dzialek %}", K["lp"]),
        *[_komorka(sterujaca_wz, "", szerokosc)
          for nazwa, szerokosc in K.items() if nazwa != "lp"],
    ]))

    # --- dwa wiersze na działkę ----------------------------------------------
    nowe.append(_wiersz(wiersze[4], [
        # kropka po numerze — tak jak w wykazie budynku, gdzie stoi „1.", „2."
        _komorka(lp_wz, "{{ loop.index }}.", K["lp"], vmerge="restart"),
        _komorka(dane_wz, "Dotychczasowy", K["stan"], pogrubienie=True),
        _komorka(dane_wz, "{{ dzialka.numer_dotychczas }}", K["numer"]),
        _komorka(dane_wz, "{{ dzialka.pow_ewidencyjna_dotychczas }}", K["pole"]),
        _komorka(dane_wz, "{{ dzialka.ofu_dotychczas }}", K["ofu"]),
        _komorka(dane_wz, "{{ dzialka.ozu_dotychczas }}", K["ozu"]),
        _komorka(dane_wz, "{{ dzialka.ozk_dotychczas }}", K["ozk"]),
        _komorka(dane_wz, "{{ dzialka.pow_uzytkow_dotychczas }}", K["uzytki"]),
    ]))

    # dolna kreska tego wiersza jest gruba — to ona oddziela jedną działkę od następnej
    nowe.append(_wiersz(wiersze[4], [
        _komorka(lp_wz, "", K["lp"], vmerge="dalej", dol=GRUBA_KRESKA),
        _komorka(dane_wz, "Nowy", K["stan"], dol=GRUBA_KRESKA, pogrubienie=True),
        _komorka(dane_wz, "{{ dzialka.numer_nowy }}", K["numer"], dol=GRUBA_KRESKA),
        _komorka(dane_wz, "{{ dzialka.pow_ewidencyjna_nowy }}", K["pole"], dol=GRUBA_KRESKA),
        _komorka(dane_wz, "{{ dzialka.ofu_nowy }}", K["ofu"], dol=GRUBA_KRESKA),
        _komorka(dane_wz, "{{ dzialka.ozu_nowy }}", K["ozu"], dol=GRUBA_KRESKA),
        _komorka(dane_wz, "{{ dzialka.ozk_nowy }}", K["ozk"], dol=GRUBA_KRESKA),
        _komorka(dane_wz, "{{ dzialka.pow_uzytkow_nowy }}", K["uzytki"], dol=GRUBA_KRESKA),
    ]))

    nowe.append(_wiersz(wiersze[5], [
        _komorka(sterujaca_wz, "{%tr endfor %}", K["lp"]),
        *[_komorka(sterujaca_wz, "", szerokosc)
          for nazwa, szerokosc in K.items() if nazwa != "lp"],
    ]))

    # --- podmiana tabeli ------------------------------------------------------
    for wiersz in wiersze:
        stara.remove(wiersz)
    grid = stara.find(qn("w:tblGrid"))
    for kolumna in grid.findall(qn("w:gridCol")):
        grid.remove(kolumna)
    for szerokosc in KOLUMNY.values():
        kolumna = OxmlElement("w:gridCol")
        kolumna.set(qn("w:w"), str(szerokosc))
        grid.append(kolumna)
    _ustaw(stara.find(qn("w:tblPr")), "w:tblW", w=SZEROKOSC_TABELI, type="dxa")
    for wiersz in nowe:
        stara.append(wiersz)

    wyjscie.parent.mkdir(parents=True, exist_ok=True)
    dokument.save(str(wyjscie))
    return wyjscie


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--wyjscie", required=True, type=Path,
                        help="gdzie zapisać wariant (nigdy do szablony/)")
    parser.add_argument("--zrodlo", type=Path, default=ZRODLO)
    argumenty = parser.parse_args()

    if argumenty.wyjscie.resolve().parent == (KORZEN / "szablony"):
        print("Nie zapisuję do szablony/ — to katalog lustrzany, wariant idzie do dane/.",
              file=sys.stderr)
        return 1

    plik = zbuduj(argumenty.zrodlo, argumenty.wyjscie)
    print(f"{plik}: 8 kolumn, {SZEROKOSC_TABELI} twipów, dwa wiersze na działkę")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
