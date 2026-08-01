"""Nakłada wspólny wygląd na szablony operatu.

Zamysł: dokumenty w jednej teczce mają wyglądać jak komplet, a nie jak formatki robione
w różnych latach. Hierarchię buduje **rozmiar i grubość pisma**, a nie podkreślenia
i ozdobniki — stąd „minimalistycznie”.

**Podział ról: skrypt odpowiada za hierarchię, autor formatki za akcenty.** Skrypt
mówi, co jest tytułem, nagłówkiem i etykietą, jakim krojem i w jakim rozmiarze —
ale o tym, którą wartość w danym dokumencie wyróżnić pogrubieniem, decyduje brat
w Wordzie. Wymuszanie grubości w treści kasowało jego poprawki przy każdym
uruchomieniu (`pogrubienie=None` w `_ustaw` znaczy „zostaw, jak jest”).

Czego skrypt NIE rusza:

* słów — ani jednego znaku treści, także znaczników ``{{ }}`` i ``{%p %}``,
* czerwieni — numer roboty zostaje czerwony wszędzie, gdzie był (i pogrubiony:
  to jedyne miejsce, gdzie grubość narzucamy),
* pogrubień — ani w treści, ani w etykietach, ani w tabelach,
* wyrównania akapitów i rozmiarów pisma w tabelach,
* **tabulatorów poza wyrównywanym blokiem** — ciąg tabulatorów to świadome
  wcięcie autora, a nie usterka do wyprostowania,
* marginesów i rozmiaru strony.

Ta lista rosła kosztem czterech wpadek: skrypt po kolei kasował bratu pogrubienia
w treści, pogrubienia etykiet, wybrany rozmiar nagłówka i wcięcie zrobione
tabulatorami. Za każdym razem brałem jego świadomy wybór za niespójność do
naprawienia. **Jeśli wahasz się, czy coś ujednolicić — nie ujednolicaj.**

Co ujednolica:

* **krój** — jeden na cały dokument, także w stylu bazowym i w tabelach, żeby
  dopisany później w Wordzie akapit nie wyszedł inną czcionką,
* **logo** — ta sama szerokość i to samo miejsce w każdym dokumencie; proporcje
  obrazka zostają nienaruszone,
* **stopkę** — przepisywaną z `STOPKA_WZORCOWA`, żeby była identyczna wszędzie;
  jej trzy kolumny rozkładają się na szerokość **danej** strony, więc formatka
  pozioma nie dostaje ustawień pionowej (`_przystanki_stopki`),
* **hierarchię** — tytuł, nagłówek sekcji, etykieta, tekst; puste akapity dostają
  jeden mały rozmiar, więc odstępy między blokami są równe,
* **wcięcia** — akapit wcięty „pierwszym wierszem” dostaje wcięcie całego akapitu.
  To najczęstsza usterka tych formatek: druga linijka zdania wracała na margines
  i blok rozjeżdżał się w schodki,
* **kolumny w blokach „Etykieta:<tab>wartość”** — wspólne wcięcie i jeden przystanek
  wyliczony z **pomiaru najdłuższej etykiety** prawdziwym plikiem czcionki
  (`_wyrownaj_bloki`). Blok, który jest już równy, zostaje nietknięty,
* **łamańce** — zdanie rozbite na kilka akapitów wraca do jednego akapitu
  (patrz `_scal_ciagi`).

Role akapitów rozpoznaje po tym, co w dokumencie już jest, więc reguły zadziałają
też na formatkach dołożonych później.

Na koniec `_sprawdz_kolejnosc` odmawia zapisania pliku, którego Word by nie otworzył
— patrz pułapki 12d i 12e w CLAUDE.md. Zielony PDF z LibreOffice'a nie jest dowodem,
że plik jest poprawny.

    python narzedzia/ujednolic_wyglad.py                     # wszystkie wzory
    python narzedzia/ujednolic_wyglad.py szablony/inny.docx

Uruchomienie na pojedynczym pliku i tak przetwarza najpierw `STOPKA_WZORCOWA`,
bo to z niej bierze się stopka — ten plik zapisze się więc nawet wtedy, gdy nic
się w nim nie zmienia.
"""
import argparse
import copy
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from app.config import SZABLONY  # noqa: E402

# Calibri jest na każdym Windowsie z Office'em, a na Linuksie zastępuje ją
# metrycznie zgodne Carlito (pakiet fonts-crosextra-carlito) — ten sam dokument
# łamie się więc tak samo po obu stronach.
KROJ = "Calibri"

TYTUL = Pt(14)
NAGLOWEK = Pt(12)          # rozmiar wybrany przez brata; skrypt tylko go pilnuje
TRESC = Pt(10)
PUSTY = Pt(6)              # wysokość pustego akapitu = odstęp między blokami

# Rozstrzelenie tytułu w dwudziestych częściach punktu. Drobiazg, ale to on sprawia,
# że wersaliki wyglądają na złożone, a nie na przypadkiem powiększone.
TYTUL_ROZSTRZELENIE = 12
ODSTEP_NAD_TYTULEM = Pt(8)      # oddziela tytuł od logo, gdy stoi zaraz pod nim
ODSTEP_POD_TYTULEM = Pt(14)

# Logo: szerokość i miejsce takie samo w każdym dokumencie (od marginesu, nie od
# krawędzi kartki — inaczej rozjechałoby się przy innych marginesach).
LOGO_SZEROKOSC = Mm(50)
LOGO_OD_LEWEJ = Mm(5.8)
LOGO_OD_GORY = Mm(4)

# Stopka pochodzi z jednego pliku i jest przepisywana do pozostałych — patrz
# `_stopka_wzorcowa`. Ten plik jest jej źródłem prawdy.
STOPKA_WZORCOWA = "spis_tresci_wzor.docx"
STOPKA_ROZMIAR = Pt(8)
STOPKA_KOLOR = RGBColor(0x6B, 0x6B, 0x6B)

PODPIS = "Stania"

ZAMYKA_ZDANIE = (".", ":", ";", "!", "?")
DLUGOSC_PELNEGO_WIERSZA = 60    # patrz `_scal_ciagi`


def _znacznik(tekst: str) -> bool:
    return "{{" in tekst or "{%" in tekst


def _czerwony(bieg) -> bool:
    return (bieg.font.color is not None and bieg.font.color.type is not None
            and str(bieg.font.color.rgb) == "FF0000")


def _rozmiar(akapit) -> float:
    return max((b.font.size.pt for b in akapit.runs if b.font.size), default=0)


# --- struktura -------------------------------------------------------------------

def _scal_ciagi(dokument) -> int:
    """Skleja zdanie rozbite na osobne akapity.

    W formatkach pisanych „na oko” dłuższy opis bywa wprowadzony Enterem zamiast
    zawinięciem wiersza. Word traktuje każdy kawałek jak osobny akapit, więc tekst
    czyta się jak kolumna urwanych linijek, a przy innej czcionce łamie się w losowych
    miejscach.

    Sklejamy tylko wtedy, gdy zdanie jest **bez wątpienia** urwane w połowie: poprzedni
    akapit jest wyrównany do lewej, ma długość pełnego wiersza i nie kończy się kropką
    ani dwukropkiem, a dalszy ciąg zaczyna się **małą literą**.

    Ten ostatni warunek doszedł później i jest najważniejszy. Bez niego reguła sklejała
    wiersz, który brat celowo złamał — opis punktu pomiarowego i tolerancje w nawiasach
    kwadratowych pod spodem. Krótkie, wyśrodkowane wpisy (adres firmy, NIP) chroni
    wymagana długość poprzedniego akapitu.
    """
    scalone = 0
    poprzedni = None
    for akapit in list(dokument.paragraphs):
        tekst = akapit.text.strip()
        do_lewej = akapit.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT,
                                        WD_ALIGN_PARAGRAPH.JUSTIFY)

        # scalamy tylko czysty tekst: w akapicie z odsyłaczem, zakładką albo polem
        # Worda przeniesienie samych biegów zgubiłoby resztę bez śladu
        czysty = all(_nazwa(e) in ("pPr", "r", "bookmarkStart", "bookmarkEnd", "proofErr")
                     for e in akapit._p)

        moze = (poprzedni is not None and tekst and do_lewej and czysty
                and not _znacznik(tekst) and not _znacznik(poprzedni.text)
                and len(poprzedni.text.strip()) >= DLUGOSC_PELNEGO_WIERSZA
                and not poprzedni.text.strip().endswith(ZAMYKA_ZDANIE)
                # dalszy ciąg musi zaczynać się małą literą — wtedy to bez wątpienia
                # środek zdania. Wiersz zaczynający się wielką literą, cyfrą albo
                # nawiasem bywa **celowo** osobny (tolerancje pomiaru w nawiasach
                # kwadratowych pod opisem punktu) i sklejenie go psuje układ
                and tekst[:1].islower())
        if moze:
            ostatni = poprzedni.runs[-1]
            if not ostatni.text.endswith(" "):
                ostatni.text += " "
            for bieg in akapit.runs:                       # przenosimy gotowe biegi,
                poprzedni._p.append(bieg._element)         # więc formaty zostają
            akapit._element.getparent().remove(akapit._element)
            scalone += 1
            continue

        if tekst:
            poprzedni = akapit
    return scalone


# --- wcięcia i tabulatory --------------------------------------------------------

def _popraw_wciecie(akapit) -> bool:
    """Wcięcie „pierwszego wiersza” zamienia na wcięcie całego akapitu.

    Zwisu (wartość ujemna, użytego w numerowanym spisie) nie rusza.
    """
    ustawienia = akapit.paragraph_format
    pierwszy = ustawienia.first_line_indent
    if pierwszy is None or pierwszy <= 0:
        return False
    ustawienia.left_indent = (ustawienia.left_indent or 0) + pierwszy
    ustawienia.first_line_indent = None
    return True


KROJE_PLIKI = (
    "C:/Windows/Fonts/calibri.ttf",
    "/usr/share/fonts/truetype/crosextra/Carlito-Regular.ttf",
)
KROJE_PLIKI_POGRUBIONE = (
    "C:/Windows/Fonts/calibrib.ttf",
    "/usr/share/fonts/truetype/crosextra/Carlito-Bold.ttf",
)
ODSTEP_KOLUMNY = Mm(4)          # prześwit między najdłuższą etykietą a wartością
PRZERWA_W_BLOKU = 2             # tyle pustych akapitów bloku jeszcze nie rozdziela


def _plik_kroju(pogrubiony: bool):
    for sciezka in (KROJE_PLIKI_POGRUBIONE if pogrubiony else KROJE_PLIKI):
        if Path(sciezka).exists():
            return sciezka
    return None


KOMUNIKAT_BRAK_KROJU = (
    "Nie znaleziono pliku czcionki Calibri ani Carlito, a bez pomiaru nie da się\n"
    "policzyć, gdzie postawić kolumnę wartości. Oszacowanie z liczby znaków wychodzi\n"
    "inne niż prawdziwy pomiar (zmierzone: przystanek 3781 → 4371 twipów, czyli ponad\n"
    "centymetr) i przestawiłoby układ w formatkach, których nikt nie prosił o zmianę.\n"
    "Na Linuksie: sudo apt install fonts-crosextra-carlito"
)


def _szerokosc(tekst: str, pogrubiony: bool = False) -> float:
    """Szerokość napisu w EMU, przy rozmiarze tekstu zwykłego.

    Mierzymy **prawdziwym plikiem czcionki** (Calibri albo metrycznie zgodne Carlito),
    bo od tego zależy, gdzie postawić kolumnę wartości. Bez czcionki nie zgadujemy:
    skrypt ma odmówić, a nie po cichu przesunąć bratu kolumny w całym operacie.
    """
    plik = _plik_kroju(pogrubiony)
    if plik is None:
        raise SystemExit(KOMUNIKAT_BRAK_KROJU)
    from PIL import ImageFont                                  # noqa: PLC0415
    czcionka = ImageFont.truetype(plik, 100)
    punkty = czcionka.getlength(tekst) / 100 * TRESC.pt
    return int(Emu(int(punkty * 12700)))


def _etykieta_i_wartosc(akapit) -> tuple[str, str] | None:
    """Rozbija „Etykieta:<tab>wartość” na części. `None`, gdy to nie taki akapit."""
    if "\t" not in akapit.text:
        return None
    etykieta, _, wartosc = akapit.text.partition("\t")
    if not etykieta.strip().endswith(":"):
        return None
    return etykieta.strip(), wartosc.replace("\t", " ").strip()


def _zostaw_jeden_tabulator(akapit) -> bool:
    """Ciąg tabulatorów → jeden. O tym, gdzie ląduje wartość, decyduje przystanek.

    Word trzyma każdy tabulator w osobnym biegu tekstu, więc `"\\t\\t" in bieg.text`
    bywa fałszem, choć w akapicie stoją obok siebie dwa — nadmiarowe trzeba usuwać,
    idąc przez cały akapit i pamiętając ostatni znak z poprzedniego biegu.
    """
    if "\t\t" not in akapit.text:
        return False
    po_tabulatorze = False
    for bieg in akapit.runs:
        zostawione = []
        for znak in bieg.text:
            if znak == "\t" and po_tabulatorze:
                continue
            po_tabulatorze = znak == "\t"
            zostawione.append(znak)
        bieg.text = "".join(zostawione)
    return True


def _bloki_etykieta_wartosc(dokument) -> list[list]:
    """Grupy sąsiadujących wierszy „Etykieta:<tab>wartość”.

    Puste akapity pomiędzy nie przerywają grupy — w bloku podpisu rozdzielają wiersze,
    a i tak należą do tego samego układu.
    """
    from docx.text.paragraph import Paragraph                  # noqa: PLC0415

    grupy, biezaca, puste_z_rzedu = [], [], 0

    def zamknij():
        nonlocal biezaca, puste_z_rzedu
        if biezaca:
            grupy.append(biezaca)
        biezaca, puste_z_rzedu = [], 0

    # Idziemy po **zawartości dokumentu**, nie po samych akapitach: tabela stojąca
    # między blokiem nagłówkowym a podpisem znaczy, że to dwa osobne układy. Pętla
    # po `dokument.paragraphs` w ogóle jej nie widzi i skleiłaby oba bloki w jeden.
    for element in dokument.element.body:
        if _nazwa(element) == "tbl":
            zamknij()
            continue
        if _nazwa(element) != "p":
            continue
        akapit = Paragraph(element, dokument)
        if _etykieta_i_wartosc(akapit):
            biezaca.append(akapit)
            puste_z_rzedu = 0
        elif not akapit.text.strip():
            puste_z_rzedu += 1
            if puste_z_rzedu > PRZERWA_W_BLOKU:      # duża przerwa = osobny układ
                zamknij()
        else:
            zamknij()
    zamknij()
    return [g for g in grupy if len(g) > 1]


def _wyrownaj_bloki(dokument) -> int:
    """Ustawia wartości bloku w jednej kolumnie, za najdłuższą etykietą.

    Rozjazd bierze się stąd, że każdy wiersz ma własną liczbę tabulatorów i własne
    wcięcie, a przystanek — jeśli w ogóle jest — bywa ustawiony bliżej niż kończy się
    najdłuższa etykieta. Wtedy tabulator ją przeskakuje i ląduje na przypadkowym
    przystanku domyślnym.

    **Nie ruszamy bloków, które są już równe** — jeśli wszystkie wiersze mają ten sam
    przystanek i mieści się za nim najdłuższa etykieta, to znaczy, że ktoś ustawił go
    świadomie w Wordzie.
    """
    sekcja = dokument.sections[0]
    szerokosc_tekstu = sekcja.page_width - sekcja.left_margin - sekcja.right_margin

    zmienione = 0
    for grupa in _bloki_etykieta_wartosc(dokument):
        przystanki = [[t.position for t in p.paragraph_format.tab_stops] for p in grupa]
        etykiety = [_etykieta_i_wartosc(p)[0] for p in grupa]
        najdluzsza = max(_szerokosc(e, True) for e in etykiety)

        rowny = (all(len(p) == 1 for p in przystanki)
                 and len({p[0] for p in przystanki}) == 1)
        if rowny:
            wciecia = {int(p.paragraph_format.left_indent or 0) for p in grupa}
            if len(wciecia) == 1 and przystanki[0][0] >= int(wciecia.pop()) + najdluzsza:
                continue                                    # już równo i z zapasem

        # wcięcie wspólne dla całej grupy: bierzemy najmniejsze, żeby nic nie uciekło
        # poza margines (wcięcie pierwszego wiersza jest już zamienione na wcięcie
        # akapitu, więc porównujemy jabłka z jabłkami)
        wciecie = min(int(p.paragraph_format.left_indent or 0) for p in grupa)
        kolumna = najdluzsza + int(ODSTEP_KOLUMNY)

        # blok ma się zmieścić na stronie: jeśli za kolumną nie starcza miejsca
        # na najdłuższą wartość, przesuwamy cały blok w lewo
        najdluzsza_wartosc = max(_szerokosc(_etykieta_i_wartosc(p)[1], True) for p in grupa)
        # z zapasem szerokości prześwitu: pomiar jest dokładny co do czcionki, ale
        # Word łamie wiersz odrobinę wcześniej, niż wynika z sumy szerokości znaków
        nadmiar = (wciecie + kolumna + najdluzsza_wartosc
                   + int(ODSTEP_KOLUMNY) - int(szerokosc_tekstu))
        if nadmiar > 0:
            wciecie = max(0, wciecie - nadmiar)

        for akapit in grupa:
            ustawienia = akapit.paragraph_format
            ustawienia.left_indent = Emu(wciecie)
            ustawienia.first_line_indent = None
            _zostaw_jeden_tabulator(akapit)
            ustawienia.tab_stops.clear_all()
            ustawienia.tab_stops.add_tab_stop(Emu(wciecie + kolumna))
            zmienione += 1
    return zmienione


# --- typografia ------------------------------------------------------------------

# Kolejność dzieci w OOXML jest częścią schematu, nie kwestią gustu: Word odmawia
# otwarcia pliku, w którym element stoi nie tam, gdzie trzeba, a LibreOffice otwiera
# go bez mrugnięcia. Dlatego wszystko dokładamy **przed** pierwszym elementem, który
# wg schematu ma iść po nim — i na końcu sprawdzamy to jeszcze raz (`_sprawdz_kolejnosc`).
PO_ROZSTRZELENIU = ("w", "kern", "position", "sz", "szCs", "highlight", "u", "effect",
                    "bdr", "shd", "fitText", "vertAlign", "rtl", "cs", "em", "lang",
                    "eastAsianLayout", "specVanish", "oMath")
PO_OBLEWANIU = ("docPr", "cNvGraphicFramePr", "graphic", "sizeRelH", "sizeRelV")
# `w:tabs` w `w:pPr` stoi zaraz za `w:pBdr` i `w:shd`
PO_PRZYSTANKACH = ("suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
                   "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi",
                   "adjustRightInd", "snapToGrid", "spacing", "ind",
                   "contextualSpacing", "jc", "textAlignment", "outlineLvl",
                   "rPr", "sectPr")

# Te sposoby oblewania wymagają atrybutu `wrapText` — bez niego Word odrzuca plik.
OBLEWANIE_Z_TEKSTEM = ("wrapSquare", "wrapTight", "wrapThrough")


def _nazwa(element) -> str:
    return element.tag.split("}")[-1]


def _wstaw_przed(rodzic, element, nastepcy: tuple[str, ...]) -> None:
    for dziecko in rodzic:
        if _nazwa(dziecko) in nastepcy:
            dziecko.addprevious(element)
            return
    rodzic.append(element)


def _rozstrzel(bieg, wartosc: int) -> None:
    rPr = bieg._element.get_or_add_rPr()
    for stary in rPr.findall(qn("w:spacing")):
        rPr.remove(stary)
    if wartosc:
        element = OxmlElement("w:spacing")
        element.set(qn("w:val"), str(wartosc))
        _wstaw_przed(rPr, element, PO_ROZSTRZELENIU)


def _ustaw(akapit, rozmiar, pogrubienie, *, rozstrzelenie=0) -> None:
    """`pogrubienie=None` znaczy „zostaw, jak jest w formatce”.

    Podział ról: skrypt odpowiada za hierarchię (co jest tytułem, nagłówkiem,
    etykietą), a autor formatki za akcenty wewnątrz tekstu — to on wie, która
    wartość jest w danym dokumencie ważna. Wymuszanie grubości w treści kasowało
    pogrubienia dokładane w Wordzie i wracały przy każdym uruchomieniu skryptu.
    """
    for bieg in akapit.runs:
        bieg.font.name = KROJ
        bieg.font.size = rozmiar
        # czerwień to numer roboty — zostaje czerwona i zawsze pogrubiona, bo to
        # jedyna rzecz w dokumencie, której szuka się wzrokiem
        if _czerwony(bieg):
            bieg.bold = True
        elif pogrubienie is not None:
            bieg.bold = pogrubienie
        bieg.underline = False
        _rozstrzel(bieg, rozstrzelenie)


# --- logo ------------------------------------------------------------------------

def _wlasny_akapit(dokument, akapit):
    """Przenosi logo do własnego, pustego akapitu tuż przed bieżącym.

    Kotwica wrzucona w akapit z tekstem (tak było w sprawozdaniu — logo siedziało
    w akapicie tytułu) potrafi wypchnąć ten tekst *nad* obrazek. Logo ma stać
    najwyżej, więc dostaje własny akapit.
    """
    if not akapit.text.strip():
        return akapit
    nowy = akapit.insert_paragraph_before()
    for bieg in akapit.runs:
        if bieg._element.findall(".//{*}anchor"):
            nowy._p.append(bieg._element)
    return nowy


def _oblewanie(dokument, akapit_logo) -> str:
    """„Z boku” tylko wtedy, gdy obok logo ma naprawdę coś stanąć.

    Rozpoznajemy to po najbliższym niepustym akapicie: wyrównany do prawej to blok
    firmówki (data, numery) — ten ma stać obok logo. Cokolwiek innego, na przykład
    wyśrodkowany tytuł, ma iść pod spodem, więc obrazek zajmuje własny pas.
    """
    # porównujemy po elemencie XML: `dokument.paragraphs` przy każdym wywołaniu
    # tworzy nowe obiekty opakowujące, więc porównanie po tożsamości nie zadziała
    akapity = [p._p for p in dokument.paragraphs]
    for dalszy in dokument.paragraphs[akapity.index(akapit_logo._p) + 1:]:
        if dalszy.text.strip():
            return ("wrapSquare" if dalszy.alignment == WD_ALIGN_PARAGRAPH.RIGHT
                    else "wrapTopAndBottom")
    return "wrapTopAndBottom"


def _ustaw_logo(dokument) -> int:
    """Ta sama szerokość i to samo miejsce w każdym dokumencie.

    Wysokość liczymy z proporcji obrazka, żeby logo się nie rozjechało.
    """
    zmienione = 0
    for akapit in list(dokument.paragraphs):
        if not akapit._p.findall(".//{*}anchor"):
            continue
        akapit = _wlasny_akapit(dokument, akapit)
        for kotwica in akapit._p.findall(".//{*}anchor"):
            rozmiar = kotwica.find(qn("wp:extent"))
            szer, wys = int(rozmiar.get("cx")), int(rozmiar.get("cy"))
            nowa_szer = int(LOGO_SZEROKOSC)
            nowa_wys = int(round(wys * nowa_szer / szer))

            rozmiar.set("cx", str(nowa_szer))
            rozmiar.set("cy", str(nowa_wys))
            # Rozmiar obrazka siedzi w <a:xfrm><a:ext>. UWAGA: <a:ext> z atrybutem
            # `uri` to zupełnie co innego — element listy rozszerzeń, przypadkiem
            # o tej samej nazwie. Dopisanie mu cx/cy daje plik, którego Word
            # nie otworzy (a LibreOffice otworzy bez słowa).
            for xfrm in kotwica.findall(".//" + qn("a:xfrm")):
                for ext in xfrm.findall(qn("a:ext")):
                    ext.set("cx", str(nowa_szer))
                    ext.set("cy", str(nowa_wys))

            for os_, wartosc in ((qn("wp:positionH"), LOGO_OD_LEWEJ),
                                 (qn("wp:positionV"), LOGO_OD_GORY)):
                polozenie = kotwica.find(os_)
                polozenie.set("relativeFrom", "margin")
                for align in polozenie.findall(qn("wp:align")):  # „przy krawędzi”
                    polozenie.remove(align)                      # zastępujemy liczbą
                przesuniecie = polozenie.find(qn("wp:posOffset"))
                if przesuniecie is None:
                    przesuniecie = OxmlElement("wp:posOffset")
                    polozenie.append(przesuniecie)
                przesuniecie.text = str(int(wartosc))

            zadany = _oblewanie(dokument, akapit)
            for stare in [e for e in kotwica if _nazwa(e).startswith("wrap")]:
                kotwica.remove(stare)
            nowe = OxmlElement(f"wp:{zadany}")
            if zadany in OBLEWANIE_Z_TEKSTEM:      # atrybut wymagany przez schemat
                nowe.set("wrapText", "bothSides")
            _wstaw_przed(kotwica, nowe, PO_OBLEWANIU)
            zmienione += 1
    return zmienione


# --- całość ----------------------------------------------------------------------

def _stopka_biegi(akapit) -> None:
    for bieg in akapit.runs:
        bieg.font.name = KROJ
        bieg.font.size = STOPKA_ROZMIAR
        bieg.bold = False
        bieg.font.color.rgb = STOPKA_KOLOR
        rPr = bieg._element.get_or_add_rPr()
        for caps in rPr.findall(qn("w:caps")):     # wersaliki bywają też na biegu
            rPr.remove(caps)


def _stopka_wzorcowa() -> list | None:
    """Akapity stopki z pliku wzorcowego — jedno źródło prawdy dla wszystkich formatek.

    Każda formatka przychodzi z własną kopią firmówki i drobne różnice same się w nich
    zalęgają: raz kreska jest obramowaniem akapitu, raz wklejonym obrazkiem, w jednej
    nazwa ulicy jest z wielkiej litery, w innej z małej. Zamiast poprawiać to w kółko
    ręcznie, każdy dokument dostaje stopkę przepisaną z `spis_tresci_wzor.docx`.
    """
    wzorzec = SZABLONY / STOPKA_WZORCOWA
    if not wzorzec.exists():
        return None
    dokument = docx.Document(wzorzec)
    for sekcja in dokument.sections:
        akapity = [p for p in sekcja.first_page_footer.paragraphs if p.text.strip()]
        # przepisujemy tylko czysty tekst: obrazek niósłby ze sobą powiązanie z paczki
        # źródłowej, którego w pliku docelowym nie ma
        if akapity and not any(p._p.findall(".//{*}drawing") for p in akapity):
            return [copy.deepcopy(p._p) for p in akapity]
    return None


def _przystanki_stopki(dokument) -> None:
    """Rozkłada trzy kolumny stopki na pełną szerokość **tej** strony.

    Stopka jest jedna dla wszystkich formatek, a przystanki miała wpisane na sztywno
    (80 mm na środek, 160 mm do prawej) — czyli pod kartkę pionową. Na poziomej,
    gdzie tekst ma 248 mm, dane kończyły się w dwóch trzecich szerokości i wyglądało
    to, jakby ktoś wkleił stopkę z innego dokumentu. Liczymy je więc z rozmiaru strony.
    """
    sekcja = dokument.sections[0]
    szerokosc = int(sekcja.page_width) - int(sekcja.left_margin) - int(sekcja.right_margin)
    polowa = szerokosc // 2

    for styl in dokument.styles:
        if styl.style_id not in ("Stopka", "Footer"):
            continue
        pPr = styl.element.find(qn("w:pPr"))
        if pPr is None:
            continue
        przystanki = pPr.find(qn("w:tabs"))
        if przystanki is None:
            przystanki = OxmlElement("w:tabs")
            _wstaw_przed(pPr, przystanki, PO_PRZYSTANKACH)
        for stary in list(przystanki):
            przystanki.remove(stary)
        for wyrownanie, pozycja in (("center", polowa), ("right", szerokosc)):
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), wyrownanie)
            tab.set(qn("w:pos"), str(int(pozycja / 914400 * 1440)))    # EMU -> twipy
            przystanki.append(tab)


def _ujednolic_stopke(dokument, wzorcowe=None) -> int:
    """Jedna stopka: na każdej stronie, w każdym dokumencie, bez numeracji stron.

    Numer strony siedział w zwykłej stopce (stopka pierwszej strony miała dane firmy).
    Skoro obie mają wyglądać tak samo, zwykła dostaje kopię tej pierwszej — a pole
    z numerem znika razem z jej dotychczasową zawartością.
    """
    for styl in dokument.styles:
        if styl.style_id not in ("Stopka", "Footer"):
            continue
        rPr = styl.element.find(qn("w:rPr"))
        if rPr is not None:
            for caps in rPr.findall(qn("w:caps")):   # to od nich stopka krzyczała
                rPr.remove(caps)
        pPr = styl.element.find(qn("w:pPr"))
        krawedzie = pPr.find(qn("w:pBdr")) if pPr is not None else None
        if krawedzie is not None:
            gora = krawedzie.find(qn("w:top"))
            if gora is not None:                     # gruba kreska -> włos
                gora.set(qn("w:sz"), "2")
                gora.set(qn("w:color"), "BFBFBF")

    _przystanki_stopki(dokument)

    zmienione = 0
    for sekcja in dokument.sections:
        if wzorcowe:
            wzorcowa_stopka = wzorcowe
        else:
            wlasne = [p for p in sekcja.first_page_footer.paragraphs if p.text.strip()]
            if not wlasne:
                continue
            for akapit in wlasne:
                _stopka_biegi(akapit)
            wzorcowa_stopka = [copy.deepcopy(p._p) for p in wlasne]

        # wszystkie trzy stopki dostają to samo: pierwsza strona, dalsze i parzyste.
        # Ta ostatnia nie jest teraz używana, ale trzymała pole z numerem strony
        # i odezwałaby się przy pierwszej zmianie ustawień
        for cel in (sekcja.first_page_footer, sekcja.footer, sekcja.even_page_footer):
            for akapit in list(cel.paragraphs):
                akapit._p.getparent().remove(akapit._p)
            for akapit in wzorcowa_stopka:
                cel._element.append(copy.deepcopy(akapit))
            zmienione += 1
    return zmienione


def _ujednolic_tabele(dokument) -> int:
    """Komórki tabel: ten sam krój, ale **rozmiar i grubość zostają**.

    Formularz z kilkunastoma wierszami bywa złożony 7-punktową czcionką, żeby zmieścić
    się na stronie — narzucenie mu rozmiaru tekstu zwykłego rozsypałoby tabelę.
    Grubość w tabeli też niesie znaczenie (nagłówek kolumny), a nie sam wygląd.
    """
    komorki = 0
    for tabela in dokument.tables:
        for wiersz in tabela.rows:
            for komorka in wiersz.cells:
                for akapit in komorka.paragraphs:
                    for bieg in akapit.runs:
                        bieg.font.name = KROJ
                komorki += 1
    return komorki


def _sprawdz_kolejnosc(dokument) -> list[str]:
    """Sprawdza, czy dołożone elementy stoją tam, gdzie wymaga tego schemat OOXML.

    Bez tego można w dobrej wierze zapisać plik, który otwiera się na Linuksie
    i **nie otwiera się w Wordzie** — a to wychodzi dopiero u brata. Kosztowało to
    już jedno wydanie, więc niech pilnuje tego kod, a nie pamięć.
    """
    zarzuty = []
    for kotwica in dokument.element.body.findall(".//{*}anchor"):
        nazwy = [_nazwa(e) for e in kotwica]
        oblewania = [n for n in nazwy if n.startswith("wrap")]
        for oblewanie in oblewania:
            pozniejsze = [n for n in PO_OBLEWANIU if n in nazwy]
            if pozniejsze and nazwy.index(oblewanie) > nazwy.index(pozniejsze[0]):
                zarzuty.append(f"<wp:{oblewanie}> stoi za <wp:{pozniejsze[0]}>")

        for oblewanie in oblewania:
            if oblewanie in OBLEWANIE_Z_TEKSTEM:
                element = next(e for e in kotwica if _nazwa(e) == oblewanie)
                if not any(k.endswith("wrapText") for k in element.attrib):
                    zarzuty.append(f"<wp:{oblewanie}> bez wymaganego wrapText")

    for rPr in dokument.element.body.findall(".//{*}rPr"):
        nazwy = [_nazwa(e) for e in rPr]
        if "spacing" not in nazwy:
            continue
        pozniejsze = [n for n in PO_ROZSTRZELENIU if n in nazwy]
        if pozniejsze and nazwy.index("spacing") > nazwy.index(pozniejsze[0]):
            zarzuty.append(f"<w:spacing> stoi za <w:{pozniejsze[0]}>")

    # <a:ext> z atrybutem `uri` to element listy rozszerzeń, a nie rozmiar — cx/cy
    # nie mają tam prawa być
    for ext in dokument.element.body.findall(".//{*}ext"):
        if ext.get("uri") and ("cx" in ext.attrib or "cy" in ext.attrib):
            zarzuty.append("<a:ext uri=…> ma cx/cy, a to element rozszerzenia")
    return zarzuty


def ujednolic(plik: Path, stopka_wzorcowa=None) -> dict[str, int]:
    # Sprawdzamy krój **przed** otwarciem dokumentu, a nie dopiero przy pierwszym
    # pomiarze: inaczej plik bez bloków „Etykieta: wartość” przeszedłby bez słowa,
    # a następny w kolejce padłby w połowie serii.
    if _plik_kroju(False) is None or _plik_kroju(True) is None:
        raise SystemExit(KOMUNIKAT_BRAK_KROJU)

    dokument = docx.Document(plik)
    licznik = dict.fromkeys(
        ("tytul", "naglowek", "etykieta", "tresc", "podpis", "pusty",
         "scalone", "wciecia", "kolumny", "logo", "komorki",
         "stopka"), 0)

    # styl bazowy: cokolwiek dopiszesz później w Wordzie, wyjdzie w tym samym kroju
    normalny = dokument.styles["Normal"]
    normalny.font.name = KROJ
    normalny.font.size = TRESC

    licznik["scalone"] = _scal_ciagi(dokument)
    licznik["logo"] = _ustaw_logo(dokument)
    licznik["komorki"] = _ujednolic_tabele(dokument)
    # plik wzorcowy normalizuje własną stopkę; reszta dostaje jego kopię
    licznik["stopka"] = _ujednolic_stopke(
        dokument, None if plik.name == STOPKA_WZORCOWA else stopka_wzorcowa)

    # Wcięcia muszą być poprawione **przed** wyrównaniem bloków: dopóki wiersz ma
    # wcięcie „pierwszego wiersza”, jego lewa krawędź jest inna niż się wydaje.
    # Tabulatorów poza wyrównywanymi blokami **nie ruszamy**. Ciąg tabulatorów bywa
    # świadomym wcięciem — brat wsuwa nimi wiersz z tolerancjami pod kolumnę wartości
    # („[dl – 0.02 m] / [dh – 0.03 m]” pod opisem punktu, pięć tabulatorów). Zwijanie
    # ich do jednego przesuwało ten wiersz na lewo przy każdym uruchomieniu skryptu.
    # Wewnątrz bloku to co innego: tam sami stawiamy przystanek, więc jeden tabulator
    # jest dokładnie tym, czego trzeba — i tam zwijanie zostaje.
    for akapit in dokument.paragraphs:
        if akapit.text.strip():
            licznik["wciecia"] += _popraw_wciecie(akapit)
    licznik["kolumny"] = _wyrownaj_bloki(dokument)

    akapity = dokument.paragraphs
    z_trescia = [p for p in akapity if p.text.strip()]
    najwiekszy = max((_rozmiar(p) for p in z_trescia), default=0)
    ostatni = z_trescia[-1] if z_trescia else None
    tytuly = [p._p for p in z_trescia
              if _rozmiar(p) == najwiekszy and najwiekszy > TRESC.pt]

    for akapit in akapity:
        tekst = akapit.text.strip()

        if not tekst:
            # Pusty akapit w ramce (`w:framePr`) to nie odstęp między blokami, tylko
            # podkładka podnosząca przypięty podpis — patrz narzedzia/popraw_szablon.py.
            # Zmiana jego rozmiaru przesunęłaby podpis, więc go nie ruszamy.
            if akapit._p.find(qn("w:pPr") + "/" + qn("w:framePr")) is None:
                for bieg in akapit.runs:
                    bieg.font.name, bieg.font.size = KROJ, PUSTY
                licznik["pusty"] += 1
            continue

        if _rozmiar(akapit) == najwiekszy and najwiekszy > TRESC.pt:
            _ustaw(akapit, TYTUL, True, rozstrzelenie=TYTUL_ROZSTRZELENIE)
            akapit.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Tytuł bywa złożony z kilku akapitów („WYKAZ ZMIAN DANYCH EWIDENCYJNYCH”
            # / „DOTYCZĄCYCH BUDYNKU”). Odstęp należy się całemu blokowi, a nie każdemu
            # wierszowi z osobna — inaczej tytuł rozjeżdża się w pionie.
            akapit.paragraph_format.space_before = (
                ODSTEP_NAD_TYTULEM if akapit._p is tytuly[0] else Pt(0))
            akapit.paragraph_format.space_after = (
                ODSTEP_POD_TYTULEM if akapit._p is tytuly[-1] else Pt(0))
            akapit.paragraph_format.keep_with_next = True
            licznik["tytul"] += 1
        elif _rozmiar(akapit) > TRESC.pt:
            _ustaw(akapit, NAGLOWEK, True)
            akapit.paragraph_format.keep_with_next = True
            licznik["naglowek"] += 1
        elif akapit is ostatni and PODPIS in tekst:
            _ustaw(akapit, TRESC, None)
            akapit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            akapit.paragraph_format.left_indent = None
            akapit.paragraph_format.first_line_indent = None
            licznik["podpis"] += 1
        elif tekst.endswith(":"):
            # Grubości etykiety też nie narzucamy: brat pisze etykiety zwykłym pismem,
            # a pogrubia wartości — i to on ma o tym decydować. Rola „etykieta” zostaje,
            # bo niesie coś innego: etykieta bez swojej wartości na dole strony wygląda
            # na urwany dokument.
            _ustaw(akapit, TRESC, None)
            akapit.paragraph_format.keep_with_next = True
            licznik["etykieta"] += 1
        else:
            _ustaw(akapit, TRESC, None)
            licznik["tresc"] += 1

    zarzuty = _sprawdz_kolejnosc(dokument)
    if zarzuty:
        raise SystemExit(f"{plik.name}: nie zapisuję, bo Word odrzuci taki plik — "
                         + "; ".join(sorted(set(zarzuty))))

    dokument.save(plik)
    return licznik


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("szablony", nargs="*", help="pliki .docx; domyślnie wszystkie wzory")
    argumenty = parser.parse_args()

    pliki = [Path(s) for s in argumenty.szablony] or sorted(SZABLONY.glob("*_wzor.docx"))

    # plik wzorcowy przepuszczamy najpierw, żeby reszta dostała już poprawioną stopkę
    wzorcowy = SZABLONY / STOPKA_WZORCOWA
    if wzorcowy.exists() and not any(p.name == STOPKA_WZORCOWA for p in pliki):
        ujednolic(wzorcowy)
    wzorcowa = _stopka_wzorcowa()
    for plik in pliki:
        if not plik.exists():
            print(f"Nie ma pliku {plik}")
            continue
        podsumowanie = ujednolic(plik, wzorcowa)
        print(f"{plik.name}: "
              + ", ".join(f"{n}: {i}" for n, i in podsumowanie.items() if i))
