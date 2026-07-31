"""Tworzy szkielet sprawozdania technicznego (szablony/sprawozdanie_techniczne_wzor.docx).

Sprawozdanie dokłada się do operatu checkboxem w formularzu i korzysta z tych samych
danych co spis treści — nie ma własnych pól poza opisem przebiegu i listą baz danych.

**Teksty w tym pliku to zaślepki do podmiany.** Brat wkleja swoją formatkę i przenosi
do niej znaczniki; tutaj chodzi tylko o to, żeby miał gotowe nazwy i działającą logikę
warunków.

Pułapka docxtpl: akapit ze znacznikiem `{%p ... %}` znika w całości, więc warunek,
treść i `endif` muszą być w **osobnych akapitach** — inaczej treść przepada razem
ze znacznikiem.

    python narzedzia/utworz_wzor_sprawozdania.py
"""
import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from app.config import SZABLONY  # noqa: E402

NAZWA = "sprawozdanie_techniczne_wzor.docx"


def akapit(dokument, tekst, rozmiar=11, pogrubienie=False, wyrownanie=None, odstep_przed=0):
    element = dokument.add_paragraph()
    element.paragraph_format.space_before = Pt(odstep_przed)
    bieg = element.add_run(tekst)
    bieg.bold = pogrubienie
    bieg.font.size = Pt(rozmiar)
    if wyrownanie is not None:
        element.alignment = wyrownanie
    return element


def zbuduj() -> Path:
    dokument = Document()
    styl = dokument.styles["Normal"]
    styl.font.name = "Times New Roman"
    styl.font.size = Pt(11)
    SRODEK = WD_ALIGN_PARAGRAPH.CENTER

    akapit(dokument, "SPRAWOZDANIE TECHNICZNE", 16, pogrubienie=True, wyrownanie=SRODEK,
           odstep_przed=24)
    akapit(dokument, "do operatu nr {{ nr_operatu }}", 11, wyrownanie=SRODEK)

    akapit(dokument, "Nr roboty: {{ nr_roboty }}", 11, odstep_przed=24)
    akapit(dokument, "Rodzaj pracy: {{ rodzaj_pracy }}", 11)
    akapit(dokument, "Jednostka ewidencyjna: [{{ polozenie_gmina_teryt }}] "
                     "{{ polozenie_gmina }}", 11)
    akapit(dokument, "Obręb: [{{ polozenie_obreb_teryt }}] {{ polozenie_obreb }}", 11)
    akapit(dokument, "Działki nr: {{ nr_dzialki }}", 11)
    akapit(dokument, "Prace wykonano w okresie {{ data_zgloszenia }} – "
                     "{{ data_zakonczenia }}.", 11)

    # --- 1. przebieg prac -----------------------------------------------------
    akapit(dokument, "1. Przebieg wykonanych prac", 12, pogrubienie=True, odstep_przed=20)
    akapit(dokument, "{%p if opis_przebiegu_jest %}")
    akapit(dokument, "W ramach zgłoszonej pracy geodezyjnej wykonano: "
                     "{{ opis_przebiegu }}")
    akapit(dokument, "{%p else %}")
    akapit(dokument, "W ramach zgłoszonej pracy geodezyjnej wykonano: brak")
    akapit(dokument, "{%p endif %}")

    # --- 2. zmiany w bazach danych --------------------------------------------
    akapit(dokument, "2. Zmiany w bazach danych", 12, pogrubienie=True, odstep_przed=20)
    akapit(dokument, "{%p if opis_przebiegu_jest and bazy %}")
    akapit(dokument, "Wyniki pracy geodezyjnej powodują zmiany w bazach danych: "
                     "{{ bazy_pliki | join(', ') }}")
    akapit(dokument, "{%p else %}")
    akapit(dokument, "Wyniki pracy geodezyjnej powodują zmiany w bazach danych: brak")
    akapit(dokument, "{%p endif %}")

    akapit(dokument, "{{ data_zakonczenia }}", 11, odstep_przed=36)

    plik = SZABLONY / NAZWA
    dokument.save(plik)
    return plik


OPIS_POL = {
    "nazwa": "Sprawozdanie techniczne",
    "opis": "Dokłada się do operatu. Korzysta z danych wpisanych przy spisie treści.",
    "pola": [],          # własnych pól nie ma — wszystko przychodzi z formularza operatu
}


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--nadpisz", action="store_true",
                        help="nadpisz istniejący .docx (UWAGA: skasuje prawdziwą formatkę)")
    argumenty = parser.parse_args()

    docelowy = SZABLONY / NAZWA
    if docelowy.exists() and not argumenty.nadpisz:
        print(f"{docelowy} już istnieje — zostawiam bez zmian.")
        print("Nadpisanie zaślepką: --nadpisz")
    else:
        print("Utworzono:", zbuduj())

    opis = docelowy.with_suffix(".json")
    opis.write_text(json.dumps(OPIS_POL, ensure_ascii=False, indent=2) + "\n",
                    encoding="utf-8")
    print("Zapisano opis pól:", opis)
