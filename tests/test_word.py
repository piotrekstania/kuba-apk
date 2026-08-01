"""Ścieżka przez Microsoft Word (COM) — jedyna część, której nie sprawdzi CI.

Na runnerze GitHuba nie ma Worda, więc `.github/workflows/testy.yml` pomija marker
`word`. Te testy uruchamia się **ręcznie na Windowsie z Office** przed każdym wydaniem
dotykającym `app/pdf.py`:

    .venv\\Scripts\\pytest -m word -v

Nie puszczaj ich równolegle (`-n auto`) — Word to jedna aplikacja na komputerze
i równoległość testowałaby wtedy samą siebie, a nie kod.
"""
from __future__ import annotations

import subprocess
import sys
import time
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter

from app import operaty, pdf

pytestmark = [
    pytest.mark.word,
    pytest.mark.skipif(sys.platform != "win32", reason="ścieżka wordowa tylko na Windowsie"),
    pytest.mark.skipif(pdf.dostepny_konwerter() != "word", reason="Word niewykryty"),
]


# --- pomocnicze --------------------------------------------------------------

def dokument(sciezka: Path, tekst: str = "Operat techniczny 001/2026",
             naglowek: str | None = None) -> Path:
    plik = Document()
    if naglowek:
        plik.add_heading(naglowek, level=1)
    plik.add_paragraph(tekst)
    sciezka.parent.mkdir(parents=True, exist_ok=True)
    plik.save(sciezka)
    return sciezka


def tekst_pdf(plik: Path) -> str:
    return "".join(strona.extract_text() for strona in PdfReader(str(plik)).pages)


def liczba_wordow() -> int:
    wynik = subprocess.run(["tasklist", "/FI", "IMAGENAME eq WINWORD.EXE"],
                           capture_output=True, text=True)
    return wynik.stdout.count("WINWORD.EXE")


def pusty_pdf(cel: Path) -> None:
    zapis = PdfWriter()
    zapis.add_blank_page(width=595, height=842)
    Path(cel).parent.mkdir(parents=True, exist_ok=True)
    with open(cel, "wb") as wyjscie:
        zapis.write(wyjscie)


def wybuchnij(*_args, **_kwargs):
    raise RuntimeError("Word czeka z otwartym oknem dialogowym")


# --- sama konwersja ----------------------------------------------------------

def test_konwersja_pojedyncza(tmp_path):
    """PDF powstaje, ma treść dokumentu i nie zostawia po sobie pliku roboczego."""
    zrodlo = dokument(tmp_path / "operat.docx")

    wynik = pdf.docx_na_pdf(zrodlo, tmp_path / "operat.pdf")

    assert wynik.exists()
    assert "001/2026" in tekst_pdf(wynik)
    # Word pisze PDF kawałkami do pliku `*.czesciowy` i dopiero gotowy podmienia
    # na docelowy. Gdyby roboczy został, miniatury czytałyby urwany dokument.
    assert list(tmp_path.glob("*.czesciowy")) == []


def test_nie_zostaje_wiszacy_word(tmp_path):
    """Bez `Quit()` w `finally` procesy WINWORD.EXE mnożą się w tle po każdej robocie.

    Objaw u brata: komputer zwalnia, a Word „nie chce się otworzyć”.
    """
    przed = liczba_wordow()

    pdf.docx_na_pdf(dokument(tmp_path / "a.docx"), tmp_path / "a.pdf")

    time.sleep(1.5)                     # Word znika z listy procesów chwilę po Quit()
    assert liczba_wordow() <= przed


def test_konwersja_z_watku_roboczego(tmp_path):
    """Główny powód, dla którego nie używamy docx2pdf.

    Trasy `/scal` i miniatury wykonują się w puli wątków FastAPI, a tam bez
    `pythoncom.CoInitialize()` leci „CoInitialize has not been called”.
    """
    zrodlo = dokument(tmp_path / "watek.docx")

    with ThreadPoolExecutor(max_workers=1) as pula:
        wynik = pula.submit(pdf.docx_na_pdf, zrodlo, tmp_path / "watek.pdf").result()

    assert wynik.exists()
    assert "001/2026" in tekst_pdf(wynik)


def test_dwie_konwersje_naraz(tmp_path):
    """Pilnuje `_BLOKADA_KONWERSJI` — bez niej dwie instancje Worda zamykały sobie sesję."""
    zrodla = [dokument(tmp_path / f"r{i}.docx", f"Dokument numer {i}") for i in range(2)]

    with ThreadPoolExecutor(max_workers=2) as pula:
        wyniki = list(pula.map(lambda z: pdf.docx_na_pdf(z, z.with_suffix(".pdf")), zrodla))

    for indeks, wynik in enumerate(wyniki):
        assert wynik.exists()
        assert f"Dokument numer {indeks}" in tekst_pdf(wynik)      # żaden nie jest urwany


def test_wsad_jednym_uruchomieniem(tmp_path):
    """Wsad ma być wyraźnie szybszy niż konwersje po kolei — najdroższy jest start Worda.

    Asercja z dużym zapasem: to ma łapać regresję „wróciliśmy do startu na dokument”,
    a nie mierzyć wydajność.
    """
    osobno = [dokument(tmp_path / "osobno" / f"d{i}.docx", f"Dokument {i}") for i in range(4)]
    wsadowo = [dokument(tmp_path / "wsad" / f"d{i}.docx", f"Dokument {i}") for i in range(4)]

    start = time.perf_counter()
    for zrodlo in osobno:
        pdf.docx_na_pdf(zrodlo, zrodlo.with_suffix(".pdf"))
    czas_osobno = time.perf_counter() - start

    start = time.perf_counter()
    zrobione = pdf.docx_na_pdf_wsad([(z, z.with_suffix(".pdf")) for z in wsadowo])
    czas_wsadu = time.perf_counter() - start

    assert len(zrobione) == 4
    for indeks, zrodlo in enumerate(wsadowo):
        assert f"Dokument {indeks}" in tekst_pdf(zrodlo.with_suffix(".pdf"))
    assert czas_wsadu < czas_osobno * 0.8, (
        f"wsad {czas_wsadu:.2f} s, osobno {czas_osobno:.2f} s — "
        "wygląda, jakby Word startował do każdego dokumentu z osobna")


def test_zakladki_z_naglowkow(tmp_path):
    """Pilnuje `CreateBookmarks` — inaczej nikt nie zauważy braku, aż ktoś otworzy operat."""
    zrodlo = dokument(tmp_path / "z_naglowkiem.docx", naglowek="Sprawozdanie techniczne")

    wynik = pdf.docx_na_pdf(zrodlo, tmp_path / "z_naglowkiem.pdf")

    assert PdfReader(str(wynik)).outline


# --- awaria Worda ------------------------------------------------------------

def test_awaria_worda_schodzi_na_libreoffice(tmp_path, monkeypatch):
    """Gdy Word padnie, a obok jest LibreOffice — PDF ma i tak powstać.

    LibreOffice jest tu udawany: na maszynie docelowej (i u brata) go nie ma,
    a sprawdzamy rozgałęzienie w `_konwertuj`, nie samego LibreOffice'a.
    """
    zrodlo = dokument(tmp_path / "awaria.docx")
    monkeypatch.setattr(pdf, "_wordem_wsad", wybuchnij)
    monkeypatch.setattr(pdf, "sciezka_libreoffice", lambda: r"C:\udawany\soffice.exe")
    monkeypatch.setattr(pdf, "_konwersja_libreoffice",
                        lambda zrodlo, cel: pusty_pdf(cel))

    assert pdf.docx_na_pdf(zrodlo, tmp_path / "awaria.pdf").exists()


def test_awaria_worda_bez_libreoffice_mowi_po_polsku(tmp_path, monkeypatch):
    """Bez LibreOffice'a: po polsku i z podpowiedzią, że Word pewnie czeka z oknem."""
    zrodlo = dokument(tmp_path / "awaria2.docx")
    monkeypatch.setattr(pdf, "_wordem_wsad", wybuchnij)
    monkeypatch.setattr(pdf, "sciezka_libreoffice", lambda: None)

    with pytest.raises(pdf.BrakKonwertera) as awaria:
        pdf.docx_na_pdf(zrodlo, tmp_path / "awaria2.pdf")

    komunikat = str(awaria.value)
    assert "Word" in komunikat and "oknem" in komunikat
    assert "Traceback" not in komunikat


# --- cała ścieżka przez HTTP -------------------------------------------------

@pytest.fixture
def klient_z_wordem(srodowisko, monkeypatch):
    """TestClient **bez** atrapy konwersji — tu chodzi o prawdziwego Worda.

    Fixture `klient` z conftest.py podmienia konwersję na atrapę, więc do tego
    testu się nie nadaje.
    """
    from fastapi.testclient import TestClient

    from app import main, teryt

    monkeypatch.setattr(teryt, "pusto", lambda: False)
    monkeypatch.setattr(main.teryt, "pusto", lambda: False)
    with TestClient(main.app, raise_server_exceptions=False) as klient_testowy:
        klient_testowy.srodowisko = srodowisko
        yield klient_testowy


def test_trasa_scal_prawdziwym_wordem(klient_z_wordem):
    """Operat z dwoma dokumentami → POST /scal → sklejony PDF o nazwie numeru roboty.

    Sprawdza całą ścieżkę razem z wątkami FastAPI: konwersję obu dokumentów Wordem,
    kolejność sklejania i nazwę pliku wymaganą przez ośrodek.
    """
    katalog, _ = operaty.zaloz("001/2026", "GK.6640.123.2026", "spis_tresci_wzor", {})
    dokument(katalog / "spis_tresci.docx", "Spis tresci operatu")
    dokument(katalog / "sprawozdanie.docx", "Sprawozdanie techniczne roboty")

    # Lista wartości pod jednym kluczem, a nie lista krotek: httpx bierze listę krotek
    # za surową treść żądania i formularz dociera do trasy pusty.
    odpowiedz = klient_z_wordem.post(
        f"/scal/{katalog.name}",
        data={"plik": ["spis_tresci.docx", "sprawozdanie.docx"]},
        follow_redirects=False)

    assert odpowiedz.status_code == 303
    assert "blad=" not in odpowiedz.headers["location"]

    wynik = katalog / operaty.nazwa_wyniku(katalog)
    assert wynik.name == "GK.6640.123.2026.pdf"          # dokładnie numer roboty (KERG)
    assert wynik.exists()
    tresc = tekst_pdf(wynik)
    assert "Spis tresci operatu" in tresc
    assert "Sprawozdanie techniczne roboty" in tresc
