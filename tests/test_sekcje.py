"""Powtarzalne sekcje formularza — wykazy zmian danych budynku.

W jednym operacie takich wykazów bywa kilka, każdy z własnym kompletem 15 atrybutów
w dwóch stanach. To za dużo na wiersz tabeli (30 kolumn), więc powtarzamy **komplet
pól**: jedna karta w formularzu = jeden wykaz = jedna tabela w gotowym dokumencie.

Do formatki jedzie lista słowników — dokładnie taka sama jak z pola typu `tabela` —
więc `{%p for %}` w Wordzie obsługuje jedno i drugie tak samo.
"""
from __future__ import annotations

import pathlib
import tempfile

from docx import Document

from app import db, generator, szablony

PODPOLA = [{"klucz": "adres_dotychczas", "etykieta": "Adres — dotychczasowy"},
           {"klucz": "adres_nowy", "etykieta": "Adres — nowy"}]


def _operat_z_sekcjami(klient):
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "wykazy", "etykieta": "Wykazy", "typ": "sekcje",
                        "etykieta_pozycji": "Wykaz", "podpola": PODPOLA}]})


def _wyslij(klient, **pola):
    dane = {"pole__nr_roboty": "GK.1", "notatka": ""}
    dane.update(pola)
    return klient.post("/generuj/spis_tresci_wzor", data=dane, follow_redirects=False)


# --- odczyt z formularza -----------------------------------------------------

def test_przenumerowanie_kart_obejmuje_pola_wielolinijkowe():
    """OFU/OZU/OZK są `<textarea>`, a wzorzec karty ma sztywny indeks 0.

    `przenumeruj()` łapiące tylko `input, select` zostawiało textarea klona pod
    indeksem pierwszej karty — pusta wartość drugiej działki wysyłała się jako
    `sek__wykazy_dzialek__0__ofu_dotychczas` i kasowała użytki działki 1 (przy
    odczycie ostatni wygrywa). Złapane na żywym operacie: OFU wyszło puste mimo
    wypełnienia. Zachowania JS nie dotkniemy bez przeglądarki, więc pilnujemy
    selektorów — obu, bo fokus po dodaniu karty ma trafić i w textarea.
    """
    from app.config import WEB

    szablon = (WEB / "templates" / "formularz.html").read_text(encoding="utf-8")

    # w pliku są dwie funkcje `przenumeruj`: pierwsza obsługuje wiersze tabel (tam są
    # wyłącznie inputy), sekcje przenumerowuje ta ostatnia
    przenumeruj = szablon.split("function przenumeruj")[-1].split("}")[0]
    assert "'input, select, textarea'" in przenumeruj, \
        "przenumerowanie pomija textarea — użytki drugiej karty nadpiszą pierwszą"


def test_sekcje_wracaja_jako_lista_slownikow(klient):
    """Ten sam kształt co przy tabeli — żeby formatka nie musiała ich rozróżniać."""
    _operat_z_sekcjami(klient)

    _wyslij(klient,
            **{"sek__wykazy__0__adres_dotychczas": "Polna 7",
               "sek__wykazy__0__adres_nowy": "Polna 7a",
               "sek__wykazy__1__adres_nowy": "Leśna 3"})

    import json
    zapisane = json.loads(db.dokumenty()[0]["dane_json"])["wykazy"]
    assert zapisane == [{"adres_dotychczas": "Polna 7", "adres_nowy": "Polna 7a"},
                        {"adres_nowy": "Leśna 3"}]


def test_pusta_sekcja_nie_jedzie_do_dokumentu(klient):
    """Formularz startuje z jedną pustą kartą; niewypełniona nie ma po co dawać
    pustej strony w wykazie."""
    _operat_z_sekcjami(klient)

    _wyslij(klient, **{"sek__wykazy__0__adres_dotychczas": "Polna 7",
                       "sek__wykazy__1__adres_dotychczas": "",
                       "sek__wykazy__1__adres_nowy": ""})

    import json
    assert len(json.loads(db.dokumenty()[0]["dane_json"])["wykazy"]) == 1


def test_kolejnosc_sekcji_idzie_po_indeksie_a_nie_po_kolejnosci_pol(klient):
    """Przeglądarka wysyła pola w kolejności DOM-u; po usunięciu karty numery
    przenumerowuje JS, ale kolejność ma wynikać z indeksu, nie z wysyłki."""
    _operat_z_sekcjami(klient)

    _wyslij(klient, **{"sek__wykazy__1__adres_nowy": "drugi",
                       "sek__wykazy__0__adres_nowy": "pierwszy"})

    import json
    assert [w["adres_nowy"] for w in json.loads(db.dokumenty()[0]["dane_json"])["wykazy"]] \
        == ["pierwszy", "drugi"]


# --- formularz ---------------------------------------------------------------

def test_formularz_pokazuje_zapisane_sekcje(klient):
    _operat_z_sekcjami(klient)
    _wyslij(klient, **{"sek__wykazy__0__adres_nowy": "Polna 7",
                       "sek__wykazy__1__adres_nowy": "Leśna 3"})

    formularz = klient.get(f"/nowy/spis_tresci_wzor?edytuj={db.dokumenty()[0]['id']}").text

    assert 'name="sek__wykazy__0__adres_nowy"' in formularz
    assert 'name="sek__wykazy__1__adres_nowy"' in formularz
    assert "Polna 7" in formularz and "Leśna 3" in formularz


def test_pusty_formularz_ma_jedna_karte_i_wzorzec(klient):
    """Bez ani jednej karty formularz wygląda, jakby czegoś brakowało; wzorzec
    w `<template>` sprawia, że JS nie musi znać listy podpól."""
    _operat_z_sekcjami(klient)

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert formularz.count('class="sekcja"') == 2, "jedna karta + wzorzec do klonowania"
    assert "wzorzec-sekcji" in formularz


# --- droga do dokumentu ------------------------------------------------------

def _wykaz_z_danymi(wykazy: list[dict]) -> Document:
    glowny = szablony.szablon_po_id("spis_tresci_wzor")
    sz = szablony.szablon_po_id("wykaz_zmian_budynku_wzor")
    kontekst = generator.przygotuj_kontekst(
        glowny, {"nr_roboty": "G.1", "wykazy_budynkow": wykazy}, {})
    return Document(generator.dopisz_dokument(sz, kontekst,
                                              pathlib.Path(tempfile.mkdtemp())))


def test_kazdy_wykaz_to_osobna_tabela(baza):
    """Sedno tej zmiany: kilka wykazów w jednym operacie, każdy ze swoimi danymi."""
    d = _wykaz_z_danymi([
        {"identyfikator_dzialki_dotychczas": "123/4", "uwagi_nowy": "przebudowa"},
        {"identyfikator_dzialki_dotychczas": "123/5", "uwagi_nowy": "wyburzony"},
    ])

    assert len(d.tables) == 2
    assert d.tables[0].rows[1].cells[3].text == "123/4"
    assert d.tables[1].rows[1].cells[3].text == "123/5"
    assert d.tables[0].rows[15].cells[4].text == "przebudowa"
    assert d.tables[1].rows[15].cells[4].text == "wyburzony"


def test_naglowek_powtarza_sie_przy_kazdym_wykazie(baza):
    """Każdy wykaz to osobna kartka do ośrodka — musi mieć swój numer roboty."""
    d = _wykaz_z_danymi([{"adres_nowy": "a"}, {"adres_nowy": "b"}])

    assert sum(1 for p in d.paragraphs if "G.1" in p.text) == 2
    assert sum(1 for p in d.paragraphs if "WYKAZ ZMIAN" in p.text) == 2


def test_lamanie_strony_miedzy_wykazami_a_nie_po_ostatnim(baza):
    """Łamanie po ostatnim wykazie zostawiałoby w operacie pustą kartkę."""
    import zipfile
    glowny = szablony.szablon_po_id("spis_tresci_wzor")
    sz = szablony.szablon_po_id("wykaz_zmian_budynku_wzor")

    def lamania(ile: int) -> int:
        kontekst = generator.przygotuj_kontekst(
            glowny, {"nr_roboty": "G.1",
                     "wykazy_budynkow": [{"adres_nowy": str(i)} for i in range(ile)]}, {})
        plik = generator.dopisz_dokument(sz, kontekst, pathlib.Path(tempfile.mkdtemp()))
        return zipfile.ZipFile(plik).read("word/document.xml").decode().count('w:type="page"')

    assert lamania(1) == 0
    assert lamania(2) == 1
    assert lamania(3) == 2


def test_wartosc_dziedziczy_krój_z_etykiety_obok(baza):
    """Wartość ma wyglądać jak opis w sąsiedniej kolumnie, a nie jak domyślna
    czcionka dokumentu — to ta sama pułapka co przy opisie przebiegu prac."""
    d = _wykaz_z_danymi([{"identyfikator_dzialki_dotychczas": "123/4"}])

    wartosc = d.tables[0].rows[1].cells[3].paragraphs[0].runs[0]
    etykieta = d.tables[0].rows[1].cells[1].paragraphs[0].runs[0]
    assert (wartosc.font.name, wartosc.font.size) == (etykieta.font.name, etykieta.font.size)


def test_formatka_nie_zostawia_niewypelnionych_znacznikow(baza):
    d = _wykaz_z_danymi([{"adres_nowy": "Polna 7"}])
    tekst = "\n".join(p.text for p in d.paragraphs)
    tekst += "\n".join(c.text for t in d.tables for w in t.rows for c in w.cells)

    assert "{{" not in tekst and "{%" not in tekst


# --- układ tabelaryczny ------------------------------------------------------

def test_nazwa_atrybutu_pojawia_sie_raz_a_nie_przy_kazdym_polu(klient):
    """Pierwsza wersja powtarzała nazwę przy obu stanach („Adres — dotychczasowy”,
    „Adres — nowy”) i przy piętnastu atrybutach robiła się ściana tekstu.

    Teraz sekcja wygląda jak tabela w dokumencie: nazwa raz, obok pole na każdy stan.
    """
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "wykazy", "etykieta": "Wykazy", "typ": "sekcje",
                        "kolumny": [{"klucz": "dotychczas", "etykieta": "Stan dotychczasowy"},
                                    {"klucz": "nowy", "etykieta": "Stan nowy"}],
                        "podpola": [
                            {"klucz": "adres_dotychczas", "wiersz": "Adres budynku",
                             "kolumna": "dotychczas"},
                            {"klucz": "adres_nowy", "wiersz": "Adres budynku",
                             "kolumna": "nowy"}]}]})

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert formularz.count("Adres budynku</th>") == 2, "raz w karcie, raz we wzorcu"
    assert "Stan dotychczasowy" in formularz and "Stan nowy" in formularz
    # nazwy pól się nie zmieniły, więc odczyt danych działa tak samo
    assert 'name="sek__wykazy__0__adres_dotychczas"' in formularz


def test_grupowanie_zachowuje_kolejnosc_atrybutow():
    """Kolejność ma być ta sama co w dokumencie — dlatego grupujemy w Pythonie,
    a nie `groupby` w Jinji, które najpierw sortuje."""
    pole = szablony.Pole(
        klucz="w", etykieta="W", typ="sekcje",
        podpola=[{"klucz": "b_nowy", "wiersz": "Beta", "kolumna": "nowy"},
                 {"klucz": "a_nowy", "wiersz": "Alfa", "kolumna": "nowy"},
                 {"klucz": "b_stary", "wiersz": "Beta", "kolumna": "dotychczas"}])

    wiersze = pole.wiersze_sekcji

    assert [w["etykieta"] for w in wiersze] == ["Beta", "Alfa"]
    assert sorted(wiersze[0]["pola"]) == ["dotychczas", "nowy"]


# --- pusty wykaz nie powstaje ------------------------------------------------

def _operat_z_wykazem(klient):
    """Operat główny plus dokument dodatkowy, który bez danych nie ma sensu."""
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "wykazy", "etykieta": "Wykazy", "typ": "sekcje",
                        "podpola": PODPOLA},
                       {"klucz": "dokumenty_wykazy", "etykieta": "Wygeneruj",
                        "typ": "dokumenty", "tylko": ["wykaz_wzor"]}]})
    klient.srodowisko.dodaj_szablon(
        "wykaz_wzor", ["{%p for w in wykazy %}", "{{ w.adres_nowy }}", "{%p endfor %}"],
        opis={"nazwa": "Wykaz zmian", "wymaga": "wykazy"})


def _pliki_operatu(klient) -> list[str]:
    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    return sorted(p.name for p in katalog.iterdir() if p.suffix == ".docx")


def test_wykaz_bez_danych_w_ogole_nie_powstaje(klient):
    """Dokument będący samą pętlą po pustej liście wychodził jako plik bez jednej
    litery — w składaniu operatu pusty kafelek, po którym nie wiadomo, czy to usterka."""
    _operat_z_wykazem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__dokumenty_wykazy": "wykaz_wzor",
                      "sek__wykazy__0__adres_nowy": ""},
                follow_redirects=False)

    assert _pliki_operatu(klient) == ["spis_tresci.docx"]


def test_program_mowi_dlaczego_wykazu_nie_ma(klient):
    """Ciche pominięcie jest gorsze niż pusty plik: brat zaznaczył dokument i ma prawo
    wiedzieć, czemu go nie widzi."""
    _operat_z_wykazem(klient)

    odpowiedz = klient.post("/generuj/spis_tresci_wzor",
                            data={"pole__nr_roboty": "GK.1", "notatka": "",
                                  "pole__dokumenty_wykazy": "wykaz_wzor",
                                  "sek__wykazy__0__adres_nowy": ""},
                            follow_redirects=False)

    from urllib.parse import unquote
    komunikat = unquote(odpowiedz.headers["location"])
    assert "nie powstał, bo nie wypełniłeś ani jednej pozycji" in komunikat
    assert "Wykaz zmian" in komunikat, "brat ma wiedzieć, którego dokumentu to dotyczy"


def test_wykaz_z_danymi_powstaje_normalnie(klient):
    _operat_z_wykazem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__dokumenty_wykazy": "wykaz_wzor",
                      "sek__wykazy__0__adres_nowy": "Polna 7"},
                follow_redirects=False)

    assert _pliki_operatu(klient) == ["spis_tresci.docx", "wykaz.docx"]


# --- poprawianie operatu sprząta dokumenty odznaczone w tej rundzie ----------
#
# Poprawianie nadpisuje pliki nazwa po nazwie, więc dokument odznaczony (albo pominięty
# przez `wymaga`) zostawał w katalogu z danymi z poprzedniej rundy. Szedł potem do
# scalonego PDF-a, choć spis treści o nim milczał — a przy pustym wykazie program
# ogłaszał „nie powstał”, mając jego starą wersję na dysku.

def _wygeneruj_z_wykazem(klient) -> int:
    _operat_z_wykazem(klient)
    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__dokumenty_wykazy": "wykaz_wzor",
                      "sek__wykazy__0__adres_nowy": "Polna 7"},
                follow_redirects=False)
    assert _pliki_operatu(klient) == ["spis_tresci.docx", "wykaz.docx"]
    return db.dokumenty()[0]["id"]


def test_poprawianie_sprzata_odznaczony_dokument(klient):
    """Odznaczony przy poprawianiu dokument znika z katalogu razem z podglądem."""
    identyfikator = _wygeneruj_z_wykazem(klient)
    katalog = db.dokumenty()[0]["katalog"]
    podglad = klient.srodowisko.dane / "podglad" / katalog / "wykaz.pdf"
    podglad.parent.mkdir(parents=True, exist_ok=True)
    podglad.write_bytes(b"stary podglad")

    klient.post(f"/generuj/spis_tresci_wzor?edytuj={identyfikator}",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "sek__wykazy__0__adres_nowy": "Polna 7"},
                follow_redirects=False)

    assert _pliki_operatu(klient) == ["spis_tresci.docx"], \
        "stary wykaz został w katalogu i poszedłby do scalonego PDF-a"
    assert not podglad.exists(), "podgląd skasowanego dokumentu został w dane/podglad"


def test_poprawianie_z_oproznionym_wykazem_sprzata_stary_plik(klient):
    """Komunikat „nie powstał” nie może kłamać: stary plik z poprzednimi danymi
    ma zniknąć, skoro tym razem nie ma go z czego zrobić."""
    identyfikator = _wygeneruj_z_wykazem(klient)

    klient.post(f"/generuj/spis_tresci_wzor?edytuj={identyfikator}",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__dokumenty_wykazy": "wykaz_wzor",
                      "sek__wykazy__0__adres_nowy": ""},
                follow_redirects=False)

    assert _pliki_operatu(klient) == ["spis_tresci.docx"]


def test_poprawianie_nie_rusza_plikow_dolozonych_recznie(klient):
    """Sprzątanie obejmuje wyłącznie nazwy nadawane przez program — mapa dołożona
    Eksploratorem, nawet jako .docx, zostaje."""
    identyfikator = _wygeneruj_z_wykazem(klient)
    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    (katalog / "mapa_do_celow_projektowych.docx").write_bytes(b"od brata")

    klient.post(f"/generuj/spis_tresci_wzor?edytuj={identyfikator}",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "sek__wykazy__0__adres_nowy": "Polna 7"},
                follow_redirects=False)

    pliki = _pliki_operatu(klient)
    assert "mapa_do_celow_projektowych.docx" in pliki
    assert "wykaz.docx" not in pliki


def test_dokument_bez_deklaracji_wymaga_powstaje_zawsze(klient):
    """`wymaga` jest wpisem w `.json`, a nie regułą zgadującą po treści — dokument,
    który jej nie ma, zachowuje się dokładnie jak dotąd."""
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "dokumenty_inne", "etykieta": "Wygeneruj",
                        "typ": "dokumenty", "tylko": ["notatka_wzor"]}]})
    klient.srodowisko.dodaj_szablon("notatka_wzor", ["{{ nr_roboty }}"],
                                    opis={"nazwa": "Notatka"})

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__dokumenty_inne": "notatka_wzor"}, follow_redirects=False)

    assert _pliki_operatu(klient) == ["notatka.docx", "spis_tresci.docx"]


# --- spis treści jest jedynym włącznikiem dokumentów -------------------------
#
# Wcześniej pozycja w spisie treści i checkbox „Wygeneruj” w karcie dokumentu pytały
# o to samo dwa razy — dało się je ustawić sprzecznie: dokument w spisie, a pliku brak.

def _operat_ze_spisem(klient):
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "spis_tresci", "etykieta": "Co wchodzi", "grupa": "Spis treści",
                        "typ": "wybor_wielokrotny",
                        "opcje": ["Spis treści", "Sprawozdanie techniczne", "Mapa"],
                        "zawsze": ["Spis treści"],
                        "dokumenty": {"Sprawozdanie techniczne": "sprawozdanie_wzor"}},
                       {"klucz": "uwagi", "etykieta": "Uwagi", "typ": "textarea",
                        "aktywne_gdy": "spis_tresci:Sprawozdanie techniczne"}]})
    klient.srodowisko.dodaj_szablon("sprawozdanie_wzor", ["{{ nr_roboty }} {{ uwagi }}"],
                                    opis={"nazwa": "Sprawozdanie techniczne"})


def test_zaznaczenie_w_spisie_tresci_generuje_dokument(klient):
    _operat_ze_spisem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__spis_tresci": "Sprawozdanie techniczne",
                      "pole__uwagi": "treść"},
                follow_redirects=False)

    assert _pliki_operatu(klient) == ["spis_tresci.docx", "sprawozdanie.docx"]


def test_brak_pozycji_w_spisie_tresci_to_brak_dokumentu(klient):
    """Sedno zmiany: nie ma osobnego „wygeneruj”, więc nie da się mieć dokumentu
    w spisie treści bez pliku ani pliku bez pozycji w spisie."""
    _operat_ze_spisem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__spis_tresci": "Mapa", "pole__uwagi": ""},
                follow_redirects=False)

    assert _pliki_operatu(klient) == ["spis_tresci.docx"]


def test_pola_dokumentu_wisza_na_pozycji_ze_spisu_tresci(klient):
    """Zaznaczenie w spisie treści ma **włączać pola** tego dokumentu."""
    _operat_ze_spisem(klient)

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert 'data-aktywne-gdy="spis_tresci:Sprawozdanie techniczne"' in formularz


def test_dokument_ze_spisu_nie_wraca_do_karty_inne_dokumenty(klient):
    """Szablon, który ma już włącznik w spisie treści, nie może dać się włączyć
    drugi raz gdzie indziej — inaczej dwa miejsca mówiłyby co innego."""
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "spis_tresci", "etykieta": "Co wchodzi",
                        "typ": "wybor_wielokrotny", "opcje": ["Sprawozdanie techniczne"],
                        "dokumenty": {"Sprawozdanie techniczne": "sprawozdanie_wzor"}},
                       {"klucz": "dokumenty", "etykieta": "Wygeneruj też",
                        "typ": "dokumenty"}]})
    klient.srodowisko.dodaj_szablon("sprawozdanie_wzor", ["{{ nr_roboty }}"],
                                    opis={"nazwa": "Sprawozdanie techniczne"})
    klient.srodowisko.dodaj_szablon("inny_wzor", ["{{ nr_roboty }}"],
                                    opis={"nazwa": "Inny dokument"})

    formularz = klient.get("/nowy/spis_tresci_wzor").text
    lista_innych = formularz.split('name="pole__dokumenty"')

    assert "Inny dokument" in formularz, "szablon bez pozycji w spisie ma zostać do wyboru"
    assert formularz.count('value="sprawozdanie_wzor"') == 0, \
        "sprawozdanie ma włącznik w spisie treści — w „Innych dokumentach” nie ma czego dublować"


# --- wykaz działki: każda działka to OSOBNA STRONA ---------------------------
#
# Decyzja z 19.08.2026 (druga tego dnia — pionowa tabela z dwoma wierszami na działkę
# nie utrzymała się w praniu): wykaz działki wygląda dokładnie jak wykaz budynku.
# Atrybut w wierszu, dwie kolumny stanów, jedna działka na stronę, numer działki
# w nagłówku **swojej** strony. Formatkę buduje z formatki budynku
# `narzedzia/utworz_wykaz_dzialki.py`, żeby oba dokumenty nie rozjechały się z czasem.

def _wykaz_dzialek(dzialki: list[dict]) -> Document:
    glowny = szablony.szablon_po_id("spis_tresci_wzor")
    sz = szablony.szablon_po_id("wykaz_zmian_dzialki_wzor")
    kontekst = generator.przygotuj_kontekst(
        glowny, {"nr_roboty": "G.1", "polozenie": {"obreb": "247301_1.0112"},
                 "wykazy_dzialek": dzialki}, {})
    kontekst["polozenie_obreb_teryt"] = "247301_1.0112"
    return Document(generator.dopisz_dokument(sz, kontekst,
                                              pathlib.Path(tempfile.mkdtemp())))


def test_kazda_dzialka_dostaje_osobna_strone(baza):
    """Tak samo jak przy budynku: jedna pozycja, jedna strona, własna tabela."""
    d = _wykaz_dzialek([{"dzialka": "119/10", "numer_dotychczas": "119/10"},
                        {"dzialka": "119/11", "numer_dotychczas": "119/11"}])

    assert len(d.tables) == 2, "każda działka ma własną tabelę"
    assert [t.rows[2].cells[2].text.strip() for t in d.tables] == ["119/10", "119/11"]
    lamania = sum(1 for p in d.paragraphs
                  for b in p.runs for _ in b._r.findall(
                      "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"))
    assert lamania == 1, "łamanie strony między działkami, ale nie po ostatniej"


def test_numer_dzialki_wchodzi_do_naglowka_swojej_strony(baza):
    """Numer jest podpolem każdego wykazu, nie jednym polem na cały dokument — dlatego
    na drugiej stronie stoi drugi numer. Identyfikatorem jest dopiero całość: obręb,
    kropka i ten numer."""
    d = _wykaz_dzialek([{"dzialka": "119/10"}, {"dzialka": "119/11"}])

    naglowki = [p for p in d.paragraphs if "Identyfikator działki" in p.text]
    assert [p.text.split("\t")[-1] for p in naglowki] == ["[247301_1.0112.119/10]",
                                                          "[247301_1.0112.119/11]"]
    # wartość ma być pogrubiona, tak jak w trzech wierszach nad nią — pierwsza wersja
    # skryptu sklejała etykietę i wartość w jeden bieg i to pogrubienie ginęło
    wartosc = [b for b in naglowki[0].runs if "247301" in b.text]
    assert wartosc and all(b.bold for b in wartosc), "numer działki stracił pogrubienie"
    assert not any(b.bold for b in naglowki[0].runs if "Identyfikator" in b.text), \
        "etykieta ma zostać zwykła"


def test_atrybuty_dzialki_stoja_w_wierszach(baza):
    """Nagłówek ma dwa piętra: „STAN DOTYCHCZASOWY” i „STAN NOWY”, a pod każdym cztery
    podkolumny użytków. Wiersze bez podziału (numer, pole powierzchni) mają wartość
    scaloną przez wszystkie cztery."""
    d = _wykaz_dzialek([{
        "dzialka": "119/10", "numer_dotychczas": "119/10", "numer_nowy": "119/11",
        "pow_ewidencyjna_dotychczas": "5.5241", "pow_ewidencyjna_nowy": "5.5241",
        "ofu_dotychczas": "R", "ofu_nowy": "Ba",
        "ozu_dotychczas": "IVa", "ozk_dotychczas": "II",
        "pow_uzytkow_dotychczas": "5.5241", "pow_uzytkow_nowy": "5.5241"}]) 

    tabela = d.tables[0]
    assert len(tabela.rows) == 2 + 3, "dwa wiersze nagłówka, trzy atrybuty"
    gora = [k.text.strip() for k in tabela.rows[0].cells]
    assert gora[0] == "L.p." and gora[1] == "Oznaczenie atrybutu działki"
    assert gora[2:6] == ["STAN DOTYCHCZASOWY"] * 4, "stan scalony nad podkolumnami"
    assert gora[6:] == ["STAN NOWY"] * 4
    assert [k.text.strip() for k in tabela.rows[1].cells][2:] == \
        ["OFU", "OZU", "OZK", "PPU"] * 2

    opisy = [w.cells[1].text.strip() for w in tabela.rows[2:]]
    assert opisy[0] == "Numer działki"
    assert "Użytki gruntowe" in opisy[2]
    # wiersz bez podziału: jedna wartość na całą szerokość stanu
    numer = tabela.rows[2]
    assert [k.text.strip() for k in numer.cells[2:6]] == ["119/10"] * 4
    assert [k.text.strip() for k in numer.cells[6:]] == ["119/11"] * 4
    # wiersz użytków: cztery wartości obok siebie w każdym stanie
    uzytki = tabela.rows[4]
    assert [k.text.strip() for k in uzytki.cells[2:6]] == ["R", "IVa", "II", "5.5241"]
    assert [k.text.strip() for k in uzytki.cells[6:]] == ["Ba", "", "", "5.5241"]


def test_kilka_uzytkow_stoi_jeden_pod_drugim(baza):
    """Jedna działka bywa podzielona na kilka użytków. Każdy to jedna linijka we
    **wszystkich czterech** kolumnach naraz — dlatego OFU, OZU, OZK i PPU stoją obok
    siebie, a nie w osobnych wierszach tabeli."""
    d = _wykaz_dzialek([{"dzialka": "119/10", "ofu_dotychczas": "R\nR\nW/R",
                         "ozk_dotychczas": "II\nIIIa\nIIIa",
                         "pow_uzytkow_dotychczas": "0.4013\n5.1014\n0.0138"}])

    uzytki = next(w for w in d.tables[0].rows if "Użytki gruntowe" in w.cells[1].text)
    assert uzytki.cells[2].text.strip().splitlines() == ["R", "R", "W/R"]
    assert uzytki.cells[4].text.strip().splitlines() == ["II", "IIIa", "IIIa"]
    assert uzytki.cells[5].text.strip().splitlines() == ["0.4013", "5.1014", "0.0138"]


def _vAlign(komorka) -> str:
    from docx.oxml.ns import qn
    element = komorka._tc.get_or_add_tcPr().find(qn("w:vAlign"))
    return element.get(qn("w:val")) if element is not None else "top"


def test_pojedyncza_wartosc_stoi_na_srodku_komorki(baza):
    """Wiersz z jedną wartością po każdej stronie ma ją równo z nazwą atrybutu obok.

    Wyrównanie do góry jest potrzebne tylko przy kilku linijkach; w pozostałych wierszach
    kleiło liczbę do górnej krawędzi, podczas gdy dwuwierszowa etykieta obok stała
    na środku — i wyglądało to na usterkę.
    """
    d = _wykaz_dzialek([{"dzialka": "119/80", "pow_ewidencyjna_dotychczas": "5.5241",
                         "pow_ewidencyjna_nowy": "5.6000"}])

    wiersz = next(w for w in d.tables[0].rows
                  if "Pole powierzchni ewidencyjnej" in w.cells[1].text)
    assert {_vAlign(k) for k in wiersz.cells[2:]} == {"center"}


def test_kilka_linijek_ciagnie_komorki_stanow_do_gory(baza):
    """Przy kilku użytkach wartość w sąsiedniej kolumnie musi wypaść na wysokości
    **swojej** linijki — brat dosuwa ją pustymi linijkami, a to działa tylko wtedy,
    gdy komórki są wyrównane do góry. Dotyczy **wszystkich** podkolumn stanu, nie tylko
    tej wielolinijkowej: inaczej „Ba” w OFU stanu nowego pływałoby w połowie wysokości.
    Nazwa atrybutu zostaje na środku.
    """
    d = _wykaz_dzialek([{"dzialka": "119/80", "ozk_dotychczas": "II\nIIIa\nIIIa",
                         "ofu_nowy": "Ba"}])

    wiersz = next(w for w in d.tables[0].rows if "Użytki gruntowe" in w.cells[1].text)
    assert {_vAlign(k) for k in wiersz.cells[2:]} == {"top"}
    assert _vAlign(wiersz.cells[1]) == "center", "nazwa atrybutu ma zostać na środku"


def test_wykaz_dzialek_bez_danych_nie_powstaje():
    """Ta sama zasada co przy budynku — deklaracja `wymaga` w `.json` szablonu."""
    assert szablony.szablon_po_id("wykaz_zmian_dzialki_wzor").wymaga == "wykazy_dzialek"


def test_wykaz_dzialki_wyglada_jak_wykaz_budynku():
    """Oba wykazy mają być jednym kompletem: ta sama kartka, ta sama tabela, ten sam
    nagłówek. Rozjazd zaczyna się od drobiazgu, więc porównujemy wprost."""
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    budynek = Document(str(szablony.szablon_po_id("wykaz_zmian_budynku_wzor").plik))
    dzialka = Document(str(szablony.szablon_po_id("wykaz_zmian_dzialki_wzor").plik))

    assert dzialka.sections[0].orientation == WD_ORIENT.PORTRAIT
    assert dzialka.sections[0].page_width == budynek.sections[0].page_width
    # Siatka kolumn jest inna (stany dzielą się na cztery podkolumny użytków), ale
    # dwie pierwsze kolumny — L.p. i opis — mają zostać tej samej szerokości, bo to one
    # decydują, czy oba dokumenty wyglądają jak komplet.
    kolumny = lambda d: [int(k.get(qn("w:w")))                       # noqa: E731
                        for k in d.tables[0]._tbl.find(qn("w:tblGrid"))]
    assert kolumny(dzialka)[0] == kolumny(budynek)[0]
    assert sum(kolumny(dzialka)) == sum(kolumny(budynek)), "inna szerokość całej tabeli"


def test_zmieniony_stan_nowy_jest_czerwony_i_pogrubiony(baza):
    """Po tym ośrodek czyta, co się zmieniło — brat zaznaczał to dotąd ręcznie.

    Wartość taka sama jak dotychczasowa zostaje czarna; różna (także wpisana tam,
    gdzie wcześniej było pusto) idzie na czerwono i grubo.
    """
    d = _wykaz_dzialek([{"dzialka": "119/10",
                         "numer_dotychczas": "119/10", "numer_nowy": "119/10",
                         "pow_ewidencyjna_dotychczas": "5.5241",
                         "pow_ewidencyjna_nowy": "5.6000",
                         "ofu_nowy": "Ba"}])

    def biegi(wiersz):
        return [b for p in wiersz.cells[-1].paragraphs for b in p.runs if b.text.strip()]

    tabela = d.tables[0]
    bez_zmiany, ze_zmiana = tabela.rows[2], tabela.rows[3]
    assert [b.text for b in biegi(bez_zmiany)] == ["119/10"]
    assert not any(b.bold for b in biegi(bez_zmiany)), "niezmieniony stan ma zostać zwykły"
    assert all(b.bold for b in biegi(ze_zmiana))
    assert all(str(b.font.color.rgb) == "FF0000" for b in biegi(ze_zmiana))
    # wartość wpisana tam, gdzie dotąd było pusto, to też zmiana
    uzytki = next(w for w in tabela.rows if "Użytki gruntowe" in w.cells[1].text)
    ofu_nowy = [b for p in uzytki.cells[6].paragraphs for b in p.runs if b.text.strip()]
    assert all(b.bold and str(b.font.color.rgb) == "FF0000" for b in ofu_nowy)


def test_wykaz_budynku_tez_czerwieni_zmiany(baza):
    """Ta sama zasada w obu wykazach — inaczej brat musiałby pamiętać, gdzie działa."""
    sz = szablony.szablon_po_id("wykaz_zmian_budynku_wzor")
    d = Document(generator.dopisz_dokument(sz, {
        "nr_roboty": "G.1",
        "wykazy_budynkow": [{"adres_dotychczas": "Polna 7", "adres_nowy": "Polna 7",
                             "kondygnacje_nadziemne_dotychczas": "1",
                             "kondygnacje_nadziemne_nowy": "2"}]},
        pathlib.Path(tempfile.mkdtemp())))

    def biegi(tekst_wiersza):
        wiersz = next(w for w in d.tables[0].rows if tekst_wiersza in w.cells[1].text)
        return [b for p in wiersz.cells[-1].paragraphs for b in p.runs if b.text.strip()]

    assert not any(b.bold for b in biegi("Adres budynku"))
    assert all(b.bold and str(b.font.color.rgb) == "FF0000"
               for b in biegi("Liczba kondygnacji"))


def test_w_wykazie_budynku_nie_ma_juz_odsylaczy_do_przypisow():
    """Odsyłacze (1) i (2) prowadziły do przypisów, których w formatce nie ma —
    zostały po wzorze z rozporządzenia. Indeks górny przy m² to co innego i zostaje."""
    d = Document(str(szablony.szablon_po_id("wykaz_zmian_budynku_wzor").plik))
    tresc = " ".join(k.text for w in d.tables[0].rows for k in w.cells)

    assert "(1)" not in tresc and "(2)" not in tresc
    assert "Pole zabudowy m2" in tresc.replace("\n", " "), "jednostka m² ma zostać"


def test_liczby_porzadkowe_wykazu_budynku_stoja_na_srodku_komorki():
    """Cyfry L.p. siadały na różnych wysokościach — najbardziej w wierszach niższych.

    Winne były dwie rzeczy naraz: część akapitów miała interlinię z własnego
    ustawienia, a część brała ją ze stylu, i **znacznik końca akapitu był większy
    niż cyfra** (12 pt przy 7-punktowej cyfrze). Wyśrodkowanie w pionie centruje
    wiersz tekstu, a nie same cyfry, więc rozdmuchany znacznik podnosił je nad
    środek komórki. Zmierzone na wydruku: odchyłka do 3 pt, po poprawce ±0,7 pt.
    """
    from docx.oxml.ns import qn

    d = Document(str(szablony.szablon_po_id("wykaz_zmian_budynku_wzor").plik))
    rozmiar = lambda el: (el.find(qn("w:sz")).get(qn("w:val"))     # noqa: E731
                          if el is not None and el.find(qn("w:sz")) is not None else None)

    interlinie, znaczniki = set(), set()
    for wiersz in d.tables[0]._tbl.findall(qn("w:tr")):
        komorka = wiersz.findall(qn("w:tc"))[0]
        tcPr = komorka.find(qn("w:tcPr"))
        vAlign = tcPr.find(qn("w:vAlign")) if tcPr is not None else None
        assert vAlign is not None and vAlign.get(qn("w:val")) == "center", \
            "komórka L.p. bez wyśrodkowania w pionie"

        for akapit in komorka.findall(qn("w:p")):
            pPr = akapit.find(qn("w:pPr"))
            jc = pPr.find(qn("w:jc")) if pPr is not None else None
            assert jc is not None and jc.get(qn("w:val")) == "center", \
                "cyfra L.p. nie jest wyśrodkowana w poziomie"
            odstep = pPr.find(qn("w:spacing"))
            interlinie.add((odstep.get(qn("w:line")), odstep.get(qn("w:lineRule")))
                           if odstep is not None else None)
            znacznik = rozmiar(pPr.find(qn("w:rPr")))
            znaczniki.add(znacznik)
            bieg = akapit.find(qn("w:r"))
            if bieg is not None:
                assert znacznik == rozmiar(bieg.find(qn("w:rPr"))), \
                    "znacznik akapitu większy niż cyfra — podniesie ją nad środek"

    assert len(interlinie) == 1, f"różne interlinie w kolumnie L.p.: {interlinie}"
    assert len(znaczniki) == 1, f"różne rozmiary znaczników akapitu: {znaczniki}"


# --- listy wyboru w sekcji (KŚT) ---------------------------------------------

# Klasyfikacja Środków Trwałych, podgrupy 10 i 11 — słowo w słowo z rozporządzenia
# (KŚT 2016, Dz.U. 2016 poz. 1864; w KŚT 2010 te same nazwy). Brat wpisywał to z ręki,
# a nazwa rodzaju musi zgadzać się co do znaku, bo idzie do dokumentu składanego
# w ośrodku. Podgrupy 12 (rodzaje 121 i 122) **nie ma na liście**: to lokale, a nie
# budynki — geodety w wykazie zmian danych budynku nie dotyczą (decyzja brata).
KST = [
    "101 - BUDYNKI PRZEMYSŁOWE",
    "102 - BUDYNKI TRANSPORTU I ŁĄCZNOŚCI",
    "103 - BUDYNKI HANDLOWO-USŁUGOWE",
    "104 - ZBIORNIKI, SILOSY I BUDYNKI MAGAZYNOWE",
    "105 - BUDYNKI BIUROWE",
    "106 - BUDYNKI SZPITALI I INNE BUDYNKI OPIEKI ZDROWOTNEJ",
    "107 - BUDYNKI OŚWIATY, NAUKI I KULTURY ORAZ BUDYNKI SPORTOWE",
    "108 - BUDYNKI PRODUKCYJNE, USŁUGOWE I GOSPODARCZE DLA ROLNICTWA",
    "109 - POZOSTAŁE BUDYNKI NIEMIESZKALNE",
    "110 - BUDYNKI MIESZKALNE",
]


def _podpola_kst() -> list[dict]:
    pole = next(p for p in szablony.szablon_po_id("spis_tresci_wzor").pola
                if p.klucz == "wykazy_budynkow")
    return [pod for pod in pole.podpola if pod["klucz"].startswith("rodzaj_kst")]


def test_rodzaj_kst_ma_liste_z_rozporzadzenia():
    """Nazwy rodzajów muszą zgadzać się co do znaku — dokument idzie do ośrodka."""
    podpola = _podpola_kst()

    assert len(podpola) == 2, "KŚT ma być w obu stanach"
    for pod in podpola:
        assert pod["opcje"][0] == "", "pierwsza pozycja pusta — wiersz wolno zostawić bez wpisu"
        assert pod["opcje"][1:] == KST
        assert not [o for o in pod["opcje"] if o.startswith(("121", "122"))], \
            "lokale to nie budynki — podgrupa 12 nie ma czego szukać w tym wykazie"


def test_oba_stany_kst_biora_te_sama_liste():
    """Dwie kopie listy w `.json` rozjechałyby się przy pierwszej poprawce, a wykaz
    z dwiema wersjami tej samej klasyfikacji to dokument nie do przyjęcia."""
    dotychczas, nowy = _podpola_kst()

    assert dotychczas["opcje"] == nowy["opcje"]


def test_lista_wyboru_w_sekcji_to_select_a_nie_pole_tekstowe(klient):
    """Nazwa rodzaju ma 60 znaków i trzy przecinki — z ręki nikt tego nie wpisze
    bez literówki."""
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "listy": {"kst": ["", "101 - BUDYNKI PRZEMYSŁOWE", "110 - BUDYNKI MIESZKALNE"]},
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "wykazy", "etykieta": "Wykazy", "typ": "sekcje",
                        "kolumny": [{"klucz": "dotychczas", "etykieta": "Stan dotychczasowy"},
                                    {"klucz": "nowy", "etykieta": "Stan nowy"}],
                        "podpola": [
                            {"klucz": "kst_dotychczas", "wiersz": "Rodzaj budynku według KŚT",
                             "kolumna": "dotychczas", "opcje": "kst"},
                            {"klucz": "kst_nowy", "wiersz": "Rodzaj budynku według KŚT",
                             "kolumna": "nowy", "opcje": "kst"},
                            {"klucz": "uwagi_nowy", "wiersz": "Uwagi", "kolumna": "nowy"}]}]})

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert formularz.count('<option value="101 - BUDYNKI PRZEMYSŁOWE"') == 4, \
        "dwie kolumny w karcie i dwie we wzorcu do klonowania"
    assert 'name="sek__wykazy__0__uwagi_nowy"' in formularz
    assert '<input name="sek__wykazy__0__uwagi_nowy"' in formularz, \
        "podpole bez listy ma zostać zwykłym polem tekstowym"


def test_wartosc_spoza_listy_nie_ginie_przy_poprawianiu(klient):
    """Operaty sprzed tej zmiany mają w tym miejscu tekst wpisany z ręki. Samo otwarcie
    „Popraw ten operat” podstawiłoby pierwszą pozycję listy i skasowało jego wpis."""
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "listy": {"kst": ["", "101 - BUDYNKI PRZEMYSŁOWE"]},
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "wykazy", "etykieta": "Wykazy", "typ": "sekcje",
                        "podpola": [{"klucz": "kst_nowy", "etykieta": "KŚT",
                                     "opcje": "kst"}]}]})
    _wyslij(klient, **{"sek__wykazy__0__kst_nowy": "budynek gospodarczy"})
    wpis = db.dokumenty()[0]

    formularz = klient.get(f"/nowy/spis_tresci_wzor?edytuj={wpis['id']}").text

    assert '<option value="budynek gospodarczy" selected>' in formularz


# --- strona operatu ----------------------------------------------------------

def _operat_z_kolumnami(klient):
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "wykazy", "etykieta": "Wykazy", "typ": "sekcje",
                        "etykieta_pozycji": "Wykaz",
                        "kolumny": [{"klucz": "dotychczas", "etykieta": "Stan dotychczasowy"},
                                    {"klucz": "nowy", "etykieta": "Stan nowy"}],
                        "podpola": [
                            {"klucz": "adres_dotychczas", "wiersz": "Adres budynku",
                             "kolumna": "dotychczas"},
                            {"klucz": "adres_nowy", "wiersz": "Adres budynku",
                             "kolumna": "nowy"},
                            {"klucz": "pole_dotychczas", "wiersz": "Pole zabudowy",
                             "kolumna": "dotychczas"},
                            {"klucz": "pole_nowy", "wiersz": "Pole zabudowy",
                             "kolumna": "nowy"}]}]})


def test_strona_operatu_pokazuje_wpisane_wiersze(klient):
    """„2 wierszy” nie mówiło nic — a po to właśnie wchodzi się w gotowy operat.

    Wykazów w operacie bywa kilka, więc każdy dostaje własny nagłówek z numerem;
    inaczej nie wiadomo, który adres należy do którego budynku.
    """
    _operat_z_kolumnami(klient)

    _wyslij(klient, **{"sek__wykazy__0__adres_nowy": "Polna 7a",
                       "sek__wykazy__1__pole_dotychczas": "148",
                       "sek__wykazy__1__pole_nowy": "162"})
    strona = klient.get(f"/dokument/{db.dokumenty()[0]['id']}").text

    assert "wierszy" not in strona, "została sama liczba zamiast danych"
    assert "Polna 7a" in strona and "148" in strona and "162" in strona
    assert "Wykaz 1" in strona and "Wykaz 2" in strona
    assert strona.index("Polna 7a") < strona.index("Wykaz 2"), \
        "dane wpadły do niewłaściwego wykazu"


def test_strona_operatu_pomija_niewypelnione_atrybuty(klient):
    """Wykaz budynku ma piętnaście atrybutów, a wypełnione bywają dwa.

    Wypisywanie wszystkich dałoby ścianę kresek zasłaniającą to, co istotne.
    """
    _operat_z_kolumnami(klient)

    _wyslij(klient, **{"sek__wykazy__0__adres_nowy": "Polna 7a"})
    strona = klient.get(f"/dokument/{db.dokumenty()[0]['id']}").text

    assert "Adres budynku" in strona
    assert "Pole zabudowy" not in strona, "pusty atrybut trafił na stronę operatu"


def test_sekcja_bez_kolumn_tez_sie_pokazuje(klient):
    """Nie każda sekcja ma układ tabelaryczny — płaską wypisujemy etykietami podpól."""
    _operat_z_sekcjami(klient)          # `podpola` bez `wiersz`/`kolumna`

    _wyslij(klient, **{"sek__wykazy__0__adres_nowy": "Leśna 3"})
    strona = klient.get(f"/dokument/{db.dokumenty()[0]['id']}").text

    assert "Adres — nowy" in strona and "Leśna 3" in strona
    assert "Adres — dotychczasowy" not in strona


def test_plaska_sekcja_nie_wypisuje_wartosci_dwa_razy(klient):
    """Podpola bez `wiersz` idą do bloku wspólnego nad tabelą — wypisane drugi raz
    w tabeli robiły z podglądu echo: każda wartość stała na stronie operatu dwukrotnie."""
    _operat_z_sekcjami(klient)

    _wyslij(klient, **{"sek__wykazy__0__adres_nowy": "Leśna 3"})
    strona = klient.get(f"/dokument/{db.dokumenty()[0]['id']}").text

    assert strona.count("Leśna 3") == 1, "wartość płaskiej sekcji wypisana podwójnie"


def test_tabele_wykazow_mieszcza_sie_w_marginesach():
    """Tabela szersza od szerokości tekstu **choćby o 23 twipy** ucieka na drugą stronę
    (pułapka 12g). Porównanie działki z budynkiem tego nie łapie — obie formatki mogą
    być za szerokie o tyle samo — więc strażnik musi być bezwzględny."""
    from docx.oxml.ns import qn

    for nazwa in ("wykaz_zmian_budynku_wzor", "wykaz_zmian_dzialki_wzor"):
        d = Document(str(szablony.szablon_po_id(nazwa).plik))
        sekcja = d.sections[0]
        szerokosc_tekstu = int((sekcja.page_width - sekcja.left_margin
                                - sekcja.right_margin) / 914400 * 1440)
        for tabela in d.tables:
            kolumny = [int(k.get(qn("w:w"))) for k in tabela._tbl.find(qn("w:tblGrid"))]
            assert sum(kolumny) <= szerokosc_tekstu, \
                f"{nazwa}: tabela ma {sum(kolumny)} twipów przy {szerokosc_tekstu} tekstu"


# --- puste linijki w polach wielolinijkowych ---------------------------------
#
# OFU/OZU/OZK stoją w dwóch kolumnach obok siebie i brat wyrównuje wartość pustymi
# enterami do właściwej linijki sąsiedniej kolumny — wpis w stanie nowym potrafi
# dotyczyć dopiero drugiego użytku. Zjadały je dwa miejsca: `strip()` przy odczycie
# formularza i parser HTML, który ignoruje pierwszy znak nowej linii po <textarea>.

def test_puste_linijki_na_poczatku_wartosci_przezywaja_odczyt(klient):
    """Przeglądarka wysyła `\r\n`; do historii ma wejść `\n` — z pustymi linijkami
    na początku, bez ogonowych na końcu."""
    _operat_z_sekcjami(klient)

    _wyslij(klient, **{"sek__wykazy__0__adres_nowy": "\r\n\r\nsdf\r\n"})

    import json
    zapisane = json.loads(db.dokumenty()[0]["dane_json"])["wykazy"]
    assert zapisane == [{"adres_nowy": "\n\nsdf"}]


def test_formularz_nie_zjada_pustej_linijki_przy_poprawianiu(klient):
    """Parser HTML ignoruje pierwszy znak nowej linii po `<textarea>` — bez dodatkowego
    złamania w szablonie każda runda „Popraw ten operat” zjadałaby jedną pustą linijkę
    z początku wartości, po cichu i kumulatywnie."""
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "wykazy", "etykieta": "Wykazy", "typ": "sekcje",
                        "podpola": [{"klucz": "ofu_nowy", "etykieta": "OFU",
                                     "typ": "textarea"}]}]})
    _wyslij(klient, **{"sek__wykazy__0__ofu_nowy": "\r\n\r\nsdf"})

    formularz = klient.get(f"/nowy/spis_tresci_wzor?edytuj={db.dokumenty()[0]['id']}").text

    # dodatkowe złamanie z szablonu + dwie puste linijki brata; parser zje to pierwsze
    assert ">\n\n\nsdf</textarea>" in formularz, \
        "pierwsza pusta linijka zniknie przy każdym otwarciu formularza"


def test_puste_linijki_wchodza_do_dokumentu(baza):
    """Obie drogi do formatki: stan nowy przez `{{r }}`/RichText, dotychczasowy przez
    zwykłe `{{ }}` — puste linijki mają wyjść w komórce jako złamania wiersza."""
    d = _wykaz_dzialek([{"dzialka": "119/80",
                         "ofu_dotychczas": "sdf\nsdf\nsdf",
                         "ofu_nowy": "\n\nsdf"}])

    komorki = {c.text: c for t in d.tables for w in t.rows for c in w.cells
               if "sdf" in c.text}
    assert "\n\nsdf" in komorki, "puste linijki stanu nowego nie doszły do dokumentu"
    assert komorki["\n\nsdf"]._tc.xml.count("<w:br/>") == 2
    assert "sdf\nsdf\nsdf" in komorki


