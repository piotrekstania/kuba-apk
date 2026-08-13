"""Biblioteka gotowych opisów do sprawozdania (Ustawienia).

Pierwszy krok większej przebudowy: na razie to sama biblioteka — dokładanie, lista
i usuwanie. **Nic jeszcze nie wstawia opisu do dokumentu**, więc testy sprawdzają
wyłącznie to, co program dziś obiecuje.

Rzecz, o którą tu najbardziej chodzi: opisy siedzą w bazie, czyli w `dane/` — a to
znaczy, że przeżywają aktualizację programu (`szablony/` jest lustrzane i kasowane,
`dane/` nietykalne). Brat pisze je raz.
"""
from __future__ import annotations

import json

from app import db


def _dodaj(klient, nazwa: str, opis: str):
    return klient.post("/ustawienia/opisy", data={"nazwa": nazwa, "opis": opis},
                       follow_redirects=False)


def test_dodany_opis_widac_w_ustawieniach(klient):
    odpowiedz = _dodaj(klient, "Mapa do celów projektowych",
                       "Pomiar sytuacyjno-wysokościowy wykonano metodą RTN GNSS.")

    assert odpowiedz.status_code == 303
    strona = klient.get("/ustawienia").text
    assert "Mapa do celów projektowych" in strona
    assert "metodą RTN GNSS" in strona


def test_opis_zapisuje_sie_z_obydwoma_polami(klient):
    _dodaj(klient, "Rozgraniczenie", "Ustalono przebieg granic.")

    zapisane = db.opisy_sprawozdania()

    assert len(zapisane) == 1
    assert zapisane[0]["nazwa"] == "Rozgraniczenie"
    assert zapisane[0]["opis"] == "Ustalono przebieg granic."


def test_lamanie_wierszy_zostaje(klient):
    """Opis bywa kilkoma akapitami — zlepiony w jeden przestaje być tym, co zapisywał."""
    _dodaj(klient, "Podział", "Pierwszy akapit.\n\nDrugi akapit.")

    assert db.opisy_sprawozdania()[0]["opis"] == "Pierwszy akapit.\n\nDrugi akapit."


def test_mozna_dolozyc_kilka(klient):
    _dodaj(klient, "Podział", "a")
    _dodaj(klient, "Rozgraniczenie", "b")
    _dodaj(klient, "Mapa do celów projektowych", "c")

    assert len(db.opisy_sprawozdania()) == 3


def test_lista_idzie_po_nazwie(klient):
    """Lista rośnie, a szuka się w niej po nazwie — nie po kolejności dodawania."""
    _dodaj(klient, "Rozgraniczenie", "a")
    _dodaj(klient, "mapa do celów projektowych", "b")
    _dodaj(klient, "Podział", "c")

    nazwy = [w["nazwa"] for w in db.opisy_sprawozdania()]

    assert nazwy == ["mapa do celów projektowych", "Podział", "Rozgraniczenie"], \
        "wielkość liter nie może rozbijać porządku alfabetycznego"


# --- czego nie zapisujemy ----------------------------------------------------

def test_pusta_nazwa_nie_przechodzi(klient):
    """Opis bez nazwy jest nie do rozpoznania na liście."""
    odpowiedz = _dodaj(klient, "   ", "Jakaś treść.")

    assert db.opisy_sprawozdania() == []
    assert "blad=" in odpowiedz.headers["location"]


def test_pusta_tresc_nie_przechodzi(klient):
    """Pozycja, która niczego nie wstawi, to zagadka na liście, a nie opis."""
    odpowiedz = _dodaj(klient, "Podział", "  ")

    assert db.opisy_sprawozdania() == []
    assert "blad=" in odpowiedz.headers["location"]


def test_blad_mowi_po_polsku(klient):
    """Brat nie odróżni „422 Unprocessable Entity” od awarii dysku."""
    _dodaj(klient, "", "")

    strona = klient.get("/ustawienia?blad=Uzupe%C5%82nij%20nazw%C4%99").text

    assert "Uzupełnij nazwę" in strona


def test_biale_znaki_obcinamy(klient):
    _dodaj(klient, "  Podział  ", "  Ustalono granice.  ")

    zapisany = db.opisy_sprawozdania()[0]

    assert zapisany["nazwa"] == "Podział"
    assert zapisany["opis"] == "Ustalono granice."


def test_biale_znaki_obcina_takze_sam_zapis(srodowisko):
    """Obcinanie stoi w warstwie danych, nie tylko w trasie.

    Trasa i tak przycina, żeby sprawdzić, czy pole nie jest puste — ale przez tę
    funkcję będą wołać kolejne kroki przebudowy, a wtedy „ Podział ” i „Podział”
    byłyby na liście dwiema pozycjami nie do odróżnienia okiem.
    """
    db.dodaj_opis_sprawozdania("  Podział  ", "  Ustalono granice.  ")

    zapisany = db.opisy_sprawozdania()[0]

    assert zapisany["nazwa"] == "Podział"
    assert zapisany["opis"] == "Ustalono granice."


# --- usuwanie ----------------------------------------------------------------

def test_usuwanie_opisu(klient):
    _dodaj(klient, "Podział", "a")
    _dodaj(klient, "Rozgraniczenie", "b")
    do_usuniecia = [w for w in db.opisy_sprawozdania() if w["nazwa"] == "Podział"][0]

    klient.post("/ustawienia/opisy/usun", data={"opis": str(do_usuniecia["id"])},
                follow_redirects=False)

    assert [w["nazwa"] for w in db.opisy_sprawozdania()] == ["Rozgraniczenie"]


def test_usuwanie_nieistniejacego_nie_wywala_strony(klient):
    """Dwa razy kliknięte „Usuń” (albo stara karta) nie mogą dać angielskiego błędu."""
    odpowiedz = klient.post("/ustawienia/opisy/usun", data={"opis": "999"},
                            follow_redirects=False)

    assert odpowiedz.status_code == 303


def test_smieciowy_identyfikator_nie_wywala_strony(klient):
    odpowiedz = klient.post("/ustawienia/opisy/usun", data={"opis": "abc"},
                            follow_redirects=False)

    assert odpowiedz.status_code == 303
    assert db.opisy_sprawozdania() == []


# --- to, po co w ogóle trzymamy to w bazie -----------------------------------

def test_opisy_przezywaja_aktualizacje_programu(srodowisko, tmp_path, monkeypatch):
    """Sedno wyboru miejsca: `dane/` jest nietykalne, `szablony/` lustrzane.

    Gdyby opisy leżały w plikach obok szablonów, zniknęłyby przy najbliższej
    aktualizacji — plik, którego nie ma w repozytorium, jest u brata kasowany —
    a dowiedziałby się o tym w środku roboty. Aktualizacja jedzie tu naprawdę,
    z paczki podstawionej przez `file://`.
    """
    from app import aktualizacja
    from tests.test_aktualizacja import _paczka, _podstaw_github

    db.dodaj_opis_sprawozdania("Podział", "Ustalono przebieg granic.")
    (srodowisko.katalog / "WERSJA").write_text("2026.01.01.1\nStara.", encoding="utf-8")
    (srodowisko.katalog / "app").mkdir(exist_ok=True)
    (srodowisko.katalog / "app" / "main.py").write_text("# stary kod", encoding="utf-8")

    paczka = _paczka(tmp_path, "2026.09.09.9\nNowości.", {"app/main.py": "# nowy kod"})
    _podstaw_github(monkeypatch, tmp_path, "2026.09.09.9\nNowości.", paczka)

    assert aktualizacja.sprawdz_i_zaktualizuj() is True
    assert (srodowisko.katalog / "app" / "main.py").read_text() == "# nowy kod", \
        "aktualizacja miała naprawdę przejść — inaczej test niczego nie dowodzi"

    zapisane = db.opisy_sprawozdania()
    assert [(w["nazwa"], w["opis"]) for w in zapisane] == \
        [("Podział", "Ustalono przebieg granic.")]


# --- biblioteka w formularzu operatu -----------------------------------------
#
# Drugi krok: gotowy opis daje się wkleić do pola „Przebieg wykonanych prac”. Który
# to pole, mówi `"biblioteka"` w `.json` szablonu — klucza pola nie ma w kodzie, więc
# bibliotekę da się przepiąć pod inne pole bez ruszania programu.

def _operat_z_opisem(klient):
    # `{%p if %}` kasuje **cały akapit**, w którym stoi, więc warunek, treść i `endif`
    # muszą być trzema osobnymi akapitami — tak samo jak w prawdziwej formatce.
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor",
        ["{{ nr_roboty }}",
         "{%p if opis_przebiegu_jest %}",
         "Przebieg wykonanych prac: {{ opis_przebiegu }}",
         "{%p endif %}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "opis_przebiegu", "etykieta": "Przebieg wykonanych prac",
                        "typ": "textarea", "biblioteka": "sprawozdanie"}]})


def test_lista_gotowych_opisow_jest_w_formularzu(klient):
    db.dodaj_opis_sprawozdania("Mapa do celów projektowych", "Pomiar metodą RTN GNSS.")
    _operat_z_opisem(klient)

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert "Mapa do celów projektowych" in formularz
    assert "Pomiar metodą RTN GNSS." in formularz, "treść musi dojechać, bo wkleja ją JS"
    assert 'data-wklej="p_opis_przebiegu"' in formularz


def test_bez_zapisanych_opisow_formularz_odsyla_do_ustawien(klient):
    """Pusta lista rozwijana niczego nie tłumaczy — ma być powiedziane, gdzie je dodać."""
    _operat_z_opisem(klient)

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert "/ustawienia#opisy" in formularz
    # samego napisu „data-wklej” szukać nie można — jest w skrypcie na dole strony;
    # chodzi o atrybut z wartością, czyli o wyrenderowany przycisk
    assert 'data-wklej="' not in formularz, "nie ma czego ładować, więc nie ma przycisku"


def test_pole_bez_biblioteki_nie_dostaje_listy(klient):
    """Biblioteka pokazuje się tam, gdzie mówi o niej `.json` — a nie przy każdym polu."""
    db.dodaj_opis_sprawozdania("Podział", "Ustalono granice.")
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }} {{ uwagi }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "uwagi", "etykieta": "Uwagi", "typ": "textarea"}]})

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert 'data-wklej="' not in formularz
    assert "Ustalono granice." not in formularz


# --- checkbox „Opis przebiegu” zniknął, a dokument zachowuje się jak przedtem --

def test_nie_ma_juz_checkboxa_opis_przebiegu(klient):
    """Opis jest zawsze — pytanie „czy jest opis” widać po samej treści opisu."""
    _operat_z_opisem(klient)

    formularz = klient.get("/nowy/spis_tresci_wzor").text

    assert "opis_przebiegu_jest" not in formularz


def test_wypelniony_opis_wchodzi_do_dokumentu(klient):
    """`{%p if opis_przebiegu_jest %}` w formatce ma dalej działać — bez checkboxa."""
    from docx import Document
    _operat_z_opisem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__opis_przebiegu": "Pomiar metodą RTN GNSS."},
                follow_redirects=False)

    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    tresc = "\n".join(a.text for a in Document(katalog / "spis_tresci.docx").paragraphs)
    assert "Pomiar metodą RTN GNSS." in tresc


def test_pusty_opis_wycina_sekcje_z_dokumentu(klient):
    """Tak jak przy odznaczonym checkboxie: pusto = warunek fałszywy, sekcja znika.

    To jest ta część, która musiała przetrwać usunięcie checkboxa — inaczej formatka
    brata zaczęłaby po cichu wypisywać pusty opis zamiast „brak”.
    """
    from docx import Document
    _operat_z_opisem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "", "pole__opis_przebiegu": ""},
                follow_redirects=False)

    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    tresc = "\n".join(a.text for a in Document(katalog / "spis_tresci.docx").paragraphs)
    assert "GK.1" in tresc, "reszta dokumentu ma zostać"
    # nie „czy pusto”, tylko czy zniknął **cały akapit** z warunku: przy `_jest`
    # zawsze prawdziwym zostałaby sama etykieta z pustką po dwukropku
    assert "Przebieg wykonanych prac" not in tresc


# --- formatowanie: od formularza po plik Worda -------------------------------

def _operat_z_formatowaniem(klient):
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor",
        ["{{ nr_roboty }}",
         "{%p if opis_przebiegu_jest %}",
         "Przebieg: {{r opis_przebiegu }}",
         "{%p endif %}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "opis_przebiegu", "etykieta": "Przebieg", "typ": "textarea",
                        "formatowanie": True, "biblioteka": "sprawozdanie"}]})


def _akapit_opisu(klient):
    from docx import Document
    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    for akapit in Document(katalog / "spis_tresci.docx").paragraphs:
        if akapit.text.startswith("Przebieg:"):
            return akapit
    return None


def test_pogrubienie_z_formularza_dojezdza_do_dokumentu(klient):
    _operat_z_formatowaniem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__opis_przebiegu": "<b>Pomiar</b> wykonano <i>metodą RTN</i>."},
                follow_redirects=False)

    biegi = {b.text: (b.bold, b.italic) for b in _akapit_opisu(klient).runs if b.text}
    assert biegi["Pomiar"] == (True, None)
    assert biegi["metodą RTN"] == (None, True)


def test_smieci_z_wklejki_nie_wchodza_do_dokumentu(klient):
    """Kopia z Worda ciągnie style i czasem skrypty — obcinamy je przy zapisie."""
    _operat_z_formatowaniem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__opis_przebiegu":
                          '<script>alert(1)</script><span style="color:red">Opis</span>'},
                follow_redirects=False)

    assert json.loads(db.dokumenty()[0]["dane_json"])["opis_przebiegu"] == "Opis"
    assert "alert" not in _akapit_opisu(klient).text


def test_pusty_sformatowany_opis_dalej_wycina_sekcje(klient):
    """Pułapka tej zmiany: wartość przestaje być napisem, a pusty `RichText` jest
    prawdziwy jak każdy obiekt — więc `opis_przebiegu_jest` liczymy, zanim ją podmienimy.
    Bez tego formatka nigdy więcej nie napisałaby „brak”."""
    _operat_z_formatowaniem(klient)

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__opis_przebiegu": "<b></b><br>"},
                follow_redirects=False)

    assert _akapit_opisu(klient) is None, "pusty opis miał wyciąć całą sekcję"


def test_formatowanie_wraca_do_formularza_przy_poprawianiu(klient):
    _operat_z_formatowaniem(klient)
    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__opis_przebiegu": "<b>Pomiar</b> RTN"},
                follow_redirects=False)

    formularz = klient.get(
        f"/nowy/spis_tresci_wzor?edytuj={db.dokumenty()[0]['id']}").text

    assert "<b>Pomiar</b> RTN" in formularz, "edytor ma dostać HTML, a nie zjedzone znaczniki"
    assert 'contenteditable="true"' in formularz


def test_formatowanie_widac_na_stronie_operatu(klient):
    _operat_z_formatowaniem(klient)
    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__opis_przebiegu": "<b>Pomiar</b> RTN"},
                follow_redirects=False)

    strona = klient.get(f"/dokument/{db.dokumenty()[0]['id']}").text

    assert "<b>Pomiar</b> RTN" in strona


def test_pole_bez_formatowania_nadal_eskejpuje(klient):
    """Sąsiednie pola biorą się wprost z tego, co ktoś wpisał — nawias trójkątny
    w uwagach nie ma prawa stać się znacznikiem."""
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }} {{ uwagi }}"],
        opis={"nazwa": "Operat", "glowny": True,
              "pola": [{"klucz": "nr_roboty", "etykieta": "Nr roboty"},
                       {"klucz": "uwagi", "etykieta": "Uwagi", "typ": "textarea"}]})

    klient.post("/generuj/spis_tresci_wzor",
                data={"pole__nr_roboty": "GK.1", "notatka": "",
                      "pole__uwagi": "<b>to nie jest pogrubienie</b>"},
                follow_redirects=False)

    strona = klient.get(f"/dokument/{db.dokumenty()[0]['id']}").text
    assert "&lt;b&gt;to nie jest pogrubienie" in strona


def test_opis_w_bibliotece_tez_jest_czyszczony(klient):
    klient.post("/ustawienia/opisy",
                data={"nazwa": "Wklejka", "opis": '<b>Grube</b><script>x()</script>'},
                follow_redirects=False)

    assert db.opisy_sprawozdania()[0]["opis"] == "<b>Grube</b>"


def test_opis_z_samych_znacznikow_nie_przechodzi(klient):
    """Pusty edytor zostawia po sobie `<br>` — to nadal jest pusty opis."""
    odpowiedz = klient.post("/ustawienia/opisy", data={"nazwa": "Puste", "opis": "<br><b></b>"},
                            follow_redirects=False)

    assert db.opisy_sprawozdania() == []
    assert "blad=" in odpowiedz.headers["location"]


# --- czyszczenie wklejki po stronie serwera ----------------------------------

def test_trasa_czyszczaca_wklejke(klient):
    """Edytor woła to przy każdym wklejeniu.

    Czyszczenie mogłoby siedzieć w JS, ale wtedy lista dozwolonych znaczników istniałaby
    dwa razy i po pierwszej zmianie zaczęłaby się rozjeżdżać. Jest jedna, ta z testami.
    """
    odpowiedz = klient.post("/tekst/oczysc", json={
        "html": '<span style="font-weight:700;font-family:Arial">Grube</span>'
                ' i <span style="color:red">czerwone</span>'})

    assert odpowiedz.json()["html"] == "<b>Grube</b> i czerwone"


def test_trasa_czyszczaca_znosi_pustke(klient):
    assert klient.post("/tekst/oczysc", json={}).json()["html"] == ""
