"""Wypełnianie szablonu, numeracja i nazwy plików."""
from __future__ import annotations

import json
from datetime import date

import pytest

from app import db, generator, szablony

ROK = date.today().year


def _szablon_operatu(srodowisko, nazwa="operat_wzor"):
    srodowisko.dodaj_szablon(
        nazwa,
        ["Robota: {{ nr_roboty }}", "Operat: {{ nr_operatu }}", "Data: {{ data_zakonczenia }}",
         "Słownie: {{ data_zakonczenia_slownie }}", "Uwagi: {{ uwagi }}"],
        opis={"nazwa": "Operat", "glowny": True, "licznik": "operat",
              "wzor_nazwy": "Operat_{nr_roboty}",
              "pola": [
                  {"klucz": "nr_roboty", "wymagane": True},
                  {"klucz": "nr_operatu", "typ": "auto_numer", "domyslnie": "{numer3}/{rok}"},
                  {"klucz": "data_zakonczenia", "typ": "date"},
                  {"klucz": "uwagi", "typ": "textarea"},
              ]})
    return szablony.szablon_po_id(nazwa)


# --- nazwy plików ------------------------------------------------------------

def test_bezpieczna_nazwa_usuwa_ukosnik_i_znaki_zakazane():
    assert generator.bezpieczna_nazwa("001/2026") == "001-2026"
    assert generator.bezpieczna_nazwa('a<b>c:d"e|f?g*h') == "abcdefgh"
    assert generator.bezpieczna_nazwa("  ") == "dokument"


def test_bezpieczna_nazwa_zachowuje_polskie_znaki():
    """„Sułkowice” nie może się zrobić „Sukowice”.

    Litery bez odpowiednika ASCII (ł, Ł) gubią się przy rozkładzie NFKD, bo nie są
    literą + znakiem diakrytycznym. Nazwa pliku przestaje wtedy przypominać to,
    co brat wpisał w formularzu.
    """
    assert generator.bezpieczna_nazwa("Sułkowice") == "Sulkowice"
    assert generator.bezpieczna_nazwa("Łódź") == "Lodz"
    assert generator.bezpieczna_nazwa("Kraków-Podgórze") == "Krakow-Podgorze"


# --- daty --------------------------------------------------------------------

def test_formaty_dat():
    assert generator.data_pl("2026-07-31") == "31.07.2026"
    assert generator.data_slownie("2026-07-31") == "31 lipca 2026 r."
    assert generator.data_slownie("2026-01-01") == "1 stycznia 2026 r."


def test_zla_data_wraca_bez_zmian():
    """Formularz może przysłać cokolwiek — konwersja nie ma prawa wywalić generowania."""
    assert generator.data_pl("wczoraj") == "wczoraj"
    assert generator.data_slownie("wczoraj") == "wczoraj"
    assert generator.data_pl("") == ""


# --- kontekst ----------------------------------------------------------------

def test_kontekst_dokłada_daty_i_numer(srodowisko):
    szablon = _szablon_operatu(srodowisko)
    kontekst = generator.przygotuj_kontekst(
        szablon, {"nr_roboty": "GK.1.2026", "data_zakonczenia": "2026-07-31"}, {})
    assert kontekst["nr_operatu"] == f"001/{ROK}"
    assert kontekst["data_zakonczenia"] == "31.07.2026"
    assert kontekst["data_zakonczenia_iso"] == "2026-07-31"
    assert kontekst["data_zakonczenia_slownie"] == "31 lipca 2026 r."
    assert kontekst["rok"] == str(ROK)


def test_podany_numer_nie_zuzywa_licznika(srodowisko):
    """Poprawianie operatu podstawia stary numer — licznik ma stać w miejscu."""
    szablon = _szablon_operatu(srodowisko)
    generator.przygotuj_kontekst(szablon, {"nr_operatu": "007/2020"}, {})
    assert db.podglad_numeru("operat", ROK) == 1


def test_dane_stale_wchodza_do_kontekstu_ale_formularz_wygrywa(srodowisko):
    szablon = _szablon_operatu(srodowisko)
    kontekst = generator.przygotuj_kontekst(
        szablon, {"nr_roboty": "z formularza"}, {"nr_roboty": "z ustawień", "firma": "ProCAD"})
    assert kontekst["nr_roboty"] == "z formularza"
    assert kontekst["firma"] == "ProCAD"


# --- numeracja ---------------------------------------------------------------

def test_numeracja_idzie_po_kolei(srodowisko):
    assert [db.nastepny_numer("operat", ROK) for _ in range(3)] == [1, 2, 3]
    assert db.podglad_numeru("operat", ROK) == 4


def test_zwolniony_numer_wraca_do_puli(srodowisko):
    """Trzy nieudane próby między dwoma dobrymi dokumentami dają 001 i 002, bez dziury."""
    pierwszy = db.nastepny_numer("operat", ROK)
    for _ in range(3):
        nieudany = db.nastepny_numer("operat", ROK)
        assert db.zwolnij_numer("operat", ROK, nieudany)
    drugi = db.nastepny_numer("operat", ROK)
    assert (pierwszy, drugi) == (1, 2)


def test_numer_nie_cofa_sie_gdy_licznik_ruszyl_dalej(srodowisko):
    """Wolimy dziurę w numeracji niż dwa operaty o tym samym numerze."""
    nieudany = db.nastepny_numer("operat", ROK)
    kolejny = db.nastepny_numer("operat", ROK)          # ktoś zdążył wygenerować dokument
    assert db.zwolnij_numer("operat", ROK, nieudany) is False
    assert db.podglad_numeru("operat", ROK) == kolejny + 1


# --- generowanie -------------------------------------------------------------

def test_generowanie_zaklada_katalog_operatu(srodowisko):
    szablon = _szablon_operatu(srodowisko)
    plik, kontekst, ostrzezenia = generator.generuj(
        szablon, {"nr_roboty": "GK.6640.1.2026", "data_zakonczenia": "2026-07-31"}, {})

    katalog = plik.parent
    assert katalog.name == f"001.{ROK}"            # ukośnik zamieniony na kropkę
    assert plik.name == "spis_tresci.docx" or plik.name == "operat.docx"
    assert (katalog / "operat.json").exists()
    opis = json.loads((katalog / "operat.json").read_text(encoding="utf-8"))
    assert opis["nr_operatu"] == f"001/{ROK}"
    assert opis["nr_roboty"] == "GK.6640.1.2026"
    assert (katalog / "GK.6640.1.2026").exists()   # pusty znacznik z numerem roboty
    assert ostrzezenia == []


def test_poprawianie_nie_zuzywa_numeru_i_wraca_do_tego_samego_katalogu(srodowisko):
    szablon = _szablon_operatu(srodowisko)
    plik, kontekst, _ = generator.generuj(szablon, {"nr_roboty": "GK.1.2026"}, {})
    poprzedni = json.loads((plik.parent / "operat.json").read_text(encoding="utf-8"))

    # tak jak z przeglądarki: pole numeru **jest** w danych, tylko puste
    plik2, kontekst2, _ = generator.generuj(
        szablon, {"nr_roboty": "GK.1.2026", "nr_operatu": ""}, {}, poprzedni)

    assert kontekst2["nr_operatu"] == kontekst["nr_operatu"]
    assert plik2.parent == plik.parent
    assert db.podglad_numeru("operat", ROK) == 2      # licznik ruszył tylko raz


def test_bledny_szablon_oddaje_numer(srodowisko):
    """Wywrotka przy wypełnianiu nie może zostawiać dziury w numeracji operatów.

    Numer trzeba zarezerwować **przed** wypełnieniem (wchodzi do treści dokumentu),
    więc gdy wypełnianie padnie, trzeba go oddać.
    """
    srodowisko.dodaj_szablon(
        "zly", ["{{ nr_operatu }}", "{{ 1/0 }}"],      # składnia dobra, wykonanie pada
        opis={"glowny": True, "licznik": "operat",
              "pola": [{"klucz": "nr_operatu", "typ": "auto_numer",
                        "domyslnie": "{numer3}/{rok}"}]})
    szablon = szablony.szablon_po_id("zly")
    with pytest.raises(Exception):
        generator.generuj(szablon, {"nr_roboty": "GK.1.2026"}, {})
    assert db.podglad_numeru("operat", ROK) == 1


def test_literowka_w_znaczniku_nie_wywala_listy_szablonow(srodowisko):
    """Formatka z zepsutą składnią Jinja nie może zabrać strony głównej.

    Brat podmienia .docx sam; gdy pomyli `{%p if %}` z `{% if %}`, ma zobaczyć jeden
    szablon z ostrzeżeniem, a nie pustą aplikację.
    """
    srodowisko.dodaj_szablon("dobry", ["{{ a }}"])
    srodowisko.dodaj_szablon("literowka", ["{% for x in %}"])
    lista = {s.id: s for s in szablony.lista_szablonow()}
    assert set(lista) == {"dobry", "literowka"}
    assert lista["literowka"].opis.startswith("⚠")
    assert not lista["dobry"].opis.startswith("⚠")


def test_dokument_dodatkowy_dostaje_ten_sam_numer(srodowisko):
    """`dopisz_dokument` nie może sięgnąć po kolejny numer z licznika."""
    glowny = _szablon_operatu(srodowisko)
    srodowisko.dodaj_szablon("sprawozdanie_wzor", ["Operat: {{ nr_operatu }}"],
                             opis={"nazwa": "Sprawozdanie", "pola": []})
    dodatkowy = szablony.szablon_po_id("sprawozdanie_wzor")

    plik, kontekst, _ = generator.generuj(glowny, {"nr_roboty": "GK.1.2026"}, {})
    drugi = generator.dopisz_dokument(dodatkowy, kontekst, plik.parent)

    assert drugi.parent == plik.parent
    assert drugi.name == "sprawozdanie.docx"
    assert db.podglad_numeru("operat", ROK) == 2

    from docx import Document
    tresc = "\n".join(p.text for p in Document(drugi).paragraphs)
    assert kontekst["nr_operatu"] in tresc


def test_wygenerowany_dokument_nie_zostawia_znacznikow(srodowisko):
    """W gotowym pliku nie ma prawa zostać {{ ani {%."""
    szablon = _szablon_operatu(srodowisko)
    plik, _, _ = generator.generuj(
        szablon, {"nr_roboty": "GK.1.2026", "data_zakonczenia": "2026-07-31",
                  "uwagi": "bez uwag"}, {})
    from docx import Document
    tresc = "\n".join(p.text for p in Document(plik).paragraphs)
    assert "{{" not in tresc and "{%" not in tresc
    assert "GK.1.2026" in tresc and "31.07.2026" in tresc
