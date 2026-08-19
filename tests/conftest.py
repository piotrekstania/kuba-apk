"""Wspólne rusztowanie testów.

Najważniejsza sztuczka: moduły robią `from .config import WYNIKI`, więc ścieżki są
**przywiązane do modułu w chwili importu**. Podmiana samego `app.config` niczego by nie
dała — trzeba podmienić nazwę w każdym module, który ją u siebie trzyma. Robi to
fixture `srodowisko`, dzięki czemu żaden test nie dotyka prawdziwych `wyniki/` i `dane/`.
"""
from __future__ import annotations

import sys
import threading
import time
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

KORZEN = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(KORZEN))

# Ile czekamy na wątki tła przy sprzątaniu testu. Limit jest z zapasem, bo czekanie
# kosztuje tylko wtedy, gdy coś naprawdę wisi: w większości testów konwerter jest
# atrapą i wątki kończą się w ułamku sekundy, ale `test_word.py` puszcza prawdziwego
# Worda przez komplet dokumentów operatu. Strażnik ma łapać wątki, które przeżyły test,
# a nie zamieniać się w nowe źródło losowych padnięć.
LIMIT_WATKOW = 60.0


def _zaczekaj_na_watki(wczesniejsze: set[threading.Thread]) -> None:
    """Czeka na wątki tła wystartowane w trakcie testu — zanim znikną podmienione ścieżki.

    Program startuje w tle dwie rzeczy: pierwsze pobranie TERYT-u (`cykl_zycia`)
    i przygotowanie podglądów PDF po wygenerowaniu operatu (`/generuj`). Oba wątki
    są demoniczne, więc test kończy się, nie oglądając się na nie, a one czytają
    `db.BAZA_DANYCH`, `operaty.WYNIKI` i resztę **w chwili wywołania**, nie startu.
    Wątek, który dożyje sprzątnięcia monkeypatcha, dostaje do ręki prawdziwe `dane/`
    i `wyniki/` autora albo katalogi kolejnego testu — i pisze tam po cichu, bo oba
    te wątki łykają każdy wyjątek. Dlatego czekamy na nie, póki podmiana jeszcze stoi.
    """
    koniec = time.monotonic() + LIMIT_WATKOW
    obce = lambda: [w for w in threading.enumerate()                      # noqa: E731
                    if w not in wczesniejsze and w is not threading.current_thread()]
    for watek in obce():
        watek.join(max(0.0, koniec - time.monotonic()))
    zostaly = [w.name for w in obce() if w.is_alive()]
    if zostaly:
        pytest.fail("wątki w tle nie skończyły się przed sprzątnięciem ścieżek testu "
                    f"(po {LIMIT_WATKOW:.0f} s żyją: {', '.join(zostaly)}) — dopóki żyją, "
                    "piszą tam, gdzie akurat pokazują ścieżki modułów")


@pytest.fixture(autouse=True)
def bez_wysylki_statystyk(monkeypatch):
    """Żaden test nie wysyła niczego do arkusza autora.

    Autouse i przez **ten sam wyłącznik co w programie**, więc chroni wszystkie testy
    naraz — także te, które stawiają własny TestClient (`test_word.py`) i przez to
    uruchamiają cykl życia aplikacji razem z wątkiem wysyłki.
    """
    from app import raport
    monkeypatch.setenv(raport.WYLACZNIK, "1")


@pytest.fixture(autouse=True)
def bez_prawdziwej_bazy(tmp_path, monkeypatch):
    """Żaden test nie pisze do bazy autora — nawet ten, który nie bierze `srodowisko`.

    Testy sprawdzające **wydane formatki** celowo nie mogą wziąć `srodowisko`: ono
    podmienia `szablony/` na pusty katalog, a one chcą prawdziwych plików brata.
    Przez to nie miały podmienionej też bazy i szły do `dane/operaty.sqlite3` autora —
    `przygotuj_kontekst` rezerwuje numer operatu, więc **każde uruchomienie testów
    podbijało mu licznik numeracji** (doszedł do 463 przy czterech operatach w historii).
    Widać to było dopiero w CI, gdzie bazy nie ma w ogóle: `no such table: liczniki`.

    Dlatego strażnik jest autouse i stoi obok `bez_wysylki_statystyk`: chroni wszystkie
    testy naraz, także te pisane w przyszłości. Testy z własnym środowiskiem nadpisują
    tę ścieżkę swoją (fixtury autouse idą pierwsze), więc niczego im nie psuje.
    """
    from app import aktualizacja, config, db
    baza = tmp_path / "straznik.sqlite3"
    monkeypatch.setattr(config, "BAZA_DANYCH", baza, raising=False)
    monkeypatch.setattr(db, "BAZA_DANYCH", baza)
    monkeypatch.setattr(aktualizacja, "BAZA_DANYCH", baza)
    db.init()


# --- pomocnicze budowanie dokumentów -----------------------------------------

def zbuduj_docx(sciezka: Path, akapity: list[str | tuple], **kwargs) -> Path:
    """Prosty dokument z listy akapitów.

    Element listy to napis albo krotka (tekst, rozmiar_pt, pogrubienie).
    """
    dokument = Document()
    for pozycja in akapity:
        tekst, rozmiar, pogrubienie = (
            pozycja if isinstance(pozycja, tuple) else (pozycja, 10, False))
        akapit = dokument.add_paragraph()
        bieg = akapit.add_run(tekst)
        bieg.font.size = Pt(rozmiar)
        bieg.bold = pogrubienie
    if kwargs.get("tabela"):
        tabela = dokument.add_table(rows=4, cols=2)
        tabela.rows[0].cells[0].text = "Nagłówek A"
        tabela.rows[0].cells[1].text = "Nagłówek B"
        tabela.rows[1].cells[0].text = "{%tr for w in wiersze %}"
        tabela.rows[2].cells[0].text = "{{ w.a }}"
        tabela.rows[2].cells[1].text = "{{ w.b }}"
        tabela.rows[3].cells[0].text = "{%tr endfor %}"
    sciezka.parent.mkdir(parents=True, exist_ok=True)
    dokument.save(sciezka)
    return sciezka


@pytest.fixture
def srodowisko(tmp_path, monkeypatch):
    """Izolowana instalacja programu: własne szablony/, wyniki/, dane/ i baza.

    Zwraca obiekt z gotowymi ścieżkami i skrótem `dodaj_szablon`.

    Sprzątając, czeka na wątki tła wystartowane w trakcie testu (`_zaczekaj_na_watki`).
    Robi to **tutaj**, a nie w `klient`, bo własny TestClient stawia też `test_word.py`,
    a wątki startują nie tylko z cyklu życia aplikacji. Kolejność jest bezpieczna:
    `monkeypatch` powstaje przed tą fixture, więc jego sprzątanie idzie po naszym
    i przez cały czas oczekiwania ścieżki są jeszcze podmienione.
    """
    from app import (aktualizacja, config, db, generator, main, operaty, opisy, pdf,
                     raport, szablony, warianty)

    watki_przed = set(threading.enumerate())

    szablony_kat = tmp_path / "szablony"
    wyniki_kat = tmp_path / "wyniki"
    dane_kat = tmp_path / "dane"
    for katalog in (szablony_kat, wyniki_kat, dane_kat):
        katalog.mkdir(parents=True, exist_ok=True)
    baza = dane_kat / "operaty.sqlite3"

    monkeypatch.setattr(config, "BAZA", tmp_path, raising=False)
    monkeypatch.setattr(config, "SZABLONY", szablony_kat, raising=False)
    monkeypatch.setattr(config, "WYNIKI", wyniki_kat, raising=False)
    monkeypatch.setattr(config, "DANE", dane_kat, raising=False)
    monkeypatch.setattr(config, "BAZA_DANYCH", baza, raising=False)

    monkeypatch.setattr(db, "BAZA_DANYCH", baza)
    monkeypatch.setattr(db, "DANE", dane_kat)
    monkeypatch.setattr(szablony, "SZABLONY", szablony_kat)
    monkeypatch.setattr(opisy, "SZABLONY", szablony_kat)
    monkeypatch.setattr(warianty, "KATALOG", dane_kat / "szablony")
    monkeypatch.setattr(operaty, "WYNIKI", wyniki_kat)
    monkeypatch.setattr(operaty, "DANE", dane_kat)
    monkeypatch.setattr(operaty, "PODGLADY", dane_kat / "podglad")
    monkeypatch.setattr(pdf, "KATALOG_ROBOCZY", dane_kat / "konwersja")
    monkeypatch.setattr(main, "WYNIKI", wyniki_kat)
    monkeypatch.setattr(main, "DANE", dane_kat)
    monkeypatch.setattr(main, "DZIENNIK_BLEDOW", dane_kat / "bledy.log")
    monkeypatch.setattr(aktualizacja, "BAZA", tmp_path)
    monkeypatch.setattr(aktualizacja, "DANE", dane_kat)
    monkeypatch.setattr(aktualizacja, "BAZA_DANYCH", baza)
    monkeypatch.setattr(raport, "BAZA", tmp_path)
    monkeypatch.setattr(raport, "DANE", dane_kat)
    monkeypatch.setattr(raport, "PLIK_ID", dane_kat / "instalacja.txt")
    monkeypatch.setattr(raport, "PLIK_ETYKIETY", dane_kat / "etykieta.txt")
    monkeypatch.setattr(aktualizacja, "PLIK_WERSJI", tmp_path / "WERSJA")
    monkeypatch.setattr(aktualizacja, "KOPIE", dane_kat / "kopie")
    monkeypatch.setattr(aktualizacja, "ZNACZNIK_NOWOSCI", dane_kat / "co_nowego.txt")

    db.init()

    class Srodowisko:
        katalog = tmp_path
        szablony = szablony_kat
        wyniki = wyniki_kat
        dane = dane_kat
        baza_danych = baza

        @staticmethod
        def dodaj_szablon(nazwa: str, akapity: list, opis: dict | None = None, **kwargs):
            import json
            plik = zbuduj_docx(szablony_kat / f"{nazwa}.docx", akapity, **kwargs)
            if opis is not None:
                plik.with_suffix(".json").write_text(
                    json.dumps(opis, ensure_ascii=False), encoding="utf-8")
            return plik

    # Generator woła teryt.jednostka() przy polach typu „teryt” — bez bazy TERYT
    # oddaje puste napisy i to nam w testach wystarczy.
    monkeypatch.setattr(generator.teryt, "jednostka", lambda identyfikator: None)
    monkeypatch.setattr(generator.teryt, "obreb", lambda identyfikator: None)
    yield Srodowisko
    _zaczekaj_na_watki(watki_przed)


@pytest.fixture
def bez_konwertera(monkeypatch):
    """Podmienia konwersję DOCX→PDF na atrapę: testy mają chodzić bez Worda i LibreOffice."""
    from app import operaty, pdf

    def atrapa(zrodlo, cel=None):
        """Udaje konwersję, ale oddaje **prawdziwy** jednostronicowy PDF.

        Zaślepka w rodzaju `b"%PDF-1.4"` przechodzi przez `exists()`, ale wykłada się
        dopiero w pypdf przy sklejaniu — test przechodziłby wtedy z niewłaściwego powodu.
        """
        from pypdf import PdfWriter
        cel = Path(cel) if cel else Path(zrodlo).with_suffix(".pdf")
        cel.parent.mkdir(parents=True, exist_ok=True)
        zapis = PdfWriter()
        zapis.add_blank_page(width=595, height=842)      # A4 w punktach
        with open(cel, "wb") as wyjscie:
            zapis.write(wyjscie)
        return cel

    monkeypatch.setattr(pdf, "docx_na_pdf", atrapa)
    monkeypatch.setattr(pdf, "docx_na_pdf_wsad",
                        lambda pary: [atrapa(z, c) for z, c in pary])
    monkeypatch.setattr(pdf, "dostepny_konwerter", lambda: "libreoffice")
    monkeypatch.setattr(operaty.pdf, "docx_na_pdf", atrapa)
    monkeypatch.setattr(operaty.pdf, "docx_na_pdf_wsad",
                        lambda pary: [atrapa(z, c) for z, c in pary])
    return atrapa


@pytest.fixture
def klient(srodowisko, bez_konwertera, monkeypatch):
    """TestClient z wyłączonym pobieraniem TERYT-u (żaden test nie rusza sieci)."""
    from fastapi.testclient import TestClient

    from app import main, teryt

    monkeypatch.setattr(teryt, "pusto", lambda: False)
    monkeypatch.setattr(main.teryt, "pusto", lambda: False)
    with TestClient(main.app, raise_server_exceptions=False) as klient_testowy:
        klient_testowy.srodowisko = srodowisko
        yield klient_testowy
