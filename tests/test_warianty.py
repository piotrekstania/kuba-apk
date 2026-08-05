"""Własne formatki użytkownika — wiele wariantów do jednego rodzaju dokumentu.

Najważniejszy test w tym pliku to `test_wlasna_formatka_przezywa_aktualizacje`.
Katalog `szablony/` jest lustrzany, więc formatka wgrana tam zostałaby u brata
skasowana przy najbliższym wydaniu — i dowiedziałby się o tym dopiero wtedy,
gdy jego operat wyszedłby ze standardowego wzoru.
"""
from __future__ import annotations

import json

from app import db, szablony, warianty
from conftest import zbuduj_docx

OPIS_OPERATU = {
    "nazwa": "Operat", "glowny": True, "licznik": "operat",
    "pola": [
        {"klucz": "nr_roboty", "etykieta": "Nr roboty", "wymagane": True},
        {"klucz": "nr_operatu", "typ": "auto_numer", "domyslnie": "{numer3}/{rok}"},
    ]}
FORMULARZ = {"pole__nr_roboty": "GK.1.2026", "pole__nr_operatu": ""}


def _kategoria(srodowisko, identyfikator="spis_tresci_wzor"):
    srodowisko.dodaj_szablon(identyfikator, ["Standardowa: {{ nr_roboty }}"],
                             opis=OPIS_OPERATU)
    return identyfikator


def _zaznaczona_formatka(html: str, kategoria: str) -> str:
    """Wartość zaznaczonej opcji w liście formatek danej kategorii.

    Czytamy to z HTML-a zamiast szukać w nim napisów: „selected” gdziekolwiek na stronie
    przechodziłoby też wtedy, gdy zaznaczona jest zupełnie inna lista.
    """
    import re
    lista = re.search(rf'<select name="wariant__{re.escape(kategoria)}">(.*?)</select>',
                      html, re.DOTALL)
    assert lista, f"nie ma listy formatek dla {kategoria}"
    wybrana = re.search(r'value="([^"]*)"\s*\n?\s*selected', lista.group(1))
    return wybrana.group(1) if wybrana else ""


def _wgraj(srodowisko, kategoria, nazwa="Moja formatka.docx", akapity=None):
    zrodlo = zbuduj_docx(srodowisko.katalog / "wgrywane" / nazwa,
                         akapity or ["Moja: {{ nr_roboty }}"])
    with open(zrodlo, "rb") as plik:
        return warianty.dodaj(kategoria, nazwa, plik)


# --- gdzie lądują pliki ------------------------------------------------------

def test_wlasna_formatka_nie_lezy_w_katalogu_szablony(srodowisko):
    kategoria = _kategoria(srodowisko)
    wariant, ostrzezenia = _wgraj(srodowisko, kategoria)

    assert ostrzezenia == []
    assert warianty.plik(wariant["id"]).parent.parent == srodowisko.dane / "szablony"
    assert [p.name for p in srodowisko.szablony.glob("*.docx")] == ["spis_tresci_wzor.docx"]


def test_wlasna_formatka_przezywa_aktualizacje(srodowisko, monkeypatch, tmp_path):
    """Katalog `szablony/` jest lustrzany — dlatego własne formatki są w `dane/`."""
    from app import aktualizacja
    from test_aktualizacja import _instalacja, _paczka, _podstaw_github

    _instalacja(srodowisko)
    kategoria = _kategoria(srodowisko)
    wariant, _ = _wgraj(srodowisko, kategoria)

    paczka = _paczka(tmp_path, "2026.09.09.9", {"szablony/spis_tresci_wzor.docx": "nowa"})
    _podstaw_github(monkeypatch, tmp_path, "2026.09.09.9", paczka)
    assert aktualizacja.sprawdz_i_zaktualizuj() is True

    assert warianty.plik(wariant["id"]) is not None, "aktualizacja zjadła własną formatkę"
    assert warianty.lista(kategoria)


def test_druga_formatka_o_tej_samej_nazwie_nie_nadpisuje_pierwszej(srodowisko):
    kategoria = _kategoria(srodowisko)
    pierwsza, _ = _wgraj(srodowisko, kategoria, "Wzor.docx")
    druga, _ = _wgraj(srodowisko, kategoria, "Wzor.docx")

    assert pierwsza["id"] != druga["id"]
    assert len(warianty.lista(kategoria)) == 2


def test_kategoria_z_adresu_nie_wyprowadza_poza_dane(srodowisko):
    _kategoria(srodowisko)
    assert warianty.plik("../../etc/passwd") is None
    assert warianty.lista("../..") == []


# --- sprawdzanie wgrywanego pliku --------------------------------------------

def test_nieznane_znaczniki_daja_ostrzezenie_ale_plik_zostaje(srodowisko):
    """Formatka różniąca się jednym polem to najczęstszy przypadek, nie błąd.

    Ostrzegamy, bo puste miejsce w gotowym operacie nie zgłosi się samo — ale
    nie blokujemy, tak samo jak przy sprawdzaniu numeru działki w ULDK.
    """
    kategoria = _kategoria(srodowisko)
    wariant, ostrzezenia = _wgraj(srodowisko, kategoria, "Inna.docx",
                                  ["{{ nr_roboty }} {{ wlasne_pole }}"])

    assert warianty.plik(wariant["id"]) is not None
    assert len(ostrzezenia) == 1
    assert "wlasne_pole" in ostrzezenia[0]


def test_plik_ktory_nie_jest_szablonem_nie_zostaje_na_dysku(srodowisko):
    import io
    import pytest

    kategoria = _kategoria(srodowisko)
    with pytest.raises(warianty.BladWariantu):
        warianty.dodaj(kategoria, "smiec.docx", io.BytesIO(b"to nie jest docx"))
    assert warianty.lista(kategoria) == []


def test_odrzucamy_plik_w_zlym_formacie(srodowisko):
    import io
    import pytest

    kategoria = _kategoria(srodowisko)
    with pytest.raises(warianty.BladWariantu) as awaria:
        warianty.dodaj(kategoria, "stara_formatka.doc", io.BytesIO(b"x"))
    assert ".docx" in str(awaria.value)


# --- wybór przy generowaniu --------------------------------------------------

def test_wybrana_formatka_wchodzi_do_dokumentu(klient):
    kategoria = _kategoria(klient.srodowisko)
    wariant, _ = _wgraj(klient.srodowisko, kategoria)

    klient.post(f"/generuj/{kategoria}",
                data={**FORMULARZ, f"wariant__{kategoria}": wariant["id"]},
                follow_redirects=False)

    from docx import Document
    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    tresc = "\n".join(p.text for p in Document(katalog / "spis_tresci.docx").paragraphs)
    assert "Moja:" in tresc
    assert "Standardowa:" not in tresc


def test_dokument_zachowuje_nazwe_kategorii(klient):
    """Plik nazywa się jak rodzaj dokumentu, nie jak wybrana formatka."""
    kategoria = _kategoria(klient.srodowisko)
    wariant, _ = _wgraj(klient.srodowisko, kategoria)

    klient.post(f"/generuj/{kategoria}",
                data={**FORMULARZ, f"wariant__{kategoria}": wariant["id"]},
                follow_redirects=False)

    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    assert (katalog / "spis_tresci.docx").exists()


def test_wybor_zapamietuje_sie_do_nastepnego_operatu(klient):
    kategoria = _kategoria(klient.srodowisko)
    wariant, _ = _wgraj(klient.srodowisko, kategoria)

    klient.post(f"/generuj/{kategoria}",
                data={**FORMULARZ, f"wariant__{kategoria}": wariant["id"]},
                follow_redirects=False)
    formularz = klient.get(f"/nowy/{kategoria}").text

    assert _zaznaczona_formatka(formularz, kategoria) == wariant["id"]


def test_poprawianie_bierze_formatke_zapisana_przy_operacie(klient):
    """Poprawka literówki nie może podmienić formatki w gotowym operacie.

    Brat generuje operat formatką A, potem ustawia domyślną na B (bo tak robi
    kolejne roboty) i wraca poprawić literówkę w tamtym pierwszym. Ma dostać A.
    """
    kategoria = _kategoria(klient.srodowisko)
    pierwsza, _ = _wgraj(klient.srodowisko, kategoria, "A.docx", ["A: {{ nr_roboty }}"])
    druga, _ = _wgraj(klient.srodowisko, kategoria, "B.docx", ["B: {{ nr_roboty }}"])

    klient.post(f"/generuj/{kategoria}",
                data={**FORMULARZ, f"wariant__{kategoria}": pierwsza["id"]},
                follow_redirects=False)
    stary = db.dokumenty()[0]
    klient.post(f"/generuj/{kategoria}",
                data={**FORMULARZ, "pole__nr_roboty": "GK.2.2026",
                      f"wariant__{kategoria}": druga["id"]},
                follow_redirects=False)

    # domyślna jest już B...
    assert warianty.domyslne(db.wczytaj_ustawienia())[kategoria] == druga["id"]
    # ...ale przy poprawianiu tamtego operatu formularz pokazuje A
    formularz = klient.get(f"/nowy/{kategoria}?edytuj={stary['id']}").text
    assert _zaznaczona_formatka(formularz, kategoria) == pierwsza["id"]

    zapisane = json.loads(db.dokument(stary["id"])["dane_json"])
    assert zapisane["warianty"][kategoria] == pierwsza["id"]


def test_skasowana_formatka_nie_wywala_generowania(klient):
    """Wariant można skasować także wtedy, gdy zrobiono nim stary operat."""
    kategoria = _kategoria(klient.srodowisko)
    wariant, _ = _wgraj(klient.srodowisko, kategoria)
    assert warianty.usun(wariant["id"]) is True

    odpowiedz = klient.post(f"/generuj/{kategoria}",
                            data={**FORMULARZ, f"wariant__{kategoria}": wariant["id"]},
                            follow_redirects=False)

    assert odpowiedz.status_code == 303
    from docx import Document
    katalog = klient.srodowisko.wyniki / db.dokumenty()[0]["katalog"]
    tresc = "\n".join(p.text for p in Document(katalog / "spis_tresci.docx").paragraphs)
    assert "Standardowa:" in tresc          # cicho wrócił do standardowej


# --- interfejs ---------------------------------------------------------------

def test_tabelka_pokazuje_sie_dopiero_gdy_jest_z_czego_wybierac(klient):
    """Bez własnych formatek formularz ma zostać taki, jaki był."""
    kategoria = _kategoria(klient.srodowisko)

    bez = klient.get(f"/nowy/{kategoria}").text
    assert "<legend>Formatki</legend>" not in bez

    _wgraj(klient.srodowisko, kategoria)
    z_wariantem = klient.get(f"/nowy/{kategoria}").text
    assert "<legend>Formatki</legend>" in z_wariantem
    assert "standardowa" in z_wariantem


def test_ustawienia_pokazuja_wgrane_formatki(klient):
    kategoria = _kategoria(klient.srodowisko)
    wariant, _ = _wgraj(klient.srodowisko, kategoria)

    strona = klient.get("/ustawienia").text

    assert wariant["nazwa"] in strona
    assert szablony.szablon_po_id(kategoria).nazwa_dokumentu in strona


def test_listy_formatek_nazywaja_dokument_a_nie_szablon(klient):
    """Szablon nazywa się „Operat”, ale plik, który z niego powstaje, to spis treści.

    Kafelek na stronie głównej i tytuł formularza mają zostać przy „Operacie” — tam
    zaczyna się cała robota. Na listach formatek chodzi już o konkretny dokument.
    """
    klient.srodowisko.dodaj_szablon(
        "spis_tresci_wzor", ["{{ nr_roboty }}"],
        opis={**OPIS_OPERATU, "nazwa": "Operat", "nazwa_dokumentu": "Spis treści"})
    _wgraj(klient.srodowisko, "spis_tresci_wzor")

    strona_glowna = klient.get("/").text
    formularz = klient.get("/nowy/spis_tresci_wzor").text
    ustawienia = klient.get("/ustawienia").text

    assert "Operat" in strona_glowna                    # kafelek zostaje operatem
    assert "<h1>Operat</h1>" in formularz
    assert "<td>Spis treści</td>" in formularz          # ...ale w tabelce formatek
    assert "Spis treści" in ustawienia
    assert ">Operat</option>" not in ustawienia


def test_wgranie_przez_formularz(klient):
    kategoria = _kategoria(klient.srodowisko)
    zrodlo = zbuduj_docx(klient.srodowisko.katalog / "z_dysku.docx",
                         ["Z dysku: {{ nr_roboty }}"])

    with open(zrodlo, "rb") as plik:
        odpowiedz = klient.post("/ustawienia/formatki",
                                data={"kategoria": kategoria},
                                files={"plik": ("z_dysku.docx", plik)},
                                follow_redirects=False)

    assert odpowiedz.status_code == 303
    assert len(warianty.lista(kategoria)) == 1


def test_wgranie_bez_pliku_nie_wywala_strony(klient):
    _kategoria(klient.srodowisko)
    odpowiedz = klient.post("/ustawienia/formatki", data={"kategoria": "spis_tresci_wzor"},
                            follow_redirects=False)
    assert odpowiedz.status_code == 303
    assert "blad" in odpowiedz.headers["location"]
