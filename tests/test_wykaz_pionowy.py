"""Wariant wykazu zmian danych działki na kartce pionowej.

Buduje go `narzedzia/utworz_wykaz_dzialki_pionowy.py`: bierze wydaną formatkę i podmienia
w niej samą tabelę na osiem kolumn, w których **każda działka zajmuje dwa wiersze**
(„dotychczasowy" nad „nowym"). Testy pilnują tego, co przy takiej przebudowie najłatwiej
zepsuć: żeby plik dało się wgrać jako własną formatkę (te same znaczniki), żeby Word go
otworzył (kolejność w OOXML) i żeby pętla naprawdę dawała parę wierszy na działkę.
"""
from __future__ import annotations

import dataclasses
import importlib.util
import pathlib
import tempfile

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app import generator, szablony, warianty

KORZEN = pathlib.Path(__file__).resolve().parent.parent


def _narzedzie(nazwa: str):
    spec = importlib.util.spec_from_file_location(nazwa, KORZEN / "narzedzia" / f"{nazwa}.py")
    modul = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(modul)
    return modul


@pytest.fixture(scope="module")
def wariant(tmp_path_factory) -> pathlib.Path:
    budowniczy = _narzedzie("utworz_wykaz_dzialki_pionowy")
    return budowniczy.zbuduj(budowniczy.ZRODLO,
                             tmp_path_factory.mktemp("pion") / "wykaz_pionowy.docx")


def test_wariant_da_sie_wgrac_zamiast_wydanej_formatki(wariant):
    """Nowy układ, te same znaczniki — inaczej wgranie kończy się ostrzeżeniem,
    a w gotowym dokumencie zostają puste miejsca."""
    stare = set(szablony._zmienne_szablonu(
        KORZEN / "szablony" / "wykaz_zmian_dzialki_wzor.docx"))

    assert set(szablony._zmienne_szablonu(wariant)) == stare
    assert warianty.nieznane_znaczniki("wykaz_zmian_dzialki_wzor", wariant) == []


def test_kartka_jest_pionowa_a_tabela_miesci_sie_w_marginesach(wariant):
    """Cały sens tego wariantu. Tabela szersza niż tekst wychodziłaby poza margines
    i Word łamałby ją na dwie strony."""
    d = Document(str(wariant))
    sekcja = d.sections[0]

    assert sekcja.orientation == WD_ORIENT.PORTRAIT
    assert sekcja.page_width < sekcja.page_height
    szerokosc_tekstu = int((sekcja.page_width - sekcja.left_margin
                            - sekcja.right_margin) / 914400 * 1440)
    szerokosc_tabeli = sum(int(k.get(qn("w:w")))
                           for k in d.tables[0]._tbl.find(qn("w:tblGrid")))
    assert szerokosc_tabeli <= szerokosc_tekstu


def test_kolumny_nie_sa_wezsze_niz_w_poziomej_formatce(wariant):
    """Węższa kartka, a mimo to więcej miejsca na dane — bo kolumn jest osiem
    zamiast trzynastu. Gdyby któraś zeszła poniżej dzisiejszej, wariant nie miałby sensu."""
    budowniczy = _narzedzie("utworz_wykaz_dzialki_pionowy")
    stara = Document(str(KORZEN / "szablony" / "wykaz_zmian_dzialki_wzor.docx"))
    dotychczas = [int(k.get(qn("w:w"))) for k in stara.tables[0]._tbl.find(qn("w:tblGrid"))]
    # w poziomej: [L.p., numer, pole, OFU, OZU, OZK, użytki, …to samo dla stanu nowego]
    najwezsze = dict(zip(("numer", "pole", "ofu", "ozu", "ozk", "uzytki"), dotychczas[1:7]))

    for nazwa, dawna in najwezsze.items():
        assert budowniczy.KOLUMNY[nazwa] >= dawna, f"kolumna {nazwa} zwęziła się"


def test_kazda_dzialka_dostaje_dwa_wiersze(wariant):
    """Pętla `{%tr for %}` powtarza **wszystko** między znacznikami, więc para wierszy
    powtarza się w całości. Numer porządkowy stoi w komórce scalonej przez oba."""
    formatka = dataclasses.replace(szablony.szablon_po_id("wykaz_zmian_dzialki_wzor"),
                                   plik=wariant)

    d = Document(generator.dopisz_dokument(
        formatka,
        {"nr_roboty": "G.1", "polozenie_obreb_teryt": "247301_1.0112",
         "wykazy_dzialek": [
             {"numer_dotychczas": "1765/311", "numer_nowy": "1765/312"},
             {"numer_nowy": "1765/313"}]},
        pathlib.Path(tempfile.mkdtemp())))

    tabela = d.tables[0]
    assert len(tabela.rows) == 2 + 2 * 2, "dwa wiersze nagłówka i po dwa na działkę"
    stany = [w.cells[1].text.strip() for w in tabela.rows[2:]]
    assert stany == ["Dotychczasowy", "Nowy", "Dotychczasowy", "Nowy"]
    assert [w.cells[2].text.strip() for w in tabela.rows[2:]] == [
        "1765/311", "1765/312", "", "1765/313"]
    # L.p. jest scalone pionowo, więc obie linijki pary pokazują ten sam numer,
    # a numer ma kropkę — tak jak w wykazie budynku
    assert [w.cells[0].text.strip() for w in tabela.rows[2:]] == ["1.", "1.", "2.", "2."]


def test_dzialki_oddziela_grubsza_kreska(wariant):
    """Bez tego dwa wiersze jednej działki zlewają się z następną parą."""
    budowniczy = _narzedzie("utworz_wykaz_dzialki_pionowy")
    d = Document(str(wariant))
    wiersze = d.tables[0]._tbl.findall(qn("w:tr"))
    dotychczasowy, nowy = wiersze[3], wiersze[4]

    def dolna(tr):
        tc = tr.findall(qn("w:tc"))[1]
        brzegi = tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        return int(brzegi.find(qn("w:bottom")).get(qn("w:sz")))

    assert dolna(nowy) == budowniczy.GRUBA_KRESKA
    assert dolna(dotychczasowy) < budowniczy.GRUBA_KRESKA, \
        "kreska wewnątrz pary ma być cieńsza niż ta między działkami"


def test_naglowek_i_kolumna_stanu_sa_pogrubione(wariant):
    """Decyzja brata: wytłuszczony cały nagłówek tabeli oraz „Dotychczasowy" i „Nowy",
    bo to one rozdzielają parę wierszy jednej działki. Dane zostają zwykłe."""
    d = Document(str(wariant))
    tabela = d.tables[0]

    def pogrubione(komorka) -> bool:
        biegi = [b for akapit in komorka.paragraphs for b in akapit.runs if b.text.strip()]
        return bool(biegi) and all(b.bold for b in biegi)

    for wiersz in tabela.rows[:2]:
        for komorka in wiersz.cells:
            if komorka.text.strip():
                assert pogrubione(komorka), \
                    f"nagłówek {komorka.text.strip()!r} bez pogrubienia"

    stany = [w.cells[1] for w in tabela.rows[3:5]]
    assert [k.text.strip() for k in stany] == ["Dotychczasowy", "Nowy"]
    assert all(pogrubione(k) for k in stany)
    # ...a dane w tej samej parze wierszy zostają zwykłe
    assert not pogrubione(tabela.rows[3].cells[2])


def test_plik_otworzy_sie_w_wordzie(wariant):
    """Zła kolejność elementów w OOXML daje plik, który LibreOffice składa bez słowa,
    a Word w ogóle nie otwiera (pułapka 12d). U brata objawia się to brakiem miniatur."""
    ujednolic = _narzedzie("ujednolic_wyglad")

    assert ujednolic._sprawdz_kolejnosc(Document(str(wariant))) == []
