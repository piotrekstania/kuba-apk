"""Skrypt ujednolicający wygląd formatek.

Dwie rzeczy, których pilnują te testy, kosztowały w tym projekcie najwięcej:

1. **Skrypt cofał świadome zmiany brata** — pogrubienia w treści, pogrubienia etykiet,
   wybrany rozmiar nagłówka i wcięcie zrobione ciągiem tabulatorów. Cztery razy.
2. **Zapisał pliki, których Word nie otwiera** (zła kolejność elementów w OOXML),
   a LibreOffice składał je do PDF-a bez mrugnięcia. Kosztowało to całe wydanie.
"""
from __future__ import annotations

import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt

KORZEN = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(KORZEN / "narzedzia"))

import ujednolic_wyglad as ujw  # noqa: E402

PRAWDZIWE = sorted((KORZEN / "szablony").glob("*.docx"))


def _tresc_xml(plik: Path) -> bytes:
    with zipfile.ZipFile(plik) as paczka:
        return paczka.read("word/document.xml")


def _dokument_probny(sciezka: Path) -> Path:
    """Formatka z wyborami autora: tytuł, nagłówek 12 pt, pogrubienia, tabulatory."""
    dokument = Document()

    tytul = dokument.add_paragraph().add_run("SPRAWOZDANIE TECHNICZNE")
    tytul.font.size = Pt(14)
    tytul.bold = True

    naglowek = dokument.add_paragraph().add_run("Spis treści")
    naglowek.font.size = Pt(12)

    # treść z pogrubionym fragmentem w środku — decyzja autora, nie skryptu
    akapit = dokument.add_paragraph()
    zwykly = akapit.add_run("Pomiar wykonano metodą ")
    zwykly.font.size = Pt(10)
    wazny = akapit.add_run("GNSS RTN")
    wazny.font.size = Pt(10)
    wazny.bold = True

    # pogrubiona etykieta bloku „Etykieta:<tab>wartość”
    etykieta = dokument.add_paragraph()
    bieg_etykiety = etykieta.add_run("Jednostka ewidencyjna:\t")
    bieg_etykiety.font.size = Pt(10)
    bieg_etykiety.bold = True
    wartosc = etykieta.add_run("Pszczyna")
    wartosc.font.size = Pt(10)

    druga = dokument.add_paragraph()
    bieg_drugiej = druga.add_run("Obręb:\t")
    bieg_drugiej.font.size = Pt(10)
    bieg_drugiej.bold = True
    druga.add_run("Baczków").font.size = Pt(10)

    # wiersz wsunięty pięcioma tabulatorami — świadome wcięcie, poza blokiem
    tolerancje = dokument.add_paragraph()
    bieg_tolerancji = tolerancje.add_run("\t\t\t\t\t[dl – 0.02 m] / [dh – 0.03 m]")
    bieg_tolerancji.font.size = Pt(10)

    dokument.add_paragraph().add_run("Jakub Stania, nr upr. 23266").font.size = Pt(10)
    sciezka.parent.mkdir(parents=True, exist_ok=True)
    dokument.save(sciezka)
    return sciezka


@pytest.fixture
def probka(tmp_path):
    return _dokument_probny(tmp_path / "probka.docx")


def _akapity(plik: Path):
    return [p for p in Document(plik).paragraphs]


def _po_tekscie(plik: Path, fragment: str):
    return next(p for p in Document(plik).paragraphs if fragment in p.text)


# --- czego skrypt nie ma prawa ruszać ----------------------------------------

def test_pogrubienie_w_tresci_zostaje(probka):
    ujw.ujednolic(probka)
    akapit = _po_tekscie(probka, "Pomiar wykonano")
    pogrubione = [b.text for b in akapit.runs if b.bold]
    assert pogrubione == ["GNSS RTN"], "skrypt skasował pogrubienie dołożone w Wordzie"


def test_pogrubiona_etykieta_zostaje(probka):
    ujw.ujednolic(probka)
    akapit = _po_tekscie(probka, "Jednostka ewidencyjna:")
    assert akapit.runs[0].bold, "skrypt zdjął pogrubienie z etykiety"


def test_rozmiar_naglowka_zostaje(probka):
    """Brat wybrał 12 pt dla nagłówka — skrypt ma to uszanować, nie „poprawić”."""
    ujw.ujednolic(probka)
    naglowek = _po_tekscie(probka, "Spis treści")
    assert naglowek.runs[0].font.size == Pt(12)


def test_tabulatory_poza_blokiem_zostaja_nietkniete(probka):
    """Ciąg tabulatorów to świadome wcięcie, nie niechlujstwo do zwinięcia."""
    przed = _po_tekscie(probka, "[dl – 0.02 m]").text.count("\t")
    ujw.ujednolic(probka)
    po = _po_tekscie(probka, "[dl – 0.02 m]").text.count("\t")
    assert po == przed == 5


def test_liczba_akapitow_sie_nie_zmienia(probka):
    """Skrypt formatuje, a nie przepisuje treści."""
    przed = [p.text for p in _akapity(probka)]
    ujw.ujednolic(probka)
    assert [p.text for p in _akapity(probka)] == przed


# --- powtarzalność -----------------------------------------------------------

def test_drugie_uruchomienie_nic_nie_zmienia(probka):
    """Idempotencja. Bez niej skrypt przy każdym przebiegu dokłada kolejny przystanek."""
    ujw.ujednolic(probka)
    po_pierwszym = _tresc_xml(probka)
    ujw.ujednolic(probka)
    assert _tresc_xml(probka) == po_pierwszym


JEST_CZCIONKA = ujw._plik_kroju(True) is not None
BEZ_CZCIONKI = pytest.mark.skipif(
    not JEST_CZCIONKA,
    reason="brak Calibri/Carlito — pomiar etykiet leci wtedy z oszacowania, "
           "więc przystanki wychodzą inne niż na maszynie, gdzie wydano formatki "
           "(sudo apt install fonts-crosextra-carlito)")


@BEZ_CZCIONKI
@pytest.mark.parametrize("plik", PRAWDZIWE, ids=lambda p: p.stem)
def test_prawdziwa_formatka_jest_juz_ustabilizowana(plik, tmp_path):
    """Puszczenie skryptu na wydaną formatkę nie może już nic w niej zmienić.

    Gdy ten test czerwienieje, znaczy to, że w repozytorium leży formatka, która
    nie przeszła `ujednolic_wyglad.py` — albo że skrypt zmienił zdanie i przy
    najbliższym uruchomieniu przestawi bratu dokument.
    """
    kopia = tmp_path / plik.name
    kopia.write_bytes(plik.read_bytes())
    przed = _tresc_xml(kopia)
    ujw.ujednolic(kopia, ujw._stopka_wzorcowa())
    assert _tresc_xml(kopia) == przed


def test_bez_czcionki_skrypt_nie_zgaduje_szerokosci(probka, monkeypatch):
    """Bez pliku Calibri/Carlito skrypt ma odmówić, a nie szacować „na oko”.

    Pomiar najdłuższej etykiety decyduje o tym, gdzie stanie kolumna wartości.
    Gdy czcionki nie ma, `_szerokosc` liczy z liczby znaków i wychodzi inna liczba
    niż na maszynie, na której formatki wydano — czyli puszczenie skryptu na
    komputerze bez czcionki **przestawia kolumny w całym operacie**, nic o tym
    nie mówiąc. Sprawdzone na wykazach zmian: przystanek 3781 → 4371 twipów.
    """
    monkeypatch.setattr(ujw, "_plik_kroju", lambda pogrubiony=False: None)
    with pytest.raises(SystemExit):
        ujw.ujednolic(probka)


# --- pliki, których Word nie otworzy -----------------------------------------

@pytest.mark.parametrize("plik", PRAWDZIWE, ids=lambda p: p.stem)
def test_prawdziwa_formatka_ma_poprawna_kolejnosc_xml(plik):
    """Wydane formatki muszą przechodzić kontrolę schematu OOXML.

    LibreOffice składa do PDF-a także pliki niepoprawne — u brata objawiło się to
    tym, że Word nie otworzył dokumentu i przestały powstawać miniatury.
    """
    assert ujw._sprawdz_kolejnosc(Document(plik)) == []


def test_wykrywa_rozstrzelenie_za_rozmiarem(tmp_path):
    """<w:spacing> po <w:sz> — plik otwiera się na Linuksie, w Wordzie nie."""
    sciezka = _dokument_probny(tmp_path / "zly.docx")
    dokument = Document(sciezka)
    bieg = dokument.paragraphs[0].runs[0]
    rPr = bieg._element.get_or_add_rPr()
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")
    rPr.append(sz)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "12")
    rPr.append(spacing)                       # na końcu, czyli za <w:sz>

    zarzuty = ujw._sprawdz_kolejnosc(dokument)
    assert any("spacing" in z for z in zarzuty), zarzuty


def test_wykrywa_oblewanie_bez_wraptext(tmp_path):
    """<wp:wrapSquare> bez atrybutu wrapText to plik nie do otwarcia."""
    dokument = Document()
    akapit = dokument.add_paragraph()
    kotwica = OxmlElement("wp:anchor")
    kotwica.append(OxmlElement("wp:wrapSquare"))      # bez wrapText
    kotwica.append(OxmlElement("wp:docPr"))
    akapit._p.append(kotwica)

    zarzuty = ujw._sprawdz_kolejnosc(dokument)
    assert any("wrapText" in z for z in zarzuty), zarzuty


def test_wykrywa_oblewanie_za_docpr(tmp_path):
    dokument = Document()
    akapit = dokument.add_paragraph()
    kotwica = OxmlElement("wp:anchor")
    kotwica.append(OxmlElement("wp:docPr"))
    oblewanie = OxmlElement("wp:wrapSquare")
    oblewanie.set("wrapText", "bothSides")
    kotwica.append(oblewanie)                          # za <wp:docPr>
    akapit._p.append(kotwica)

    zarzuty = ujw._sprawdz_kolejnosc(dokument)
    assert any("wrapSquare" in z and "docPr" in z for z in zarzuty), zarzuty


def test_wykrywa_ext_rozszerzenia_z_rozmiarem(tmp_path):
    """<a:ext uri=…> to pozycja listy rozszerzeń, a nie rozmiar obrazka."""
    dokument = Document()
    akapit = dokument.add_paragraph()
    ext = OxmlElement("a:ext")
    ext.set("uri", "{28A0092B-C50C-407E-A947-70E740481C1C}")
    ext.set("cx", "1000")
    akapit._p.append(ext)

    zarzuty = ujw._sprawdz_kolejnosc(dokument)
    assert any("ext" in z for z in zarzuty), zarzuty
