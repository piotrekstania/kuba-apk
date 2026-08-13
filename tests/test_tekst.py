"""Tekst z formatowaniem: pogrubienie, kursywa, podkreślenie.

Opis przebiegu prac brat pisał wcześniej w Wordzie i chce zachować pogrubienia, więc
program trzyma go jako fragment HTML. Dwie rzeczy muszą tu być pewne:

* **do dokumentu wchodzi tylko to, co umiemy oddać** — pogrubienie, kursywa,
  podkreślenie i złamanie wiersza. Reszta (kolory, czcionki wklejone z Worda, skrypty)
  leci do kosza, zanim cokolwiek zapiszemy;
* **`{{r pole }}` w formatce wymaga `RichText`** — zwykły napis wjechałby w to miejsce
  bez ucieczek i pierwszy `<` w opisie zrobiłby plik, którego Word nie otworzy.
"""
from __future__ import annotations

import pathlib
import tempfile
import zipfile

import pytest
from docx import Document
from docxtpl import DocxTemplate

from app import tekst


# --- co przepuszczamy --------------------------------------------------------

def test_pogrubienie_kursywa_podkreslenie_zostaja():
    assert tekst.oczysc("<b>a</b><i>b</i><u>c</u>") == "<b>a</b><i>b</i><u>c</u>"


def test_warianty_znacznikow_sprowadzamy_do_jednego():
    """Przeglądarki i Word wstawiają to samo pod różnymi nazwami."""
    assert tekst.oczysc("<strong>a</strong><em>b</em><ins>c</ins>") == "<b>a</b><i>b</i><u>c</u>"


def test_akapity_staja_sie_zlamaniem_wiersza():
    """Formatka ma jeden akapit na opis, więc prawdziwe akapity nie mają gdzie wejść."""
    assert tekst.oczysc("<p>Pierwszy</p><p>Drugi</p>") == "Pierwszy<br>Drugi"


def test_zwykly_tekst_przechodzi_bez_zmian():
    """Opisy zapisane przed wprowadzeniem formatowania mają dalej działać."""
    assert tekst.oczysc("Pomiar wykonano metodą RTN GNSS.") == "Pomiar wykonano metodą RTN GNSS."


# --- czego nie przepuszczamy -------------------------------------------------

def test_skrypt_leci_razem_z_trescia():
    """Sam znacznik zdjęty to za mało — „alert(1)” w opisie wygląda jak usterka."""
    assert tekst.oczysc("<script>alert(1)</script>Opis") == "Opis"


def test_atrybuty_znikaja():
    """`onclick` w treści, która wraca na stronę jako HTML, byłby dziurą."""
    assert tekst.oczysc('<b onclick="zle()" style="color:red">Klik</b>') == "<b>Klik</b>"


def test_kolory_i_czcionki_z_worda_nie_wchodza_do_operatu():
    """Wklejka z Worda ciągnie style, które rozjechałyby wygląd dokumentu."""
    wklejka = '<span style="font-family:Arial;color:#ff0000">Czerwone</span> i <b>grube</b>'

    assert tekst.oczysc(wklejka) == "Czerwone i <b>grube</b>"


def test_niedomkniety_znacznik_zostaje_domkniety():
    """Otwarte pogrubienie rozlałoby się na resztę strony."""
    assert tekst.oczysc("<b>Bez końca") == "<b>Bez końca</b>"


def test_nawiasy_trojkatne_z_tekstu_nie_staja_sie_znacznikami():
    assert tekst.oczysc("pole a < b oraz a &amp; b") == "pole a &lt; b oraz a &amp; b"


# --- „czy cokolwiek wpisano” -------------------------------------------------

def test_sam_znacznik_bez_tekstu_liczy_sie_jako_pusty():
    """Edytor zostawia po sobie puste znaczniki — to nadal jest pusty opis, więc
    w sprawozdaniu ma wyjść „brak”, a nie pusta dziura."""
    assert tekst.na_zwykly_tekst("<b></b><br>") == ""
    assert tekst.na_zwykly_tekst("<b>Coś</b>") == "Coś"


# --- droga do Worda ----------------------------------------------------------

def _wypelnij(znacznik: str, wartosc) -> pathlib.Path:
    kat = pathlib.Path(tempfile.mkdtemp())
    dokument = Document()
    dokument.add_paragraph(f"Przebieg: {znacznik}")
    wzor = kat / "wzor.docx"
    dokument.save(wzor)

    szablon = DocxTemplate(wzor)
    szablon.render({"opis": wartosc}, autoescape=True)
    wynik = kat / "wynik.docx"
    szablon.save(wynik)
    return wynik


def test_formatowanie_dojezdza_do_dokumentu_jako_biegi_tekstu():
    plik = _wypelnij("{{r opis }}", tekst.na_richtext(
        "<b>Pomiar</b> zwykły <i>skos</i> <u>kreska</u>"))

    akapit = Document(plik).paragraphs[0]
    style = {b.text: (b.bold, b.italic, b.underline) for b in akapit.runs if b.text}

    assert style["Pomiar"] == (True, None, None)
    assert style["skos"] == (None, True, None)
    assert style["kreska"] == (None, None, True)
    assert style[" zwykły "] == (None, None, None)


def test_zlamanie_wiersza_dojezdza_do_dokumentu():
    plik = _wypelnij("{{r opis }}", tekst.na_richtext("Pierwszy<br>Drugi"))

    assert zipfile.ZipFile(plik).read("word/document.xml").decode().count("<w:br/>") == 1


def test_richtext_nie_wstawia_biegu_w_srodek_tekstu():
    """To jest ten błąd, o który tu chodzi.

    Przy zwykłym `{{ pole }}` docxtpl wstawia `<w:r>` w środek `<w:t>` — Word takiego
    pliku **nie otworzy**, a LibreOffice łyka go bez słowa, więc zielony PDF u autora
    niczego nie dowodzi (to samo ugryzło już raz przy `ujednolic_wyglad.py`).
    Dlatego formatka ma `{{r opis_przebiegu }}`.
    """
    import re

    dobry = zipfile.ZipFile(_wypelnij("{{r opis }}", tekst.na_richtext("<b>a</b>b"))
                            ).read("word/document.xml").decode()
    assert re.search(r"<w:t[^>]*>[^<]*<w:r", dobry) is None

    zly = zipfile.ZipFile(_wypelnij("{{ opis }}", tekst.na_richtext("<b>a</b>b"))
                          ).read("word/document.xml").decode()
    assert re.search(r"<w:t[^>]*>[^<]*<w:r", zly) is not None, (
        "gdyby docxtpl przestał tak robić, ten test straciłby sens — sprawdź, czy "
        "formatka nadal musi mieć `{{r }}`")


@pytest.mark.parametrize("pusty", ["", "<b></b>", "<br>"])
def test_pusty_opis_nie_wywala_wypelniania(pusty):
    plik = _wypelnij("{{r opis }}", tekst.na_richtext(tekst.oczysc(pusty)))

    assert Document(plik).paragraphs[0].text.strip() == "Przebieg:"


# --- wklejka z Worda ---------------------------------------------------------
#
# Sedno tej obsługi: brat wkleja opisy pisane wcześniej w Wordzie. Word zapisuje
# pogrubienie najczęściej **nie** jako `<b>`, tylko jako `style="font-weight:700"` —
# bez czytania stylów wklejone formatowanie ginęłoby w całości.

WKLEJKA_Z_WORDA = (
    '<span style="font-family:Calibri;font-size:11.0pt">Pomiar </span>'
    '<span style="font-weight:700;font-family:Calibri">RTN GNSS</span>'
    '<span style="font-style:italic">, tachimetrycznie</span>'
    '<span style="text-decoration:underline">, kontrolnie</span>.'
)


def test_pogrubienie_ze_stylu_worda_jest_rozpoznane():
    assert tekst.oczysc(WKLEJKA_Z_WORDA) == (
        "Pomiar <b>RTN GNSS</b><i>, tachimetrycznie</i><u>, kontrolnie</u>.")


def test_krój_i_rozmiar_czcionki_nie_przechodza():
    """To nie jest brak, tylko decyzja: krój i rozmiary ustala formatka, a pilnuje ich
    `ujednolic_wyglad.py` — wklejone „Times New Roman 11 pt” rozjechałoby operat."""
    wynik = tekst.oczysc(WKLEJKA_Z_WORDA)

    assert "font-family" not in wynik and "font-size" not in wynik
    assert "Calibri" not in wynik


def test_kolory_tez_zostaja_za_progiem():
    assert tekst.oczysc('<span style="color:#ff0000;background:yellow">Czerwone</span>') \
        == "Czerwone"


def test_przekreslenie_i_indeksy_przechodza():
    """`RichText` umie je oddać wprost, więc nie ma powodu ich odrzucać."""
    assert tekst.oczysc("<del>skreślone</del> m<sup>2</sup> H<sub>2</sub>O") == \
        "<s>skreślone</s> m<sup>2</sup> H<sub>2</sub>O"


def test_sasiednie_kawalki_o_tym_samym_stylu_nie_mnoza_znacznikow():
    """Word tnie tekst na kawałki po swojemu — wynik ma być czytelny, a nie posiekany."""
    assert tekst.oczysc("<b>Ala</b><b> ma</b> kota") == "<b>Ala ma</b> kota"


def test_przekreslenie_i_indeksy_dojezdzaja_do_dokumentu():
    plik = _wypelnij("{{r opis }}", tekst.na_richtext(
        tekst.oczysc("<del>skreślone</del> m<sup>2</sup>")))

    biegi = {b.text: b for b in Document(plik).paragraphs[0].runs if b.text}
    assert biegi["skreślone"].font.strike is True
    assert biegi["2"].font.superscript is True


def test_bieg_nie_narzuca_kroju_ani_rozmiaru():
    """Bieg bez własnych ustawień dziedziczy je z akapitu formatki — czyli wygląda
    jak reszta dokumentu, o co w tym wszystkim chodzi."""
    plik = _wypelnij("{{r opis }}", tekst.na_richtext("<b>Pomiar</b>"))

    bieg = next(b for b in Document(plik).paragraphs[0].runs if b.text == "Pomiar")
    assert bieg.font.size is None
    assert bieg.font.name is None


# --- odstępy: to była regresja, którą brat zgłosił ----------------------------

def test_pusta_linia_przezywa():
    """Dwa `<br>` obok siebie to świadomy odstęp między akapitami opisu.

    Pierwsza wersja zwijała je do jednego „żeby nie dublować złamań” i opis zlepiał się
    po zapisaniu w jeden blok — dokładnie ta usterka, którą brat zobaczył w sprawozdaniu.
    """
    assert tekst.oczysc("Pierwsza<br><br>Trzecia") == "Pierwsza<br><br>Trzecia"


def test_pusta_linia_z_edytora_tez_przezywa():
    """Przeglądarka zapisuje pustą linię jako `<div><br></div>` — to ten sam odstęp."""
    assert tekst.oczysc("<div>Pierwsza</div><div><br></div><div>Trzecia</div>") == \
        "Pierwsza<br><br>Trzecia"


def test_granica_blokow_sie_nie_mnozy():
    """`</p><p>` to jedna granica opisana dwoma znacznikami, a nie dwie puste linie."""
    assert tekst.oczysc("<p>a</p><p>b</p>") == "a<br>b"


def test_tabulatory_zostaja():
    """Brat wyrównuje nimi wiersz z tolerancjami pod kolumnę wartości."""
    assert tekst.oczysc("Etykieta:\t\twartość") == "Etykieta:\t\twartość"


def test_tabulatory_i_puste_linie_dojezdzaja_do_dokumentu():
    """Cała ta droga naraz — tak wygląda opis, który brat naprawdę wpisuje."""
    stary = ("Pomiar kontrolny:\t1 – istniejąca studnia\n"
             "\t\t\t[dl – 0.02 m]\n"
             "\n"
             "Wszystkie punkty spełniają dokładność.")
    xml = zipfile.ZipFile(_wypelnij("{{r opis }}", tekst.na_richtext(tekst.oczysc(stary)))
                          ).read("word/document.xml").decode()

    assert xml.count("<w:tab/>") == 4
    assert xml.count("<w:br/>") == 3, "dwa złamania i pusta linia między akapitami"
